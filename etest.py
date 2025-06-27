# coding utf-8
import base64
import datetime
import json
import os
import re
import time

import folium
import openpyxl
import pandas as pd
import plotly.graph_objects as go
import pydeck as pdk
import pymysql
import streamlit as st
import streamlit.components.v1 as components
import streamlit_antd_components as sac
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor
from folium.plugins import HeatMap, MiniMap
from PIL import Image, ImageDraw, ImageFont
from st_keyup import st_keyup
from streamlit_extras.badges import badge
from streamlit_extras.metric_cards import style_metric_cards
from streamlit_folium import st_folium
from streamlit_javascript import st_javascript
from streamlit_timeline import st_timeline
from xlsxwriter.workbook import Workbook

from commFunc import (GenerExam, deepseek_AI, deepseek_AI_GenerQues,
                      execute_sql, execute_sql_and_commit, getParam,
                      getUserEDKeys, qianfan_AI, qianfan_AI_GenerQues,
                      updateActionUser, updatePyFileinfo, xunfei_xh_AI,
                      xunfei_xh_AI_fib, xunfei_xh_AI_GenerQues)
from word2picture import tywx_generate_image, xfxh_generate_image

# cSpell:ignoreRegExp /[^\s]{16,}/
# cSpell:ignoreRegExp /\b[A-Z]{3,15}\b/g
# cSpell:ignoreRegExp /\b[A-Z]\b/g


@st.fragment
def updateKeyAction(keyAction):
    # 构建查询SQL语句
    sql = f"SELECT ID from keyactionlog where userName = {st.session_state.userName} and userCName = '{st.session_state.userCName}' and userAction = '{keyAction}' and actionDate = {int(time.time())}"
    # 执行SQL查询
    if not execute_sql(cur, sql):
        # 如果查询结果为空，则执行插入操作
        # 构建插入SQL语句
        sql = f"INSERT INTO keyactionlog(userName, userCName, StationCN, userAction, actionDate) VALUES({st.session_state.userName}, '{st.session_state.userCName}', '{st.session_state.StationCN}', '{keyAction}', {int(time.time())})"
        # 执行SQL插入并提交事务
        execute_sql_and_commit(conn, cur, sql)


# noinspection PyShadowingNames
@st.fragment
def getUserCName(sUserName, sType="Digit"):
    errorInfo = ""

    # 判断sType是否为"Digit"
    if sType.capitalize() == "Digit":
        # 使用正则表达式去除非数字和小数点字符
        cop = re.compile('[^0-9^.]')
        inputStr = cop.sub('', sUserName)
        # 如果原字符串长度与过滤后的字符串长度相等，说明原字符串只包含数字和小数点
        if len(sUserName) == len(inputStr):
            sql = f"SELECT userCName, StationCN from users where userName = {sUserName}"
        else:
            sql = ""
            errorInfo = "请输入纯数字用户编码"

    # 判断sType是否为"Str"
    elif sType.capitalize() == "Str":
        sql = f"SELECT userCName, StationCN from users where userCName = '{sUserName}'"

    # 其他情况
    else:
        sql = ""

    # 如果sql不为空
    if sql != "":
        rows = execute_sql(cur, sql)
        if rows:
            st.session_state.userCName = rows[0][0]
            st.session_state.StationCN = rows[0][1]
        else:
            st.session_state.userCName = "未找到"
            st.session_state.StationCN = "未找到"

    # 如果sql为空
    else:
        if errorInfo != "":
            st.error(errorInfo)
        st.session_state.userCName = ""
        st.session_state.StationCN = ""


def is_valid_table_name(table_name):
    """简单校验表名是否为合法的数据库标识符"""
    import re
    return bool(re.match(r'^[\w\d_]+$', table_name))


def delOutdatedTable():
    tables_to_drop = []
    if st.session_state.examRandom and "examTable" in st.session_state:
        tables_to_drop.append(st.session_state.examTable)
    if "examFinalTable" in st.session_state:
        tables_to_drop.append(st.session_state.examFinalTable)
    for table_name in tables_to_drop:
        if is_valid_table_name(table_name):
            execute_sql_and_commit(conn, cur, sql=f"DROP TABLE IF EXISTS {table_name}")
        else:
            raise ValueError(f"Invalid table name detected: {table_name}")


# noinspection PyShadowingNames
def changePassword():
    # 显示密码修改页面标题
    st.write("### :red[密码修改]")
    # 创建一个带有边框的容器
    changePW = st.empty()
    with changePW.container(border=True):
        # 输入原密码
        oldPassword = st.text_input("请输入原密码", max_chars=8, type="password", autocomplete="off")
        # 输入新密码
        newPassword = st.text_input("请输入新密码", max_chars=8, type="password", autocomplete="off")
        # 再次输入新密码以确认
        confirmPassword = st.text_input("请再次输入新密码", max_chars=8, placeholder="请与上一步输入的密码一致", type="password", autocomplete="new-password")
        # 确认修改按钮
        buttonSubmit = st.button("确认修改")

    # 检查原密码是否为空
    if oldPassword:
        # 验证用户原密码
        verifyUPW = verifyUserPW(st.session_state.userName, oldPassword)
        if verifyUPW[0]:
            oldPassword = verifyUPW[1]
        # 构造SQL查询语句，验证用户名和密码是否匹配
        sql = f"SELECT ID from users where userName = {st.session_state.userName} and userPassword = '{oldPassword}'"
        if execute_sql(cur, sql):
            # 检查新密码和确认密码是否填写且一致
            if newPassword and confirmPassword and newPassword != "":
                if newPassword == confirmPassword:
                    # 确认修改按钮是否被点击
                    if buttonSubmit:
                        # 加密新密码
                        newPassword = getUserEDKeys(newPassword, "enc")
                        # 构造SQL更新语句，更新用户密码
                        sql = f"UPDATE users set userPassword = '{newPassword}' where userName = {st.session_state.userName}"
                        # 执行SQL语句并提交
                        execute_sql_and_commit(conn, cur, sql)
                        # 记录用户密码修改操作
                        updateKeyAction("用户密码修改")
                        # 显示密码修改成功提示，并要求重新登录
                        st.toast("密码修改成功, 请重新登录")
                        # 登出用户
                        logout()
                else:
                    # 显示密码不一致的错误信息
                    st.error("两次输入的密码不一致")
            else:
                # 显示新密码未填写的警告信息
                st.warning("请检查新密码")
        else:
            # 显示原密码错误的错误信息
            st.error("原密码不正确")
    else:
        st.warning("原密码不能为空")

    # 记录用户密码修改操作及时间
    updateActionUser(st.session_state.userName, "密码修改", st.session_state.loginTime)


# noinspection PyShadowingNames
@st.cache_data
def get_userName(searchUserName=""):
    searchUserNameInfo = ""
    if len(searchUserName) > 1:
        sql = f"SELECT userName, userCName, StationCN from users where userName like '{searchUserName}%'"
        rows = execute_sql(cur, sql)
        for row in rows:
            searchUserNameInfo += f"用户编码: :red[{row[0]}] 姓名: :blue[{row[1]}] 站室: :orange[{row[2]}]\n\n"
    if searchUserNameInfo != "":
        searchUserNameInfo += "\n请在用户编码栏中填写查询出的完整编码"
    return searchUserNameInfo


@st.cache_data
def get_userCName(searchUserCName=""):
    searchUserCNameInfo = ""
    if len(searchUserCName) > 1:
        sql = f"SELECT userName, userCName, StationCN from users where userCName like '{searchUserCName}%'"
        rows = execute_sql(cur, sql)
        for row in rows:
            searchUserCNameInfo += f"用户编码: :red[{row[0]}] 姓名: :blue[{row[1]}] 站室: :orange[{row[2]}]\n\n"
    else:
        searchUserCNameInfo = ":red[**请输入至少2个字**]"
    if searchUserCNameInfo != "" and "请输入至少2个字" not in searchUserCNameInfo:
        searchUserCNameInfo += "\n请在用户编码栏中填写查询出的完整编码"

    return searchUserCNameInfo


@st.fragment
def login():
    # 显示应用名称
    #st.write("## :blue[专业技能考试系统 - 离线版]")
    st.markdown(f"<font face='微软雅黑' color=purple size=20><center>**{APPNAME}**</center></font>", unsafe_allow_html=True)

    # 登录表单容器
    login = st.empty()
    with login.container(border=True):
        # 用户编码输入框
        userName = st_keyup("请输入用户编码", placeholder="请输入纯数字用户编码, 非站室名称, 如果不知编码, 请在下方输入姓名查询", max_chars=8)
        # 初始化用户姓名
        st.session_state.userCName = ""

        # 如果输入了用户编码
        if userName:
            filtered = get_userName(userName)
            # 如果未找到对应的用户
            if filtered == "":
                # 根据用户编码获取用户姓名和站室
                getUserCName(userName, "Digit")
                # 显示用户姓名和站室
                st.caption(f"用户名: :blue[{st.session_state.userCName}] 站室: :orange[{st.session_state.StationCN}]")
        else:
            filtered = ""

        # 如果用户姓名未找到或存在过滤结果
        if st.session_state.userCName == "未找到" or filtered:
            st.caption(filtered)

        # 如果用户编码为空或用户姓名未找到
        if userName == "" or st.session_state.userCName == "未找到":
            # 用户姓名输入框
            userCName = st_keyup("请输入用户姓名", placeholder="请输入用户姓名, 至少2个字, 用于查询, 非必填项", max_chars=8)
            st.session_state.userCName = ""

            # 如果输入了用户姓名
            if userCName:
                filtered = get_userCName(userCName)
                # 如果未找到对应的用户
                if filtered == "":
                    # 根据用户姓名获取用户姓名和站室
                    getUserCName(userCName, "Str")
                    # 显示用户姓名和站室
                    st.caption(f"用户名: :blue[{st.session_state.userCName}] 站室: :orange[{st.session_state.StationCN}]")
            else:
                filtered = ""

            # 如果用户姓名未找到或存在过滤结果
            if st.session_state.userCName == "未找到" or filtered:
                # 提示区域容器
                promptArea = st.empty()
                with promptArea.container():
                    # 显示过滤结果
                    st.caption(filtered)
                # 如果用户编码存在但过滤结果为空
                if userName and filtered == "":
                    promptArea.empty()

        # 用户密码输入框
        userPassword = st.text_input("请输入密码", max_chars=8, placeholder="用户初始密码为1234", type="password", autocomplete="off")

        # 模式选择
        examType = sac.segmented(
            items=[
                sac.SegmentedItem(label="练习", icon="pen"),
                sac.SegmentedItem(label="考试", icon="card-list"),
            ], align="start"
        )
        #examType = st.selectbox("请选择模式类型", ("练习", "考试"), index=0, help="各站管理员如需修改设置及查询请选择练习模式, 考试模式仅限考试")

        # 登录按钮
        buttonLogin = st.button("登录")

    # 如果点击了登录按钮
    if buttonLogin:
        # 如果用户编码和密码不为空
        if userName != "" and userPassword != "":
            # 验证用户密码
            verifyUPW = verifyUserPW(userName, userPassword)
            # 如果密码验证成功
            if verifyUPW[0]:
                userPassword = verifyUPW[1]

            # 根据选择的模式类型执行不同的逻辑
            if examType == "练习":
                st.session_state.examType = "training"
                st.session_state.examName = "练习题库"
                sql = f"SELECT userName, userCName, userType, StationCN from users where userName = {userName} and userPassword = '{userPassword}'"
            elif examType == "考试":
                st.session_state.examType = "exam"
                sql = f"SELECT userName, userCName, userType, StationCN from users where userName = {userName} and userPassword = '{userPassword}' and activeUser = 0"
            else:
                sql = ""

            # 如果SQL语句不为空
            if sql != "":
                result = execute_sql(cur, sql)
                # 如果查询结果存在
                if result:
                    st.toast(f"用户: {result[0][0]} 姓名: {result[0][1]} 登录成功, 欢迎回来")
                    login.empty()
                    st.session_state.logged_in = True
                    st.session_state.userPwRecheck = False
                    st.session_state.userName = result[0][0]
                    st.session_state.userCName = result[0][1].replace(" ", "")
                    st.session_state.userType = result[0][2]
                    st.session_state.StationCN = result[0][3]
                    st.session_state.examLimit = getParam("同场考试次数限制", st.session_state.StationCN)
                    st.session_state.debug = bool(getParam("测试模式", st.session_state.StationCN))
                    st.session_state.clockType = bool(getParam("时钟样式", st.session_state.StationCN))
                    st.session_state.curQues = 0
                    st.session_state.examChosen = False
                    st.session_state.delExam = True
                    st.session_state.tooltipColor = "#ed872d"
                    st.session_state.loginTime = int(time.time())
                    if examType == "练习":
                        st.session_state.examRandom = True
                    elif examType == "考试":
                        st.session_state.examRandom = bool(getParam("考试题库每次随机生成", st.session_state.StationCN))
                    sql = f"UPDATE users set activeUser = 1, loginTime = {st.session_state.loginTime}, activeTime_session = 0, actionUser = '空闲' where userName = {st.session_state.userName}"
                    execute_sql_and_commit(conn, cur, sql)
                    #sql = "UPDATE verinfo set pyLM = pyLM + 1 where pyFile = 'visitcounter'"
                    #execute_sql_and_commit(conn, cur, sql)
                    ClearTables()
                    # transform Key to Encrypt(temporary)
                    #print(getUserEDKeys("", "enc"))
                    st.rerun()
                else:
                    # 如果密码验证成功但登录失败
                    if verifyUPW[0]:
                        st.error("登录失败, 用户已经在别处登录, 请联系管理员解决")
                    else:
                        st.error("登录失败, 请检查用户名和密码, 若忘记密码请联系管理员重置")
        else:
            # 如果用户编码或密码为空
            st.warning("请输入用户编码和密码")


def logout():
    try:
        # 更新用户状态为未激活，并更新会话时间
        sql = f"UPDATE users set activeUser = 0, activeTime = activeTime + activeTime_session, activeTime_session = 0 where userName = {st.session_state.userName}"
        execute_sql_and_commit(conn, cur, sql)
        # 删除过时的表
        delOutdatedTable()

    finally:
        # 关闭游标
        cur.close()
        # 关闭数据库连接
        conn.close()

    # 清除会话状态中的所有键值对
    for key in st.session_state.keys():
        del st.session_state[key]

    # 重新运行当前脚本
    st.rerun()


def aboutInfo():
    st.subheader("关于本软件", divider="rainbow")
    st.subheader(":blue[Powered by Python and Streamlit]")
    logo1, logo2, logo3, logo4, logo5, logo6 = st.columns(6)
    logo7, logo8, logo9, logo10, logo11, logo12 = st.columns(6)
    with logo1:
        st.caption("Python")
        st.image("./Images/logos/python.png")
    with logo2:
        st.caption("Streamlit")
        st.image("./Images/logos/streamlit.png")
    with logo3:
        st.caption("SQLite")
        st.image("./Images/logos/sqlite.png")
    with logo4:
        st.caption("Pandas")
        st.image("./Images/logos/pandas.png")
    with logo5:
        st.caption("Ant Comp")
        st.image("./Images/logos/antd.png")
    with logo7:
        st.caption("iFlytek Spark")
        st.image("./Images/logos/xfxh.png")
    with logo8:
        st.caption("ERNIE Qianfan")
        st.image("./Images/logos/qianfan.png")
    with logo9:
        st.caption("DeepSeek")
        st.image("./Images/logos/deepseek.png")
    with logo10:
        st.caption("通义万相")
        st.image("./Images/logos/tywx.png")
    display_pypi()
    st.write("###### :violet[为了获得更好的使用体验, 请使用浅色主题]")
    verinfo, verLM, likeCM = getVerInfo()
    st.caption(f"Version: {int(verinfo / 10000)}.{int((verinfo % 10000) / 100)}.{int(verinfo / 10)} building {verinfo} Last Modified: {time.strftime('%Y-%m-%d %H:%M', time.localtime(verLM))}")
    st.caption(f"Reviews: {EMOJI[int(likeCM) - 1][0]} {likeCM} :orange[I feel {EMOJI[int(likeCM) - 1][1]}]")
    sac.divider(align="center", color="blue")
    stars = sac.rate(label='Please give me a star if you like it!', align='start')
    if stars > 0:
        st.write(f"I feel {EMOJI[int(stars) - 1][1]} {EMOJI[int(stars) - 1][0]}")
    sql = f"UPDATE verinfo set pyMC = pyMC + 1 where pyFile = 'thumbs-up-stars' and pyLM = {stars}"
    execute_sql_and_commit(conn, cur, sql)
    updateActionUser(st.session_state.userName, "浏览[关于]信息", st.session_state.loginTime)


# noinspection PyBroadException,PyUnusedLocal
def getVerInfo():
    try:
        sql = "SELECT Sum(pyMC) from verinfo"
        verinfo = execute_sql(cur, sql)[0][0]
        sql = "SELECT Max(pyLM) from verinfo"
        verLM = execute_sql(cur, sql)[0][0]
        sql = "SELECT Sum(pyLM * pyMC), Sum(pyMC) from verinfo where pyFile = 'thumbs-up-stars'"
        tmpTable = execute_sql(cur, sql)
        likeCM = round(tmpTable[0][0] / tmpTable[0][1], 1)

        return verinfo, verLM, likeCM
    except Exception as e:
        return 0, 0, 0


def display_pypi():
    pypi1, pypi2, pypi3, pypi4, pypi5, pypi6 = st.columns(6)
    with pypi1:
        badge(type="pypi", name="streamlit")
    with pypi2:
        badge(type="pypi", name="pandas")
    with pypi3:
        badge(type="pypi", name="streamlit_antd_components")
    with pypi4:
        badge(type="pypi", name="folium")
    with pypi5:
        badge(type="pypi", name="qianfan")


def aboutLicense():
    st.subheader("License", divider="green")
    st.markdown(open("./LICENSE", "r", encoding="utf-8").read())
    updateActionUser(st.session_state.userName, "浏览License信息", st.session_state.loginTime)


def actDelTable():
    for each in st.session_state.keys():
        if each.startswith("delStaticExamTable_"):
            if st.session_state[each]:
                each = each.replace("delStaticExamTable_", "")
                execute_sql_and_commit(conn, cur, sql=f"DROP TABLE IF EXISTS {each}")
                st.info(f"{each} 静态题库删除成功")


def delStaticExamTable():
    flagExistTable = False
    sql = "SELECT name from sqlite_master where type = 'table' and name like 'exam_%'"
    tempTable = execute_sql(cur, sql)
    if tempTable:
        st.subheader("删除静态题库", divider="red")
        for row in tempTable:
            if row[0].count("_") == 2:
                st.checkbox(f"{row[0]}", key=f"delStaticExamTable_{row[0]}")
                flagExistTable = True
    if flagExistTable:
        st.button("确认删除", on_click=actDelTable)
    else:
        st.info("暂无静态题库")


def resultExcel():
    # 设置子标题，并添加蓝色分割线
    st.subheader("试卷导出", divider="blue")
    # 初始化两个空列表，用于存储试卷结果
    examResultPack, examResultPack2 = [], []
    # 构造SQL查询语句，查询所有以"exam_final_"开头的表名
    sql = "SELECT name from sqlite_master where type = 'table' and name like 'exam_final_%'"
    # 执行SQL查询，获取结果
    tempTable = execute_sql(cur, sql)
    # 如果查询结果不为空
    if tempTable:
        for row in tempTable:
            # 将表名添加到examResultPack2列表中
            examResultPack2.append(row[0])
            # 提取表名中的用户信息部分
            tmp = row[0][:row[0].rfind("_")]
            tmp = tmp[tmp.rfind("_") + 1:]
            # 构造SQL查询语句，查询用户名
            sql = "SELECT userCName from users where userName = " + str(tmp)
            # 执行SQL查询，获取结果
            tempTable = execute_sql(cur, sql)
            # 如果查询结果不为空
            if tempTable:
                # 获取用户名，并替换表名中的用户信息部分，然后添加到examResultPack列表中
                tempUserCName = tempTable[0][0]
                examResultPack.append(row[0].replace("exam_final_", "").replace(tmp, tempUserCName))
            else:
                # 如果查询结果为空，则直接替换表名中的"exam_final_"部分，然后添加到examResultPack列表中
                examResultPack.append(row[0].replace("exam_final_", ""))
        # 使用st.selectbox创建一个下拉选择框，让用户选择试卷
        examResult = st.selectbox(" ", examResultPack, index=None, label_visibility="collapsed")

        # 如果用户选择了试卷
        if examResult:
            for index, value in enumerate(examResultPack):
                # 如果当前值等于用户选择的试卷名
                if value == examResult:
                    # 获取对应的表名，并跳出循环
                    examResult = examResultPack2[index]
                    break
            # 构造SQL查询语句，查询试卷内容
            sql = f"SELECT Question, qOption, qAnswer, qType, qAnalysis, userAnswer from {examResult} order by ID"
            # 执行SQL查询，获取结果
            rows = execute_sql(cur, sql)
            # 如果查询结果不为空
            if rows:
                # 将查询结果转换为DataFrame
                df = pd.DataFrame(rows)
                # 重命名列名
                df.columns = ["题目", "选项", "标准答案", "类型", "解析", "你的答案"]
                # 使用st.dataframe显示DataFrame
                st.dataframe(df)
    else:
        # 如果查询结果为空，则显示提示信息
        st.info("暂无试卷")


def examResulttoExcel():
    # 显示子标题
    st.subheader("考试成绩导出", divider="blue")

    # 初始化搜索选项列表
    searchOption = []

    # 构造SQL查询语句
    sql = f"SELECT ID, examName from examidd where StationCN = '{st.session_state.StationCN}' order by ID"

    # 执行SQL查询并获取结果
    rows = execute_sql(cur, sql)

    # 遍历查询结果，将考试名称添加到搜索选项列表中
    for row in rows:
        searchOption.append(row[1])

    # 显示考试名称选择框
    searchExamName = st.selectbox("请选择考试场次", searchOption, index=None)

    # 显示查询类型选择框
    options = st.multiselect("查询类型", ["通过", "未通过"], default=["通过", "未通过"])

    # 如果选择了考试场次
    if searchExamName:
        # 显示导出按钮
        searchButton = st.button("导出为Excel文件", type="primary")

        # 如果点击了导出按钮且选择了考试场次
        if searchButton and searchExamName:
            # 如果选择了查询类型
            if options:
                # 构造SQL查询语句
                sql = f"SELECT ID, userName, userCName, examScore, examDate, examPass from examresult where examName = '{searchExamName}' and ("

                # 遍历查询类型，构造SQL查询条件
                for each in options:
                    if each == "通过":
                        sql = sql + " examPass = 1 or "
                    elif each == "未通过":
                        sql = sql + " examPass = 0 or "

                # 去除SQL语句末尾的“ or ”
                if sql.endswith(" or "):
                    sql = sql[:-4] + ") order by ID"

                # 执行SQL查询并获取结果
                rows = execute_sql(cur, sql)

                # 构造输出文件名
                outputFile = f"./ExamResult/{searchExamName}_{time.strftime('%Y%m%d%H%M%S', time.localtime(int(time.time())))}.xlsx"

                # 如果文件已存在，则删除
                if os.path.exists(outputFile):
                    os.remove(outputFile)

                # 创建Excel工作簿和工作表
                workbook = Workbook(outputFile)
                worksheet = workbook.add_worksheet(f"{searchExamName}考试成绩")

                # 设置表头
                title = ["ID", "编码", "姓名", "成绩", "考试时间", "考试结果"]
                for index, value in enumerate(title):
                    worksheet.write(0, index, value)

                # 设置行计数器
                k = 1

                # 遍历查询结果，写入Excel表格
                for i, row in enumerate(rows):
                    for j, value in enumerate(row):
                        if j == 0:
                            value = k
                        elif j == 4:
                            value = time.strftime('%Y-%m-%d %H:%M:%S', time.localtime(int(value)))
                        elif j == 5:
                            value = "通过" if value == 1 else "未通过"
                        worksheet.write(i + 1, j, value)
                    k = k + 1

                # 关闭工作簿
                workbook.close()

                # 如果文件存在
                if os.path.exists(outputFile):
                    # 读取文件内容
                    with open(outputFile, "rb") as file:
                        content = file.read()
                    file.close()

                    # 显示下载按钮
                    buttonDL = st.download_button("点击下载", content, file_name=f"考试成绩_{outputFile[outputFile.rfind('/') + 1:]}", icon=":material/download:", type="secondary")

                    # 显示成功消息
                    st.success(f":green[[{searchExamName}]] :gray[考试成绩成功导出至程序目录下] :orange[{outputFile[2:]}]")

                    # 如果点击了下载按钮
                    if buttonDL:
                        st.toast("文件已下载至你的默认目录")
                        updateKeyAction("导出考试成绩")
                else:
                    # 显示错误消息
                    st.error(f":red[[{searchExamName}]] 考试成绩导出失败")


def ClearTables():
    try:
        # 删除 questions 表中的重复记录
        sql_delete_questions = """
            DELETE q1
            FROM questions q1
            JOIN questions q2
            ON q1.Question = q2.Question
            AND q1.qType = q2.qType
            AND q1.StationCN = q2.StationCN
            AND q1.chapterName = q2.chapterName
            WHERE q1.id > q2.id;
        """
        cur.execute(sql_delete_questions)

        # 删除 commquestions 表中的重复记录
        sql_delete_commquestions = """
            DELETE c1
            FROM commquestions c1
            JOIN commquestions c2
            ON c1.Question = c2.Question AND c1.qType = c2.qType
            WHERE c1.id > c2.id;
        """
        cur.execute(sql_delete_commquestions)

        # 删除 morepractise 表中的重复记录
        sql_delete_morepractise = """
            DELETE m1
            FROM morepractise m1
            JOIN morepractise m2
            ON m1.Question = m2.Question AND m1.qType = m2.qType AND m1.userName = m2.userName
            WHERE m1.id > m2.id;
        """
        cur.execute(sql_delete_morepractise)

        # 删除 questionaff 表中的重复记录
        sql_delete_questionaff = """
            DELETE a1
            FROM questionaff a1
            JOIN questionaff a2
            ON a1.chapterName = a2.chapterName AND a1.StationCN = a2.StationCN
            WHERE a1.id > a2.id;
        """
        cur.execute(sql_delete_questionaff)

        # 删除不在 questions 表中的 chapterName
        sql_delete_invalid_chapters = """
            DELETE FROM questionaff
            WHERE chapterName NOT IN ('公共题库', '错题集', '关注题集')
            AND chapterName NOT IN (SELECT DISTINCT(chapterName) FROM questions);
        """
        cur.execute(sql_delete_invalid_chapters)

        # 更新 users 表中的用户中文名，去除空格
        sql_update_users = """
            UPDATE users
            SET userCName = REPLACE(userCName, ' ', '')
            WHERE userCName LIKE '% %';
        """
        cur.execute(sql_update_users)

        # 去除问题字段中的换行符 - questions
        sql_update_questions = """
            UPDATE questions
            SET Question = REPLACE(Question, '\n', '')
            WHERE Question LIKE '%\n%';
        """
        cur.execute(sql_update_questions)

        # 去除问题字段中的换行符 - commquestions
        sql_update_commquestions = """
            UPDATE commquestions
            SET Question = REPLACE(Question, '\n', '')
            WHERE Question LIKE '%\n%';
        """
        cur.execute(sql_update_commquestions)

        # 去除问题字段中的换行符 - morepractise
        sql_update_morepractise = """
            UPDATE morepractise
            SET Question = REPLACE(Question, '\n', '')
            WHERE Question LIKE '%\n%';
        """
        cur.execute(sql_update_morepractise)

        # 提交事务
        conn.commit()

    except Exception as e:
        conn.rollback()
    finally:
        pass
    # 弹出提示信息，表示站室题库/公共题库/错题集/章节信息库记录清理完成
    #st.toast("站室题库/公共题库/错题集/章节信息库 记录清理完成")


def create_element(name):
    return OxmlElement(name)


def create_attribute(element, name, value):
    element.set(qn(name), value)


# noinspection PyProtectedMember
def add_page_number(run):
    fldChar1 = create_element('w:fldChar')
    create_attribute(fldChar1, 'w:fldCharType', 'begin')

    instrText = create_element('w:instrText')
    create_attribute(instrText, 'xml:space', 'preserve')
    instrText.text = "PAGE"

    fldChar2 = create_element('w:fldChar')
    create_attribute(fldChar2, 'w:fldCharType', 'end')

    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)


# noinspection PyProtectedMember,PyTypeChecker
def questoWord():
    allType, stationCName, chapterNamePack, outChapterName = [], [], [], []
    st.subheader("题库导出", divider="blue")
    sql = f"SELECT paramName, param from setup_{st.session_state.StationCN} where paramType = 'questype'"
    rows = execute_sql(cur, sql)
    for row in rows:
        allType.append(row[0])
    quesTable = st.selectbox("请选择功能类型", ("站室题库", "公共题库", "试卷", "错题集", "关注题集"), index=None)
    quesType = st.multiselect("题型", allType, default=allType)
    stationCN, headerExamName = "全站", ""
    if quesTable == "站室题库" or quesTable == "错题集" or quesTable == "关注题集":
        stationCName = getStationCNALL(flagALL=True)
        stationCN = st.select_slider("站室", stationCName, value=st.session_state.StationCN)
    elif quesTable == "试卷":
        headerExamName = st.text_input("请设置试卷名称", max_chars=20, help="文件抬头显示的试卷名称, 不填则使用默认名称")
        if "examFinalTable" in st.session_state:
            stationCN = st.session_state.StationCN
            st.write("📢:red[试卷题库如果导出文件中不包含设置的题型, 请按如下提示操作, 其他类型题库没有此限制.]")
            step = sac.steps(
                items=[
                    sac.StepsItem(title='参数设置'),
                    sac.StepsItem(title='题型设置'),
                    sac.StepsItem(title='重新生成题库'),
                    sac.StepsItem(title='试卷导出'),
                ], index=None, return_index=True
            )
            if step is not None:
                st.image(f"./Images/help/OutputFile{step}.png", caption=f"操作步骤{step + 1}")
        else:
            st.info("请先生成题库")
            quesTable = ""
    if stationCN != "全站" and quesTable == "站室题库":
        sql = f"SELECT chapterName from questionaff where StationCN = '{stationCN}' and chapterName <> '公共题库' and chapterName <> '错题集' and chapterName <> '关注题集' order by ID"
        rows = execute_sql(cur, sql)
        for row in rows:
            chapterNamePack.append(row[0])
        outChapterName = st.multiselect("章节", chapterNamePack, default=chapterNamePack)
    sac.switch(label="复核模式", on_label="On", align='start', size='md', value=False, key="sac_recheck")
    if st.session_state.sac_recheck:
        sac.switch(label="附加答题解析", on_label="On", align='start', size='md', value=False, key="sac_Analysis")
    else:
        if "sac_Analysis" not in st.session_state:
            st.session_state.sac_Analysis = False
    if quesTable and quesType:
        buttonSubmit = st.button("导出为Word文件", type="primary")
        if buttonSubmit:
            if quesTable == "站室题库":
                tablename = "questions"
            elif quesTable == "公共题库":
                tablename = "commquestions"
            elif quesTable == "试卷":
                tablename = st.session_state.examFinalTable
            elif quesTable == "错题集":
                tablename = "morepractise"
            elif quesTable == "关注题集":
                tablename = "favques"
            else:
                tablename = ""
            headerFS = getParam("抬头字体大小", st.session_state.StationCN)
            titleFS = getParam("题型字体大小", st.session_state.StationCN)
            quesFS = getParam("题目字体大小", st.session_state.StationCN)
            optionFS = getParam("选项字体大小", st.session_state.StationCN)
            answerFS = getParam("复核信息字体大小", st.session_state.StationCN)
            quesDOC = Document()
            quesDOC.styles["Normal"].font.name = "Microsoft YaHei"
            quesDOC.styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
            option, radioOption = ["A", "B", "C", "D", "E", "F", "G", "H"], ["正确", "错误"]
            blank = f"({' ' * 20})"
            pHeader = quesDOC.add_paragraph()
            pHeader.alignment = WD_ALIGN_PARAGRAPH.CENTER
            textHeader = pHeader.add_run(f"{st.session_state.StationCN} {headerExamName} {quesTable}", 0)
            #textHeader.font.name = "Microsoft YaHei"
            #textHeader.element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
            textHeader.font.size = Pt(headerFS)
            textHeader.font.bold = True
            textHeader.font.color.rgb = RGBColor(40, 106, 205)
            if quesTable == "试卷" and not st.session_state.sac_recheck and not st.session_state.sac_Analysis:
                pScore = quesDOC.add_paragraph()
                pScore.alignment = WD_ALIGN_PARAGRAPH.CENTER
                textScore = pScore.add_run("姓名:  _________                       分数:  _________", 0)
                textScore.font.size = Pt(optionFS)
                textScore.font.bold = True
            for each in quesType:
                if stationCN == "全站" or quesTable == "试卷":
                    sql = f"SELECT Question, qOption, qAnswer, qType, ID, SourceType, qAnalysis from {tablename} where qType = '{each}' order by ID"
                else:
                    if quesTable != "站室题库" and quesTable != "公共题库":
                        sql = f"SELECT Question, qOption, qAnswer, qType, ID, SourceType, qAnalysis from {tablename} where qType = '{each}' and StationCN = '{stationCN}' order by ID"
                    elif quesTable == "站室题库":
                        if outChapterName:
                            sql = f"SELECT Question, qOption, qAnswer, qType, ID, SourceType, qAnalysis from {tablename} where qType = '{each}' and StationCN = '{stationCN}' and (chapterName = "
                            for each5 in outChapterName:
                                sql += f"'{each5}' or chapterName = "
                            sql = sql[:-18] + ") order by chapterName, ID"
                        else:
                            sql = f"SELECT Question, qOption, qAnswer, qType, ID, SourceType, qAnalysis from {tablename} where qType = '{each}' and StationCN = '{stationCN}' order by chapterName, ID"
                rows = execute_sql(cur, sql)
                #st.write(f"{each} 共 {len(rows)}")
                i = 1
                if rows:
                    pTitle = quesDOC.add_paragraph()
                    textTitle = pTitle.add_run(f"{each}", 0)
                    #textTitle.font.name = "Microsoft YaHei"
                    #textTitle.element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
                    textTitle.font.size = Pt(titleFS)
                    textTitle.bold = True
                    for row in rows:
                        tmp, answer, qa, aa = "", "", [], []
                        pQues = quesDOC.add_paragraph()
                        if each == "填空题":
                            textQues = pQues.add_run(f"第{i}题   {row[0].replace('()', blank).replace('（）', blank)}", 0)
                        else:
                            textQues = pQues.add_run(f"第{i}题   {row[0]}   ({' ' * 8})", 0)
                        #textQues.font.name = "Microsoft YaHei"
                        #textQues.element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
                        #if st.session_state.sac_recheck and row[5] == "AI-LLM":
                        #textQues.font.color.rgb = RGBColor(155, 17, 30)
                        textQues.font.size = Pt(quesFS)
                        aa = row[2].replace("；", ";").split(";")
                        pOption = None
                        if each != "填空题":
                            pOption = quesDOC.add_paragraph()
                        if each == "单选题" or each == "多选题":
                            qa = row[1].replace("；", ";").split(";")
                            for each2 in qa:
                                tmp = tmp + f"{option[qa.index(each2)]}. {each2}{' ' * 8}"
                            textOption = pOption.add_run(tmp)
                            textOption.font.size = Pt(optionFS)
                        elif each == "判断题":
                            textOption = pOption.add_run(f"A. 正确{' ' * 15}B. 错误{' ' * 15}")
                            textOption.font.size = Pt(optionFS)
                        #textOption.font.name = "Microsoft YaHei"
                        #textOption.element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
                        #textOption.italic = True
                        if st.session_state.sac_recheck:
                            if each == "单选题" or each == "多选题":
                                for each3 in aa:
                                    answer = answer + f"{option[int(each3)]}" + ", "
                            elif each == "判断题":
                                answer = radioOption[int(aa[0]) ^ 1]
                            elif each == "填空题":
                                for index, value in enumerate(aa):
                                    answer = answer + f"{chr(49 + index)}.  [{value}]" + " "
                            if answer.endswith(", "):
                                answer = answer[:-2]
                            elif answer.endswith(" "):
                                answer = answer[:-1]
                            pAnswer = quesDOC.add_paragraph()
                            textAnswer = pAnswer.add_run(f"复核模式 ID: {row[4]} 正确答案: {answer}")
                            textAnswer.font.size = Pt(answerFS)
                            textAnswer.font.bold = True
                            textAnswer.font.color.rgb = RGBColor(155, 17, 30)
                            if stationCN != "全站":
                                sql = f"SELECT chapterName from questions where Question = '{row[0]}'"
                            else:
                                sql = f"SELECT chapterName from questions where Question = '{row[0]}' and StationCN = '{stationCN}'"
                            tempTable = execute_sql(cur, sql)
                            if tempTable:
                                fhQT = tempTable[0][0]
                            else:
                                sql = f"SELECT ID from commquestions where Question = '{row[0]}'"
                                if execute_sql(cur, sql):
                                    fhQT = "公共题库"
                                else:
                                    fhQT = "未知"
                            pSource = quesDOC.add_paragraph()
                            if not row[5].startswith("AI-LLM"):
                                textSource = pSource.add_run(f"试题来源: [{stationCN}] 章节名称: [{fhQT}] 试题生成类别: [{row[5]}]")
                            else:
                                textSource = pSource.add_run(f"请特别注意 试题来源: [{stationCN}] 章节名称: [{fhQT}] 试题生成类别: [{row[5]}]")
                            textSource.font.bold = True
                            textSource.font.size = Pt(answerFS)
                            if row[5].startswith("AI-LLM"):
                                textSource.font.color.rgb = RGBColor(155, 17, 30)
                                textSource.font.underline = True
                            #textSource.font.italic = True
                            if st.session_state.sac_Analysis and row[6] != "":
                                pAnalysis = quesDOC.add_paragraph()
                                if not row[5].startswith("AI-LLM"):
                                    textAnalysis = pAnalysis.add_run(f"人工解析: [{row[6].replace(':red', '').replace('[', '').replace(']', '').replace('**', '')}]")
                                else:
                                    textAnalysis = pAnalysis.add_run(f"请特别注意 A.I.解析: [{row[6].replace('**', '')}]")
                                textAnalysis.font.bold = True
                                textAnalysis.font.size = Pt(answerFS)
                                textAnalysis.font.color.rgb = RGBColor(79, 66, 181)
                                #textAnalysis.font.underline = True
                                #textAnalysis.font.italic = True
                        i += 1
            add_page_number(quesDOC.sections[0].footer.paragraphs[0].add_run())
            quesDOC.sections[0].footer.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            header = quesDOC.sections[0].header.paragraphs[0]
            header.text = f"{stationCN}\t\t{quesTable}"
            header.style = quesDOC.styles["Header"]
            if headerExamName != "":
                if st.session_state.sac_recheck:
                    outputFile = f"./QuesDoc/{stationCN}-{headerExamName}-{quesTable}-带审核信息_{time.strftime('%Y%m%d%H%M%S', time.localtime(int(time.time())))}.docx"
                else:
                    outputFile = f"./QuesDoc/{stationCN}-{headerExamName}-{quesTable}_{time.strftime('%Y%m%d%H%M%S', time.localtime(int(time.time())))}.docx"
            else:
                if st.session_state.sac_recheck:
                    outputFile = f"./QuesDoc/{stationCN}-{quesTable}-带审核信息_{time.strftime('%Y%m%d%H%M%S', time.localtime(int(time.time())))}.docx"
                else:
                    outputFile = f"./QuesDoc/{stationCN}-{quesTable}_{time.strftime('%Y%m%d%H%M%S', time.localtime(int(time.time())))}.docx"
            if os.path.exists(outputFile):
                os.remove(outputFile)
            quesDOC.save(outputFile)
            if os.path.exists(outputFile):
                if os.path.exists(outputFile):
                    with open(outputFile, "rb") as file:
                        content = file.read()
                    file.close()
                    buttonDL = st.download_button("点击下载", content, file_name=outputFile[outputFile.rfind("/") + 1:], icon=":material/download:", type="secondary")
                    st.success(f":green[[{quesTable}]] :gray[题库成功导出至程序目录下] :orange[{outputFile[2:]}]")
                    if buttonDL:
                        st.toast("文件已下载至你的默认目录")
            else:
                st.error(f":red[[{quesTable}]] 题库导出失败")


def dboutput():
    bc = sac.segmented(
        items=[
            sac.SegmentedItem(label="题库导出(Word格式)", icon="database-down"),
            #sac.SegmentedItem(label="试卷导出(DF格式)", icon="journal-arrow-down"),
            sac.SegmentedItem(label="考试成绩导出(Excel格式)", icon="layout-text-sidebar-reverse"),
        ], color="green", align="center"
    )
    if bc == "题库导出(Word格式)":
        questoWord()
    elif bc == "试卷导出(DF格式)":
        resultExcel()
    elif bc == "考试成绩导出(Excel格式)":
        examResulttoExcel()
    if bc is not None:
        updateActionUser(st.session_state.userName, bc, st.session_state.loginTime)


def actDelExamTable():
    for each in st.session_state.keys():
        if each.startswith("delExamTable_"):
            if st.session_state[each]:
                each = each.replace("delExamTable_", "")
                execute_sql_and_commit(conn, cur, sql=f"DROP TABLE IF EXISTS {each}")
                st.info(f"{each} 试卷删除成功")


def delExamTable():
    flagExistTable = False
    sql = "SELECT name from sqlite_master where type = 'table' and name like 'exam_%'"
    tempTable = execute_sql(cur, sql)
    if tempTable:
        st.subheader("删除试卷", divider="red")
        for row in tempTable:
            if row[0].count("_") == 3 or row[0].count("_") == 4:
                st.checkbox(f"{row[0]}", key=f"delExamTable_{row[0]}")
                flagExistTable = True
    if flagExistTable:
        if st.session_state.userType == "supervisor":
            if st.session_state.delExam:
                st.button("确认删除", on_click=actDelExamTable)
            else:
                st.error("试卷正在使用, 无法删除, 请先完成考试或练习后删除")
        else:
            st.error("仅Supervisor可进行此操作")
    else:
        st.info("暂无试卷")


# noinspection PyUnboundLocalVariable
def dbinputSubmit(tarTable, orgTable):
    tmpTable, sql, maxcol = "", "", 0

    # 根据目标表名设置不同的表名和SQL语句
    if tarTable == "站室题库":
        tablename = "questions"
        sql = f"INSERT INTO {tablename}(Question, qOption, qAnswer, qType, qAnalysis, StationCN, chapterName) VALUES (?, ?, ?, ?, ?, ?, ?)"
        maxcol = 7
    elif tarTable == "公共题库":
        tablename = "commquestions"
        sql = f"INSERT INTO {tablename}(Question, qOption, qAnswer, qType, qAnalysis) VALUES (?, ?, ?, ?, ?)"
        maxcol = 5

    # 如果SQL语句不为空，则执行以下操作
    if sql != "":
        st.spinner(f"正在向 [{tarTable}] 导入题库...")
        sql2 = f"SELECT Max(ID) from {tablename}"
        maxid = execute_sql(cur, sql2)[0][0]
        if maxid is None:
            maxid = 0

        # 遍历源表列表
        for each in orgTable:
            # 加载Excel文件
            listinsheet = openpyxl.load_workbook(f"./InputQues/{each}.xlsx")
            datainlist = listinsheet.active

            # 遍历Excel文件中的每一行数据
            for row in datainlist.iter_rows(min_row=2, max_col=maxcol, max_row=datainlist.max_row):
                singleQues = [cell.value for cell in row]
                if singleQues[0] is not None:
                    cur.execute(sql, singleQues)
                    conn.commit()

            # 关闭Excel文件
            listinsheet.close()

            # 如果文件名包含"_用户上传_"，则删除该文件
            if each.find("_用户上传_") != -1:
                os.remove(f"./InputQues/{each}.xlsx")

            # 拼接已处理的文件名
            tmpTable = tmpTable + each + ", "

        # 更新数据库中的空字段
        sql = f"UPDATE {tablename} set qOption = '' where qOption is Null"
        execute_sql_and_commit(conn, cur, sql)
        sql = f"UPDATE {tablename} set qAnalysis = '' where qAnalysis is Null"
        execute_sql_and_commit(conn, cur, sql)
        sql = f"UPDATE {tablename} set SourceType = '人工' where SourceType is Null"
        execute_sql_and_commit(conn, cur, sql)

        # 替换分号
        sql = f"UPDATE {tablename} set qOption = replace(qOption, '；', ';'), qAnswer = replace(qAnswer, '；', ';') where (qOption like '%；%' or qAnswer like '%；%') and (qType = '单选题' or qType = '多选题' or qType = '填空题')"
        execute_sql_and_commit(conn, cur, sql)

        # 更新题型
        sql = f"UPDATE {tablename} set qType = '单选题' where qType = '选择题' and ID > {maxid}"
        execute_sql_and_commit(conn, cur, sql)

        # 查询新添加的题目
        sql = f"SELECT ID, qOption, qAnswer, qType, Question from {tablename} where ID > {maxid} and (qType = '单选题' or qType = '多选题' or qType = '判断题')"
        rows = execute_sql(cur, sql)

        # 遍历查询结果，检查选项和答案序号是否相符
        for row in rows:
            sql = ""
            if row[3] == "单选题" or row[3] == "多选题":
                for each in row[2].split(";"):
                    if int(each) < 0 or int(each) >= len(row[1].split(";")) or len(row[1].split(";")) > 8:
                        sql = f"DELETE from {tablename} where ID = {row[0]}"
            elif row[3] == "判断题":
                if int(row[2]) < 0 or int(row[2]) > 1:
                    sql = f"DELETE from {tablename} where ID = {row[0]}"

            # 如果SQL语句不为空，则执行删除操作并显示警告
            if sql != "":
                execute_sql_and_commit(conn, cur, sql)
                st.warning(f"试题: [{row[4]}] 题型: [{row[3]}] 选项: [{row[1]}] 答案: [{row[2]}] 因为选项及答案序号不相符, 没有导入")

        # 插入章节信息
        sql = "INSERT INTO questionaff(chapterName, StationCN, chapterRatio, examChapterRatio) SELECT DISTINCT chapterName, StationCN, 5, 5 FROM questions"
        execute_sql_and_commit(conn, cur, sql)

        # 清除临时表
        ClearTables()

        # 显示成功信息
        st.success(f":green[{tmpTable[:-2]}.xlsx] 向 :red[{tarTable}] :gray[导入成功]")

        # 更新操作日志
        updateKeyAction(f"Excel文件导入试题至{tarTable}")


def dbinput():
    # 初始化输入选项列表
    inputOption = []

    # 从Streamlit获取用户选择的目标表
    targetTable = st.radio("导入至:", ("站室题库", "公共题库"), index=0, horizontal=True)

    # 从Streamlit获取用户选择的输入类型
    inputType = st.radio("文件来源:", ("服务器中文件", "上传文件"), index=0, horizontal=True)

    if targetTable:
        # 如果用户选择了“服务器中文件”作为输入类型
        if inputType == "服务器中文件":
            # 遍历"./InputQues"目录下的所有文件和文件夹
            for root, dirs, files in os.walk("./InputQues"):
                for file in files:
                    # 判断文件是否为.xlsx格式，且文件名中包含目标表名和站点名，且不是临时文件
                    if os.path.splitext(file)[1].lower() == '.xlsx' and f"{st.session_state.StationCN}_{targetTable}" in os.path.splitext(file)[0] and not os.path.splitext(file)[0].startswith("~$"):
                        # 将符合条件的文件名添加到输入选项列表中
                        inputOption.append(os.path.splitext(file)[0])

            if inputOption:
                # 如果存在可导入的文件，则显示文件选择框
                orgTable = st.multiselect("请选择导入文件", inputOption, default=None)
                if orgTable:
                    # 如果用户选择了文件，则显示导入按钮，并绑定点击事件
                    st.button("导入", on_click=dbinputSubmit, args=(targetTable, orgTable))
                else:
                    # 如果用户未选择文件，则显示提示信息
                    st.info("请选择要导入的文件")
            else:
                # 如果没有可导入的文件，则显示提示信息
                st.info("没有可导入的本站文件")

        # 如果用户选择了“上传文件”作为输入类型
        elif inputType == "上传文件":
            # 从Streamlit获取用户上传的文件
            uploaded_file = st.file_uploader("**请选择Excel文件**", type=["xlsx"])
            if uploaded_file is not None:
                # 读取文件内容
                bytes_data = uploaded_file.getvalue()
                # 生成上传文件的保存路径
                outFile = f"./InputQues/{st.session_state.StationCN}_{targetTable}_用户上传_{time.strftime('%Y%m%d%H%M%S', time.localtime(int(time.time())))}.xlsx"
                # 如果文件已存在，则删除旧文件
                if os.path.exists(outFile):
                    os.remove(outFile)
                # 将文件内容写入新文件
                with open(outFile, 'wb') as output_file:
                    output_file.write(bytes_data)
                # 如果文件成功保存，则调用dbinputSubmit函数进行导入
                if os.path.exists(outFile):
                    dbinputSubmit(targetTable, [outFile[12:-5]])
    else:
        # 如果用户未选择目标表，则显示提示信息
        st.write("请选择要导入的题库")


def dbfunc():
    if st.session_state.debug and int(st.session_state.userName) < 10:
        bc = sac.segmented(
            items=[
                sac.SegmentedItem(label="A.I.出题", icon="robot"),
                sac.SegmentedItem(label="题库导入", icon="database-up"),
                #sac.SegmentedItem(label="Word文件导入", icon="text-wrap"),
                sac.SegmentedItem(label="删除试卷", icon="trash3"),
                sac.SegmentedItem(label="删除静态题库", icon="trash3"),
                #sac.SegmentedItem(label="删除用户上传文件", icon="trash3"),
                sac.SegmentedItem(label="错题集重置", icon="journal-x"),
                sac.SegmentedItem(label="重置题库ID", icon="bootstrap-reboot"),
            ], align="center", color="red"
        )
    else:
        bc = sac.segmented(
            items=[
                sac.SegmentedItem(label="A.I.出题", icon="robot"),
                sac.SegmentedItem(label="题库导入", icon="database-up"),
                sac.SegmentedItem(label="删除试卷", icon="trash3"),
                sac.SegmentedItem(label="删除静态题库", icon="trash3"),
                #sac.SegmentedItem(label="删除用户上传文件", icon="trash3"),
            ], align="center", color="red"
        )
    if bc == "A.I.出题":
        AIGenerQues()
    elif bc == "题库导入":
        dbinput()
    elif bc == "Word文件导入":
        inputWord()
    elif bc == "删除试卷":
        delExamTable()
    elif bc == "删除静态题库":
        delStaticExamTable()
    elif bc == "删除用户上传文件":
        delUserUploadFiles()
    elif bc == "重置题库ID":
        buttonReset = st.button("重置题库ID", type="primary")
        if buttonReset:
            st.button("确认重置", type="secondary", on_click=resetTableID)
    if bc is not None:
        updateActionUser(st.session_state.userName, bc, st.session_state.loginTime)


def delUserUploadFiles():
    flagDelUserFiles = False
    for root, dirs, files in os.walk("./InputQues"):
        for file in files:
            if os.path.splitext(file)[1].lower() == '.xlsx' and "_用户上传_" in os.path.splitext(file)[0]:
                st.checkbox(os.path.splitext(file)[0], value=False, key=f"delUserFiles_{os.path.splitext(file)[0]}")
                flagDelUserFiles = True
    if flagDelUserFiles:
        buttonDel = st.button("删除", type="primary")
        if buttonDel:
            st.button("确认删除", type="secondary", on_click=actionDelUserUploadFiles)
    else:
        st.info("没有用户上传文件")


def actionDelUserUploadFiles():
    for key in st.session_state.keys():
        if key.startswith("delUserFiles_"):
            if st.session_state[key]:
                os.remove(f"./InputQues/{key.replace('delUserFiles_', '')}.xlsx")
            del st.session_state[key]
    st.success("所选文件已经删除")
    updateKeyAction("删除用户上传文件")


def resetActiveUser():
    sql = f"UPDATE users set activeUser = 0 where userName <> {st.session_state.userName}"
    execute_sql_and_commit(conn, cur, sql)
    st.success("已重置所有用户状态")
    updateKeyAction("重置所有用户状态")


# noinspection PyUnboundLocalVariable
def inputWord():
    #doc = Document("./QuesRefer/2023年全国特种设备作业人员考试题库附答案.docx")
    doc = Document("./QuesRefer/2023年特种设备作业安全管理人员证考试题库(通用版).docx")
    chapter = "特种设备安全管理员"
    #title_rule = re.compile("\\d+、")
    #title_rule = re.compile("\\d+.")
    option_rule = re.compile("\\w+、")
    ques, qAnswer, temp2, generQuesCount, qType = "", "", "", 0, ""
    if st.session_state.debug:
        os.system("cls")
    st.spinner("正在导入Word文件...")
    for i, paragraph in enumerate(doc.paragraphs[:]):
        line = paragraph.text.replace('\n', '').replace('\r', '').replace("（", "(").replace("）", ")").strip()
        if line:
            #if title_rule.search(line):
            if line[:7].find(".") != -1:
                if temp2.endswith(";"):
                    temp2 = temp2[:-1]
                    qOption = temp2
                    temp2 = ""
                if ques != "" and qAnswer != "" and qOption != "":
                    if qOption.find("正确;错误") != -1:
                        qType = "判断题"
                        qAnswer = int(qAnswer) ^ 1
                        qOption = ""
                    elif len(qAnswer) == 1:
                        qType = "单选题"
                    elif len(qAnswer) > 1:
                        qType = "多选题"
                    if st.session_state.debug:
                        print(f"Record: Q: {ques} T: {qType} O: {qOption} A: {qAnswer}")
                    sql = f"SELECT ID from questions where Question = '{ques}' and qType = '{qType}' and StationCN = '{st.session_state.StationCN}' and chapterName = '{chapter}'"
                    if not execute_sql(cur, sql):
                        sql = f"INSERT INTO questions(Question, qOption, qAnswer, qType, StationCN, chapterName, SourceType) VALUES ('{ques}', '{qOption}', '{qAnswer}', '{qType}', '{st.session_state.StationCN}', '{chapter}', '人工')"
                        execute_sql_and_commit(conn, cur, sql)
                        generQuesCount += 1
                    ques, qAnswer, qOption = "", "", ""
                if st.session_state.debug:
                    print(f"Ques:{line}")
                if line[:7].find("、") != -1:
                    ques = line[line.find("、") + 1:]
                elif line[:7].find(".") != -1:
                    ques = line[line.find(".") + 1:]
                if ques.startswith("."):
                    ques = ques[1:]
                qAnswer = ""
                while True:
                    b1 = line.find('(')
                    b2 = line.find(')')
                    if b1 != -1 and b2 != -1 and line[b1 + 1:b1 + 2] in ["A", "B", "C", "D", "E", "F"]:
                        temp = line[b1 + 1:b2]
                        ques = ques.replace(temp, " " * len(temp))
                        temp = temp.replace("、", "")
                        for each in temp:
                            qAnswer = qAnswer + str(ord(each) - 65) + ";"
                        line = line[b2 + 1:]
                    else:
                        break
                if qAnswer.endswith(";"):
                    qAnswer = qAnswer[:-1]
            elif option_rule.search(line):
                if st.session_state.debug:
                    print(f"{line}")
                temp2 = temp2 + line[2:] + ";"
            elif line.find("正确答案：") != -1:
                if st.session_state.debug:
                    print(line)
                temp = line[line.find("正确答案：") + 5:]
                for each in temp:
                    qAnswer = qAnswer + str(ord(each) - 65) + ";"
                if qAnswer.endswith(";"):
                    qAnswer = qAnswer[:-1]
        else:
            continue
    ClearTables()
    st.success(f"共生成{generQuesCount}道试题")
    updateKeyAction("导入试题")


def resetTableID():
    tables = [
        "questions", "commquestions", "morepractise", "favques",
        "examidd", "examresult", "questionaff", "studyinfo",
        "users", "keyactionlog", "setup_默认", f"setup_{st.session_state.StationCN}"
    ]

    for tablename in tables:
        try:
            # 获取当前表的所有ID并按顺序排序
            sql = f"SELECT ID FROM {tablename} ORDER BY ID"
            cur.execute(sql)
            rows = cur.fetchall()

            if not rows:
                continue

            # 更新ID字段为连续值
            for i, row in enumerate(rows):
                new_id = i + 1
                old_id = row['ID']

                update_sql = f"UPDATE {tablename} SET ID = {new_id} WHERE ID = {old_id}"
                cur.execute(update_sql)

                # 如果是 questions 或 commquestions，还需更新 studyinfo 表中的 cid
                if tablename in ["questions", "commquestions"]:
                    update_studyinfo_sql = (
                        f"UPDATE studyinfo SET cid = {new_id} "
                        f"WHERE cid = {old_id} AND questable = '{tablename}'"
                    )
                    cur.execute(update_studyinfo_sql)

            # 更新自增序列（MySQL 使用 AUTO_INCREMENT）
            if rows:
                last_id = len(rows)
                alter_sql = f"ALTER TABLE {tablename} AUTO_INCREMENT = {last_id + 1}"
                cur.execute(alter_sql)

        except Exception as e:
            conn.rollback()
            st.error(f"重置 {tablename} 表ID失败: {e}")
            continue

    conn.commit()
    st.success("题库ID重置成功")
    updateKeyAction("重置题库ID")        #st.toast(f"重置 {tablename} 表ID完毕")


# noinspection PyShadowingNames,PyUnboundLocalVariable
def AIGenerQues():
    quesPack, chars, chapterPack, dynaQuesType, generQuesCount = [], ["A", "B", "C", "D", "E", "F", "G", "H"], [], ["单选题", "多选题", "判断题", "填空题"], 0
    StationCNPack, chosenStationCN = [], st.session_state.StationCN
    temp = f"{st.session_state.StationCN}-站室题库现有: "
    for each in dynaQuesType:
        sql = f"SELECT Count(ID) from questions where qType = '{each}' and StationCN = '{st.session_state.StationCN}'"
        qCount = execute_sql(cur, sql)[0][0]
        temp = temp + ":red[" + each + "]: " + str(qCount) + "道 "
    temp = temp + "\n\n公共题库现有: "
    for each in dynaQuesType:
        sql = f"SELECT Count(ID) from commquestions where qType = '{each}'"
        qCount = execute_sql(cur, sql)[0][0]
        temp = temp + ":red[" + each + "]: " + str(qCount) + "道 "
    temp = temp.strip()
    st.caption(temp)
    table = st.radio(label="请选择要生成的题库", options=("站室题库", "公共题库"), horizontal=True, index=None)
    if table and table != "公共题库":
        sql = "SELECT Station from stations order by ID"
        rows = execute_sql(cur, sql)
        for row in rows:
            StationCNPack.append(row[0])
        chosenStationCN = st.select_slider("请选择要导入的站室", options=StationCNPack, value=st.session_state.StationCN)
        col1, col2 = st.columns(2)
        sql = f"SELECT chapterName from questionaff where StationCN = '{chosenStationCN}' and chapterName <> '公共题库' and chapterName <> '错题集' and chapterName <> '关注题集'"
        rows = execute_sql(cur, sql)
        for row in rows:
            chapterPack.append(row[0])
        with col1:
            chapter = st.selectbox(label="请选择章节", options=chapterPack, index=None)
        with col2:
            textChapter = st.text_input("请输入新章节名称", value="", placeholder="添加新的章节")
            textChapter = textChapter.strip()
    elif table == "公共题库":
        chapter, textChapter = "", ""
    quesRefer = st.text_area("请输入参考资料")
    AIModelNamePack = st.multiselect(
        "可选LLM大模型",
        ["DeepSeek", "文心千帆", "讯飞星火"],
        ["DeepSeek", "文心千帆", "讯飞星火"],
    )
    quesTypePack = st.multiselect(
        "请选择要生成的题型",
        dynaQuesType,
        dynaQuesType,
    )
    quesCount = st.number_input("请输入要生成的题目数量", min_value=1, max_value=10, value=5, step=1)
    if table is not None and quesRefer != "" and AIModelNamePack != [] and quesTypePack != []:
        buttonGener = st.button("生成试题")
        if buttonGener:
            if chapter is None and textChapter != "":
                sql = f"SELECT ID from questionaff where chapterName = '{textChapter}' and StationCN = '{chosenStationCN}'"
                if not execute_sql(cur, sql):
                    sql = f"INSERT INTO questionaff(chapterName, StationCN, chapterRatio, examChapterRatio) VALUES ('{textChapter}', '{chosenStationCN}', 5, 5)"
                    execute_sql_and_commit(conn, cur, sql)
                    st.toast(f"新的章节: :red[{textChapter}]添加完毕")
                chapter = textChapter
            if chapter is not None and table == "站室题库" or table == "公共题库":
                if st.session_state.debug:
                    os.system("cls")
                generQuesCount, displayQues, generQuesCountPack = 0, "", []
                infoArea = st.empty()
                for quesType in quesTypePack:
                    gqc = 0
                    for AIModelName in AIModelNamePack:
                        with infoArea.container(border=True):
                            st.info(f"正在使用 :red[{AIModelName}大模型] 进行:blue[{quesType}] 试题生成, 请稍等...")
                        if AIModelName == "文心千帆":
                            ques = qianfan_AI_GenerQues(quesRefer, quesType, quesCount, "ERNIE-Speed-8K")
                        elif AIModelName == "DeepSeek":
                            ques = deepseek_AI_GenerQues(quesRefer, quesType, quesCount)
                        elif AIModelName == "讯飞星火":
                            ques = xunfei_xh_AI_GenerQues(quesRefer, quesType, quesCount)
                        else:
                            ques = ""
                        quesPack = ques.split("题型")
                        for each in quesPack:
                            if each != "":
                                quesHeader, qOption, Option, qAnswer, qAnalysis, flagSuccess = "", "", [], "", "", True
                                temp = each[:10]
                                for dqt in dynaQuesType:
                                    if temp.find(dqt) != -1:
                                        quesType = dqt
                                        break
                                b1 = each.find("试题")
                                if b1 != -1:
                                    each = each[b1 + 3:]
                                    c1 = each.find("标准答案")
                                    c2 = each.find("试题解析")
                                    b2 = each.find("选项")
                                    if c1 != -1 and c2 != -1:
                                        if quesType == "填空题":
                                            quesHeader = each[:c1].replace("\n\n", "").replace("\n", "").replace("**", "").replace("无选项，填写空白处即可。", "").replace("无选项，本题为填空题", "").replace("选项：", "").strip()
                                        else:
                                            quesHeader = each[:b2].replace("\n\n", "").replace("\n", "").replace("**", "").strip()
                                        if quesHeader.startswith("*: "):
                                            quesHeader = quesHeader[3:]
                                        if quesType == "单选题" or quesType == "多选题":
                                            b3 = each.find("标准答案")
                                            if b3 != -1:
                                                Option = each[b2 + 3:b3].replace("\n\n", "").replace("正确的？ 选项：", "").replace("正确的？", "").strip().split("\n")
                                                displayOption = each[b2 + 3:b3].replace("**", "").replace("正确的？ 选项：", "").replace("正确的？", "").strip()
                                                displayOption = displayOption.replace("A. ", "\n\nA. ").replace("B. ", "\nB. ").replace("C. ", "\nC. ").replace("D. ", "\nD. ").replace("E. ", "\nE. ").replace("F. ", "\nF. ").replace("G. ", "\nG. ").replace("H. ", "\nH. ")
                                                for each2 in Option:
                                                    for each3 in chars:
                                                        each2 = each2.replace(f"{each3}.", "").strip()
                                                    qOption = qOption + each2 + ";"
                                                if AIModelName == "讯飞星火" and len(Option) == 1:
                                                    qOption = qOption[:-1]
                                                    qOption = qOption.replace("  ", ";")
                                                    revDisplayOption = str(displayOption)
                                                    revDisplayOption = revDisplayOption.replace("A. ", "\n\nA. ").replace("B. ", "\nB. ").replace("C. ", "\nC. ").replace("D. ", "\nD. ").replace("E. ", "\nE. ").replace("F. ", "\nF. ").replace("G. ", "\nG. ").replace("H. ", "\nH. ")
                                                    displayOption = revDisplayOption
                                                qOption = qOption.replace("；", ";")
                                                if qOption.endswith(";"):
                                                    qOption = qOption[:-1]
                                                if st.session_state.debug:
                                                    print(f"Option:{Option} qOption:{qOption}")
                                                b4 = each.find("试题解析")
                                                if b4 != -1:
                                                    qAnswer = each[b3 + 5:b4].replace("\n", "").replace("*", "").strip()
                                                    displayAnswer = qAnswer
                                                    qAnalysis = each[b4 + 5:].replace("\n", "").replace("*", "").strip()
                                                    if quesType == "单选题":
                                                        qAnswer = ord(qAnswer[0].upper()) - 65
                                                        if qAnswer > 7 or qAnswer < 0:
                                                            flagSuccess = False
                                                    elif quesType == "多选题":
                                                        qAnswer = qAnswer.replace("，", "").replace(",", "").replace("、", "").replace("。", "").replace(" ", "").replace("（", "(")
                                                        if qAnswer.find("(") != -1:
                                                            qAnswer = qAnswer[:qAnswer.find("(")].strip()
                                                        temp = ""
                                                        if st.session_state.debug:
                                                            print(f"未处理前的多选题标准答案:{qAnswer}")
                                                        for each4 in qAnswer:
                                                            if ord(each4.upper()) - 65 > 7 or ord(each4.upper()) - 65 < 0:
                                                                flagSuccess = False
                                                                break
                                                            else:
                                                                temp = temp + str(ord(each4) - 65) + ";"
                                                        qAnswer = temp
                                                        if qAnswer.endswith(";"):
                                                            qAnswer = qAnswer[:-1]
                                        elif quesType == "判断题":
                                            if each[c1 + 5:c2].find("正确") != -1 or each[c1 + 5:c2].find("A") != -1:
                                                qAnswer = "1"
                                                displayAnswer = "正确"
                                            else:
                                                qAnswer = "0"
                                                displayAnswer = "错误"
                                            displayOption = "A. 正确\nB. 错误\n"
                                            qAnalysis = each[c2 + 5:].replace("\n", "").replace("*", "").strip()
                                        elif quesType == "填空题":
                                            displayOption = ""
                                            qAnswer = each[c1 + 5:c2].replace("\n", "").replace("*", "").replace("无选项", "").replace("；", ";").replace("，", ";").replace("。", "").replace("、", ";").strip()
                                            if qAnswer.startswith(":"):
                                                qAnswer = qAnswer[1:]
                                            displayAnswer = qAnswer
                                            qAnalysis = each[c2 + 5:].replace("\n", "").replace("*", "").strip()
                                            i = 12
                                            while i > 0:
                                                if quesHeader.find("_" * i) != -1:
                                                    quesHeader = quesHeader.replace("_" * i, "()")
                                                i -= 1
                                    if quesHeader.startswith(":"):
                                        quesHeader = quesHeader[1:].strip()
                                    if qAnalysis.startswith(":"):
                                        qAnalysis = qAnalysis[1:].strip()
                                    if qAnalysis.endswith("---"):
                                        qAnalysis = qAnalysis[:-3].strip()
                                    if quesType == "单选题" and len(str(qAnswer)) > 1:
                                        flagSuccess = False
                                    if qOption.count(";") == 0 and (quesType == "单选题" or quesType == "多选题"):
                                        flagSuccess = False
                                    if st.session_state.debug:
                                        print(f"debug: 题目:[{quesHeader}] 选项:[{qOption}], 标准答案:[{qAnswer}] 答题解析:[{qAnalysis}]")
                                    if quesType == "填空题":
                                        quesHeader = quesHeader.replace("选项未给出，需要学生在横线上填写正确答案。", "").replace("选项:", "")
                                if qAnswer != "" and quesHeader != "" and len(str(qAnswer)) < 200 and len(quesHeader) < 200 and flagSuccess:
                                    if table == "公共题库":
                                        sql = f"SELECT ID from commquestions where Question = '{quesHeader}' and qType = '{quesType}'"
                                        if not execute_sql(cur, sql):
                                            sql = f"INSERT INTO commquestions(Question, qOption, qAnswer, qType, qAnalysis, SourceType) VALUES('{quesHeader}', '{qOption}', '{qAnswer}', '{quesType}', '{qAnalysis}', 'AI-LLM-{AIModelName}')"
                                            execute_sql_and_commit(conn, cur, sql)
                                            generQuesCount += 1
                                            gqc += 1
                                            displayQues = displayQues + f":blue[**第{generQuesCount}题:**]\n\n:red[题型: ]{quesType}\n\n:red[题目: ]{quesHeader}\n\n:red[选项: ]\n{displayOption}\n\n:red[答案: ]{displayAnswer}\n\n:red[解析: ]{qAnalysis}\n\n{'-' * 40}\n\n"
                                    elif table == "站室题库":
                                        sql = f"SELECT ID from questions where Question = '{quesHeader}' and qType = '{quesType}' and StationCN = '{chosenStationCN}' and chapterName = '{chapter}'"
                                        if not execute_sql(cur, sql):
                                            sql = f"INSERT INTO questions(Question, qOption, qAnswer, qType, qAnalysis, StationCN, chapterName, SourceType) VALUES('{quesHeader}', '{qOption}', '{qAnswer}', '{quesType}', '{qAnalysis}', '{chosenStationCN}', '{chapter}', 'AI-LLM-{AIModelName}')"
                                            execute_sql_and_commit(conn, cur, sql)
                                            generQuesCount += 1
                                            gqc += 1
                                            displayQues = displayQues + f":blue[**第{generQuesCount}题:**]\n\n:red[题型: ]{quesType}\n\n:red[题目: ]{quesHeader}\n\n:red[选项: ]\n{displayOption}\n\n:red[答案: ]{displayAnswer}\n\n:red[解析: ]{qAnalysis}\n\n{'-' * 40}\n\n"
                    generQuesCountPack.append(gqc)
                infoArea.empty()
                if generQuesCount > 0:
                    tempInfo = f"试题生成完毕, 总计生成试题{generQuesCount}道, 其中"
                    for index, value in enumerate(quesTypePack):
                        tempInfo = tempInfo + f"{value}: {generQuesCountPack[index]}道, "
                    st.success(tempInfo[:-2])
                    st.subheader("具体如下:", divider="green")
                    st.markdown(displayQues)
                    if table == "公共题库":
                        updateKeyAction(f"A.I.生成试题{generQuesCount}道至{table}题库")
                    elif table == "站室题库":
                        updateKeyAction(f"A.I.生成试题{generQuesCount}道至{table}题库{chapter}章节")
                else:
                    st.info("A.I.未生成到任何试题, 请检查参考资料是否正确或是生成的试题已经在题库中")
            else:
                st.warning("站室题库请选择章节")
    else:
        st.info("请设置各选项和添加参考资料")


def ClearMP():
    buttonSubmit = st.button(f"清空 {st.session_state.userCName} 错题集", type="primary")
    if buttonSubmit:
        bcArea = st.empty()
        with bcArea.container():
            st.button("确认清空", type="secondary", on_click=ClearMPAction, args=(bcArea,))


def ClearMPAction(bcArea):
    execute_sql_and_commit(conn, cur, sql=f"DELETE from morepractise where userName = {st.session_state.userName}")
    bcArea.empty()
    st.success("当前用户错题集已重置")
    updateKeyAction("重置当前用户错题集")


def studyinfo():
    study = sac.segmented(
        items=[
            sac.SegmentedItem(label="学习进度", icon="grid-3x2-gap"),
            sac.SegmentedItem(label="错题集", icon="list-stars"),
            sac.SegmentedItem(label="章节时间线", icon="clock-history"),
            sac.SegmentedItem(label="学习记录重置", icon="bootstrap-reboot"),
            sac.SegmentedItem(label="错题集重置", icon="journal-x"),
        ], align="center", color="red"
    )
    if study == "学习进度":
        studyinfoDetail()
    elif study == "错题集":
        displayErrorQues()
    elif study == "章节时间线":
        generTimeline()
    elif study == "学习记录重置":
        studyReset()
    elif study == "错题集重置":
        ClearMP()
    if study is not None:
        updateActionUser(st.session_state.userName, f"查看信息-{study}", st.session_state.loginTime)


def userRanking():
    study = sac.segmented(
        items=[
            sac.SegmentedItem(label="榜单", icon="bookmark-star"),
            sac.SegmentedItem(label="证书", icon="patch-check"),
            sac.SegmentedItem(label="荣誉榜", icon="mortarboard"),
        ], align="center", color="red"
    )
    if study == "榜单":
        displayUserRanking()
    elif study == "证书":
        displayCertificate()
    elif study == "荣誉榜":
        displayMedals()
    if study is not None:
        updateActionUser(st.session_state.userName, f"证书及榜单-{study}", st.session_state.loginTime)


# noinspection PyShadowingNames
def displayUserRanking():
    xData, yData, boardInfo = [], [], ""
    col1, col2, col3 = st.columns(3)
    boardType = col1.radio("榜单", options=["个人榜", "站室榜"], index=0, horizontal=True)
    heatmap = col2.radio("热力图", options=["Folium", "Pydeck"], index=0, horizontal=True)
    #maptype = col3.radio("地图", options=["OpenStreetMap", "高德"], index=0, horizontal=True)
    markertype = col3.radio("标记", options=["默认", "公司Logo"], index=1, horizontal=True)
    maptype = "高德"
    if boardType == "个人榜":
        sql = "SELECT userCName, StationCN, userRanking from users where userRanking > 0 order by userRanking DESC, ID limit 0, 10"
    elif boardType == "站室榜":
        sql = "SELECT StationCN, ID, sum(userRanking) as Count from users GROUP BY StationCN having Count > 0 order by Count DESC"
    else:
        sql = ""
    rows = execute_sql(cur, sql)
    for index, row in enumerate(rows):
        xData.append(row[0])
        yData.append(row[2])
        if boardType == "个人榜":
            boardInfo = boardInfo + f"第 {index + 1} 名: {row[0]} 站室: {row[1]} 刷题数: {row[2]}\n\n"
        elif boardType == "站室榜":
            boardInfo = boardInfo + f"第 {index + 1} 名: {row[0]} 刷题数: {row[2]}\n\n"
        else:
            boardInfo = ""
    itemArea = st.empty()
    colors = ["lightslategray",] * len(rows)
    colors[0] = "crimson"
    fig = go.Figure(data=[go.Bar(x=xData, y=yData, marker_color=colors)])
    #fig.update_layout(font=dict(family="Courier New, monospace", size=18))
    fig.update_layout(title_text=f"{boardType[:-1]}刷题榜")
    with itemArea.container(border=True):
        st.plotly_chart(fig, theme="streamlit")
        #st.bar_chart(data=pd.DataFrame({"用户": xData, "试题数": yData}), x="用户", y="试题数", color=(155, 17, 30))
    if boardType == "站室榜" and int(rows[0][2]) > 0:
        if heatmap == "Pydeck":
            data = []
            for row in rows:
                sql = f"SELECT lat, lng, Station from stations where Station = '{row[0]}'"
                tmpTable = execute_sql(cur, sql)
                for i in range(row[2]):
                    data.append([round(tmpTable[0][0] / 100, 2), round(tmpTable[0][1] / 100, 2)])
            chart_data = pd.DataFrame(data, columns=["lat", "lng"],)
            st.pydeck_chart(
                pdk.Deck(
                    map_style="road",
                    initial_view_state=pdk.ViewState(
                        latitude=data[0][0],
                        longitude=data[0][1],
                        zoom=10,
                        pitch=50,
                    ),
                    layers=[
                        pdk.Layer(
                            "HexagonLayer",
                            data=chart_data,
                            get_position="[lng, lat]",
                            radius=200,
                            elevation_scale=4,
                            elevation_range=[0, 3000],
                            pickable=True,
                            extruded=True,
                            coverage=1,
                        ),
                        pdk.Layer(
                            "ScatterplotLayer",
                            data=chart_data,
                            get_position="[lng, lat]",
                            get_color="[37, 150, 209, 160]",
                            get_radius=200,
                            coverage=1,
                        ),
                    ],
                )
            )
        elif heatmap == "Folium":
            heatData = []
            sql = "SELECT StationCN, sum(userRanking) as Ranking from users GROUP BY StationCN having Ranking > 0 order by Ranking DESC"
            rows = execute_sql(cur, sql)
            sql = f"SELECT lat, lng, Station from stations where Station = '{rows[0][0]}'"
            row = execute_sql(cur, sql)[0]
            lat = round(row[0] / 100, 2)
            lng = round(row[1] / 100, 2)
            m = None
            if maptype == "OpenStreetMap":
                m = folium.Map(location=[lat, lng], zoom_start=11, TileLayer="OpenStreetMap", control_scale=True)
            elif maptype == "高德":
                m = folium.Map(
                    location=[lat, lng],
                    tiles="https://wprd01.is.autonavi.com/appmaptile?x={x}&y={y}&z={z}&lang=zh_cn&size=1&scl=1&style=7",
                    attr='高德-路网图',
                    zoom_start=11,
                    control_scale=True,
                    )
            for row in rows:
                sql = f"SELECT lat, lng from stations where Station = '{row[0]}'"
                row2 = execute_sql(cur, sql)[0]
                lat = round(row2[0] / 100, 2)
                lng = round(row2[1] / 100, 2)
                iframe = folium.IFrame(f"{row[0]} 刷题{row[1]}道")
                popup = folium.Popup(iframe, min_width=120, max_width=300)
                icon = folium.features.CustomIcon(
                    "./Images/logos/cnaf-logo.png",
                    icon_size=(40, 40),
                    icon_anchor=(20, 40),
                    popup_anchor=(0, -40),
                )
                if markertype == "默认":
                    folium.Marker([lat, lng], popup=popup).add_to(m)
                elif markertype == "公司Logo":
                    folium.Marker([lat, lng], icon=icon, popup=popup).add_to(m)
                heatData.append([lat, lng, int(row[1])])
            HeatMap(heatData).add_to(m)
            minimap = MiniMap(
                toggle_display=True,
                width=120,
                height=120,
                minimized=True,
                )
            m.add_child(minimap)
            st_folium(m, use_container_width=True, height=430)
    st.subheader(boardInfo)


def generTimeline():
    timelineData, i = [], 1  # 初始化时间线数据和计数器
    # 构造SQL查询语句，获取指定站点和章节名的题目
    sql = f"SELECT chapterName from questionaff where StationCN = '{st.session_state.StationCN}' and chapterName <> '错题集' order by ID"
    rows = execute_sql(cur, sql)  # 执行SQL查询，获取章节名列表
    for row in rows:
        # 判断章节名是否为"公共题库"，若不是则查询该章节的题目数量
        if row[0] != "公共题库":
            sql = f"SELECT Count(ID) from questions where chapterName = '{row[0]}'"
            quesCount = execute_sql(cur, sql)[0][0]
        else:
            # 若章节名为"公共题库"，则查询公共题库的题目数量
            sql = "SELECT Count(ID) from commquestions"
            quesCount = execute_sql(cur, sql)[0][0]
        # 构造SQL查询语句，获取指定用户名和章节名的学习时间信息
        sql = f"SELECT startTime from studyinfo where userName = '{st.session_state.userName}' and chapterName = '{row[0]}' order by startTime"
        rows2 = execute_sql(cur, sql)  # 执行SQL查询，获取学习时间信息列表
        if rows2:
            # 格式化学习开始和结束时间
            trainingDate = time.strftime("%Y-%m-%d", time.localtime(rows2[0][0]))
            trainingDate2 = time.strftime("%Y-%m-%d", time.localtime(rows2[-1][0]))
            # 判断学习时间信息数量是否等于题目数量
            if len(rows2) == quesCount:
                temp = {"id": i, "content": row[0], "start": trainingDate, "end": trainingDate2}  # 记录完整学习时间范围
            else:
                temp = {"id": i, "content": row[0], "start": trainingDate, "type": "point"}  # 记录学习开始时间
            timelineData.append(temp)  # 将学习记录添加到时间线数据中
            i += 1  # 计数器自增
    #st.write(timelineData)
    if timelineData:
        timeline = st_timeline(timelineData, groups=[], options={}, height="300px")  # 构造时间线组件
        if timeline is not None:
            if "end" in timeline:
                st.write(f"章节: :green[{timeline['content']}] 练习开始时间: :blue[{timeline['start']}] 完成时间: :orange[{timeline['end']}]")  # 输出包含完成时间的章节学习记录
            else:
                st.write(f"章节: :green[{timeline['content']}] 练习开始时间: :blue[{timeline['start']}]")  # 输出仅包含开始时间的章节学习记录
    else:
        st.write(":red[暂无学习记录]")  # 输出无学习记录提示


def displayCertificate():
    flagGener, flagInfo = False, True

    # 查询考试名称
    sql = f"SELECT examName from examidd where StationCN = '{st.session_state.StationCN}' and examName <> '练习题库' order by ID"
    rows = execute_sql(cur, sql)

    for row in rows:
        # 查询考试结果
        sql = f"SELECT userCName, examScore, examDate, CertificateNum, ID from examresult where userName = '{st.session_state.userName}' and examName = '{row[0]}' and examPass = 1 order by examScore DESC limit 0, 1"
        rows2 = execute_sql(cur, sql)

        if rows2:
            flagGener = True

            # 显示打印证书提示
            if flagGener and flagInfo:
                st.write(":orange[如需打印, 请打开 :green[程序目录下Image/Certificate] 或者点击下载证书]")
                flagInfo = False

            examDetail = rows2[0]

            with st.expander(label=f"{row[0]}", expanded=False):
                # 格式化考试日期
                examDateDetail = time.strftime("%Y%m%d%H%M%S", time.localtime(examDetail[2]))

                # 获取最大证书编号
                if examDetail[3] == 0:
                    sql = "SELECT Max(CertificateNum) from examresult"
                    maxCertNum = execute_sql(cur, sql)[0][0] + 1
                else:
                    maxCertNum = examDetail[3]

                # 生成证书文件路径
                certFile = f"./Images/Certificate/Cert-Num.{str(maxCertNum).rjust(5, '0')}-{st.session_state.userName}-{examDetail[0]}-{row[0]}_{examDateDetail}.png"

                # 如果证书文件不存在
                if not os.path.exists(certFile):
                    # 根据成绩选择奖牌
                    if examDetail[1] >= 100:
                        medal = "./Images/gold-award.png"
                    elif examDetail[1] >= 90:
                        medal = "./Images/silver-award.png"
                    else:
                        medal = "./Images/bronze-award.png"

                    # 格式化考试日期
                    examDate = time.strftime("%Y-%m-%d", time.localtime(examDetail[2]))

                    # 生成证书
                    generCertificate(certFile, medal, st.session_state.userCName, row[0], examDate, maxCertNum)

                # 如果证书文件存在
                if os.path.exists(certFile):
                    # 更新考试结果中的证书编号
                    sql = f"UPDATE examresult set CertificateNum = {maxCertNum} where ID = {examDetail[4]}"
                    execute_sql_and_commit(conn, cur, sql)

                    # 显示证书图片
                    st.image(certFile)

                    # 提供证书下载按钮
                    with open(certFile, "rb") as file:
                        st.download_button(
                            label="下载证书",
                            data=file,
                            file_name=certFile[certFile.rfind("/") + 1:].replace("Cert", "证书"),
                            mime="image/png",
                            icon=":material/download:"
                        )
                    file.close()

    # 如果没有通过任何考试
    if not flagGener:
        st.info("您没有通过任何考试, 无法生成证书")


def generCertificate(certFile, medal, userCName, examName, examDate, maxCertNum):
    # 姓名位置数组
    namePosX = [866, 821, 796, 760, 726, 696]

    # 如果用户中文名的长度为2，则在其间添加空格
    if len(userCName) == 2:
        userCName = userCName[0] + " " + userCName[-1]

    # 加载字体文件
    font = ImageFont.truetype("./Fonts/msyhbd.ttf", 70)
    font2 = ImageFont.truetype("./Fonts/msyhbd.ttf", 30)
    font3 = ImageFont.truetype("./Fonts/msyhbd.ttf", 36)
    font4 = ImageFont.truetype("./Fonts/renaissance.ttf", 46)

    # 加载背景图片
    backpng = './Images/Certificate-bg.png'
    im = Image.open(backpng)

    # 加载奖牌图片
    imMedal = Image.open(medal)
    # 将奖牌图片粘贴到背景图片的指定位置
    im.paste(imMedal, (784, 860), imMedal)
    imMedal.close()

    # 创建绘图对象
    dr = ImageDraw.Draw(im)

    # 在指定位置绘制证书编号
    dr.text((160, 132), f"No.{str(maxCertNum).rjust(5, '0')}", font=font4, fill='grey')

    # 根据用户中文名的长度确定姓名的绘制位置
    if 0 <= len(userCName.replace(" ", "")) - 1 <= 5:
        dr.text((namePosX[len(userCName.replace(" ", "")) - 1], 460), userCName, font=font, fill='grey')
    else:
        dr.text((460, 460), userCName, font=font, fill='grey')

    # 在指定位置绘制考试名称
    dr.text((900 - int(len(examName) * 15), 710), examName, font=font2, fill='grey')

    # 在指定位置绘制考试日期
    dr.text((410, 940), examDate, font=font3, fill='grey')

    # 保存生成的证书图片
    im.save(certFile)
    im.close()


def displayMedals():
    # 从数据库中查询考试名称，排除练习题库，按ID排序
    sql = "SELECT examName from examidd where examName <> '练习题库' order by ID"
    # 执行SQL查询并获取结果
    rows = execute_sql(cur, sql)
    for row in rows:
        # 使用st.expander创建可折叠区域，默认不展开
        with st.expander(label=f"{row[0]}", expanded=False):
            # 创建6个等宽的列
            mcol1, mcol2, mcol3, mcol4, mcol5, mcol6 = st.columns(6)
            # 构建SQL查询语句，获取考试成绩前三名的用户信息
            sql = f"SELECT userCName, examScore, examDate from examresult where examName = '{row[0]}' and examPass = 1 order by examScore DESC limit 0, 3"
            # 执行SQL查询并获取结果
            rows2 = execute_sql(cur, sql)
            if rows2:
                # 如果查询结果不为空
                if len(rows2) > 0:
                    # 格式化日期
                    examDate = time.strftime("%Y-%m-%d", time.localtime(rows2[0][2]))
                    # 在第一列显示金牌图片
                    mcol3.image("./Images/gold-medal.png")
                    # 在第二列显示第一名用户信息
                    mcol4.write(f"##### :red[{rows2[0][0]}]")
                    mcol4.write(f"成绩: {rows2[0][1]}分")
                    mcol4.write(f"{examDate}")
                if len(rows2) > 1:
                    # 格式化日期
                    examDate = time.strftime("%Y-%m-%d", time.localtime(rows2[1][2]))
                    # 在第三列显示银牌图片
                    mcol1.image("./Images/silver-medal.png")
                    # 在第四列显示第二名用户信息
                    mcol2.write(f"##### :grey[{rows2[1][0]}]")
                    mcol2.write(f"成绩: {rows2[1][1]}分")
                    mcol2.write(f"{examDate}")
                if len(rows2) > 2:
                    # 格式化日期
                    examDate = time.strftime("%Y-%m-%d", time.localtime(rows2[2][2]))
                    # 在第五列显示铜牌图片
                    mcol5.image("./Images/bronze-medal.png")
                    # 在第六列显示第三名用户信息
                    mcol6.write(f"##### :orange[{rows2[2][0]}]")
                    mcol6.write(f"成绩: {rows2[2][1]}分")
                    mcol6.write(f"{examDate}")


def displayErrorQues():
    sql = f"SELECT Question, qOption, qAnswer, qType, qAnalysis, userAnswer, ID, WrongTime from morepractise where userAnswer <> '' and qAnswer <> userAnswer and userName = {st.session_state.userName} order by WrongTime DESC"
    rows = execute_sql(cur, sql)
    if rows:
        for row in rows:
            #st.subheader("", divider="red")
            with st.expander(label=f"题目: {row[0]} 次数: {row[7]}", expanded=False):
                if row[3] == "单选题":
                    st.write(":red[标准答案:]")
                    option, userAnswer = [], ["A", "B", "C", "D"]
                    tmp = row[1].replace("；", ";").split(";")
                    for index, each in enumerate(tmp):
                        each = each.replace("\n", "").replace("\t", "").strip()
                        option.append(f"{userAnswer[index]}. {each}")
                    st.radio(" ", option, key=f"compare_{row[6]}", index=int(row[2]), horizontal=True, label_visibility="collapsed", disabled=True)
                    st.write(f"你的答案: :red[{userAnswer[int(row[5])]}] 你的选择为: :blue[错误]")
                elif row[3] == "多选题":
                    userOption = ["A", "B", "C", "D", "E", "F", "G", "H"]
                    st.write(":red[标准答案:]")
                    option = row[1].replace("；", ";").split(";")
                    orgOption = row[2].replace("；", ";").split(";")
                    for index, value in enumerate(option):
                        value = value.replace("\n", "").replace("\t", "").strip()
                        if str(index) in orgOption:
                            st.checkbox(f"{userOption[index]}. {value}:", value=True, disabled=True)
                        else:
                            st.checkbox(f"{userOption[index]}. {value}:", value=False, disabled=True)
                    userAnswer = row[5].replace("；", ";").split(";")
                    tmp = ""
                    for each in userAnswer:
                        tmp = tmp + userOption[int(each)] + ", "
                    st.write(f"你的答案: :red[{tmp[:-2]}] 你的选择为: :blue[错误]")
                elif row[3] == "判断题":
                    st.write(":red[标准答案:]")
                    option = ["A. 正确", "B. 错误"]
                    tmp = int(row[2]) ^ 1
                    st.radio(" ", option, key=f"compare_{row[6]}", index=tmp, horizontal=True, label_visibility="collapsed", disabled=True)
                    tmp = int(row[5]) ^ 1
                    st.write(f"你的答案: :red[{option[tmp]}] 你的选择为: :blue[错误]")
                elif row[3] == "填空题":
                    option = row[2].replace("；", ";").split(";")
                    userAnswer = row[5].replace("；", ";").split(";")
                    st.write(":red[标准答案:]")
                    for index, value in enumerate(option):
                        st.write(f"第{index + 1}个填空: :green[{value}]")
                    st.write("你的答案:")
                    for index, value in enumerate(userAnswer):
                        st.write(f"第{index + 1}个填空: :red[{value}]")
                    st.write("你的填写为: :blue[错误]")
                if row[4] != "":
                    if row[4].endswith("]"):
                        #st.write(row[4])
                        st.markdown(f"答案解析: {row[4][:-1]}]")
                    else:
                        st.markdown(f"答案解析: :green[{row[4]}]")
    else:
        st.info("暂无数据")


def studyReset():
    buttonSubmit = st.button("重置学习记录", type="primary")
    if buttonSubmit:
        st.button("确认重置", type="secondary", on_click=studyResetAction)


def studyResetAction():
    sql = f"DELETE from studyinfo where userName = {st.session_state.userName}"
    execute_sql_and_commit(conn, cur, sql)
    st.success("学习记录已重置")
    updateKeyAction("重置学习记录")


# noinspection PyTypeChecker
def studyinfoDetail():
    # 创建三列布局
    scol1, scol2, scol3 = st.columns(3)

    # 查询特定条件下的题目数量
    sql = f"SELECT Count(ID) from questionaff where StationCN = '{st.session_state.StationCN}' and chapterName <> '错题集' and chapterName <> '关注题集'"
    rows = execute_sql(cur, sql)
    # 显示章节总计
    scol1.metric(label="章节总计", value=rows[0][0], help="包含公共题库, 不含错题集")

    # 查询特定条件下的题目数量
    sql = f"SELECT Count(ID) from questions where StationCN = '{st.session_state.StationCN}'"
    ct1 = execute_sql(cur, sql)[0][0]
    # 查询公共题库中的题目数量
    sql = "SELECT Count(ID) from commquestions"
    ct2 = execute_sql(cur, sql)[0][0]
    # 计算总题目数量
    ct = ct1 + ct2
    # 显示试题总计
    scol2.metric(label="试题总计", value=ct, help="包含公共题库, 不含错题集")

    # 查询已学习的试题数量
    sql = f"SELECT Count(ID) from studyinfo where userName = {st.session_state.userName}"
    rows = execute_sql(cur, sql)
    # 显示已学习试题和完成率
    scol3.metric(label="已学习试题", value=f"{rows[0][0]} - {int(rows[0][0] / ct * 100)}%", help=f"总完成率: {int(rows[0][0] / ct * 100)}%")

    # 设置度量卡片的样式
    style_metric_cards(border_left_color="#8581d9")

    # 显示帮助信息
    helpInfo = ["点击页面⤴️右上角红圈处图标, 并选择Settings", "点击Choose app theme, colors and fonts", "选择Light或是Custom Theme"]
    st.write("###### :violet[如果上面3个标签无显示内容, 请按照以下步骤改用浅色主题]")

    # 显示操作步骤
    step = sac.steps(
        items=[
            sac.StepsItem(title='页面设置'),
            sac.StepsItem(title='主题设置'),
            sac.StepsItem(title='选择主题'),
        ], index=None, return_index=True
    )
    if step is not None:
        st.image(f"./Images/help/themesetup{step}.png", caption=f"{helpInfo[step]}")

    # 显示各章节进度详情
    with st.expander("各章节进度详情", icon=":material/format_list_bulleted:", expanded=True):
        # 查询公共题库中的题目数量
        sql = "SELECT Count(ID) from commquestions"
        ct = execute_sql(cur, sql)[0][0]
        if ct > 0:
            # 查询特定章节的已学习题目数量
            sql = f"SELECT Count(ID) from studyinfo where userName = {st.session_state.userName} and chapterName = '公共题库'"
            cs = execute_sql(cur, sql)[0][0]
            # 显示公共题库的完成进度
            st.progress(value=cs / ct, text=f":blue[公共题库] 已完成 :orange[{int((cs / ct) * 100)}%]")

        # 查询特定条件下的章节名称
        sql = f"SELECT chapterName from questionaff where StationCN = '{st.session_state.StationCN}' and chapterName <> '公共题库' and chapterName <> '错题集' order by ID"
        rows = execute_sql(cur, sql)
        for row in rows:
            # 查询特定章节的题目数量
            sql = f"SELECT Count(ID) from questions where StationCN = '{st.session_state.StationCN}' and chapterName = '{row[0]}'"
            ct = execute_sql(cur, sql)[0][0]
            if ct > 0:
                # 查询特定章节的已学习题目数量
                sql = f"SELECT Count(ID) from studyinfo where userName = {st.session_state.userName} and chapterName = '{row[0]}'"
                cs = execute_sql(cur, sql)[0][0]
                # 显示各章节的完成进度
                st.progress(value=cs / ct, text=f":blue[{row[0]}] 已完成 :orange[{int((cs / ct) * 100)}%]")


def userStatus():
    # 设置子标题和分隔线颜色
    st.subheader(":violet[在线用户状态]", divider="green")

    # 判断是否需要重新检查用户密码
    if st.session_state.userPwRecheck:
        # 创建分段选择组件
        bc = sac.segmented(
            items=[
                # 在线用户状态选项
                sac.SegmentedItem(label="在线用户状态", icon="people"),
                # 重置所有用户状态选项
                sac.SegmentedItem(label="重置所有用户状态", icon="person-slash"),
            ], align="start", color="red"
        )

        # 判断用户选择的是哪个选项
        if bc == "在线用户状态":
            actionUserStatus()
        elif bc == "重置所有用户状态":
            # 创建重置所有用户状态的按钮
            buttonReset = st.button("重置所有用户状态", type="primary")
            if buttonReset:
                # 确认重置按钮
                st.button("确认重置", type="secondary", on_click=resetActiveUser)

        # 如果用户选择了某个选项，则更新用户状态
        if bc is not None:
            updateActionUser(st.session_state.userName, bc, st.session_state.loginTime)
    else:
        # 如果需要重新检查密码，则显示密码输入框
        vUserPW = st.text_input("请输入密码", max_chars=8, placeholder="请输入管理员密码, 以验证身份", type="password", autocomplete="off")

        # 判断用户是否输入了密码
        if vUserPW:
            # 验证用户输入的密码是否正确
            if verifyUserPW(st.session_state.userName, vUserPW)[0]:
                # 如果密码正确，重新运行当前函数
                st.rerun()
            else:
                # 如果密码错误，显示错误信息
                st.error("密码错误, 请重新输入")


def actionUserStatus():
    # SQL查询语句，获取活跃用户的信息
    sql = "SELECT userCName, userType, StationCN, actionUser, loginTime, activeTime_session, activeTime from users where activeUser = 1 order by loginTime desc, activeTime_session desc, activeTime desc, ID"
    # 执行SQL查询
    rows = execute_sql(cur, sql)
    # 将查询结果转换为DataFrame
    df = pd.DataFrame(rows, dtype=str)
    # 设置DataFrame的列名
    df.columns = ["姓名", "类型", "站室", "用户操作", "登录时间", "活动时间", "累计活动时间"]

    # 遍历查询结果
    for index, value in enumerate(rows):
        # 将登录时间转换为可读格式
        df.loc[index, "登录时间"] = time.strftime("%Y-%m-%d %H:%M:%S", time.localtime(int(df["登录时间"][index])))

        # 获取活动时间并转换为小时、分钟、秒格式
        activeTime = int(df.loc[index, "活动时间"])
        hTime = int(activeTime / 3600)
        mTime = int((activeTime % 3600) / 60)
        if mTime < 10:
            mTime = "0" + str(mTime)
        sTime = int(activeTime % 60)
        if sTime < 10:
            sTime = "0" + str(sTime)
        df.loc[index, "活动时间"] = f"{hTime}小时{mTime}分{sTime}秒"

        # 获取累计活动时间并转换为小时、分钟、秒格式
        activeTime = int(df.loc[index, "累计活动时间"])
        hTime = int(activeTime / 3600)
        mTime = int((activeTime % 3600) / 60)
        if mTime < 10:
            mTime = "0" + str(mTime)
        sTime = int(activeTime % 60)
        if sTime < 10:
            sTime = "0" + str(sTime)
        df.loc[index, "累计活动时间"] = f"{hTime}小时{mTime}分{sTime}秒"

    # 使用Streamlit显示DataFrame
    st.dataframe(df, use_container_width=True)


@st.fragment
def actionQuesModify(row):
    option = []
    if len(row) == 8:
        qQuestion, qOption, qAnswer, qType, qAnalysis, StationCN, chapterName, SourceType = row
    else:
        qQuestion, qOption, qAnswer, qType, qAnalysis, SourceType = row
    st.session_state.qModifyQues_qType = qType
    st.write(f"**此题为{qType}**")
    st.text_area(":blue[**题目**]", value=qQuestion, key="qModifyQues_Question")
    if qType == "单选题":
        qOption2 = qOption.split(";")
        st.session_state.qModifyQues_optionCount = len(qOption2)
        for index, value in enumerate(qOption2):
            st.text_input(f":orange[**选项{chr(65 + index)}**]", value=value, key=f"qModifyQues_{index}")
            option.append(chr(65 + index))
        st.radio(":red[**答案**]", options=option, index=int(qAnswer), key="qModifyQues_Answer", horizontal=True)
    elif qType == "多选题":
        qOption2 = qOption.split(";")
        qAnswer2 = qAnswer.split(";")
        st.session_state.qModifyQues_optionCount = len(qOption2)
        for index, value in enumerate(qOption2):
            st.text_input(f":orange[**选项{chr(65 + index)}**]", value=value, key=f"qModifyQues_{index}")
            if str(index) in qAnswer2:
                st.checkbox(":blue[**选择**]", value=True, key=f"qModifyQues_Answer_{index}")
            else:
                st.checkbox(":blue[**选择**]", value=False, key=f"qModifyQues_Answer_{index}")
    elif qType == "判断题":
        st.radio(":red[**答案**]", ["A. 正确", "B. 错误"], key="qModifyQues_Answer", index=int(qAnswer) ^ 1, horizontal=True)
    elif qType == "填空题":
        qAnswer2 = qAnswer.split(";")
        st.session_state.qModifyQues_optionCount = len(qAnswer2)
        for index, value in enumerate(qAnswer2):
            st.text_input(":orange[**答案**]", value=value, key=f"qModifyQues_Answer_{index}")
    st.text_area(":green[**答案解析**]", value=qAnalysis, key="qModifyQues_Answer_Analysis")


def quesModify():
    st.subheader(":green[试题修改]", divider="blue")
    col1, col2 = st.columns(2)
    chosenTable = col1.selectbox(":red[选择题库]", ["站室题库", "公共题库"], index=None)
    quesID = col2.number_input(":blue[题目ID]", min_value=0, step=1)
    if chosenTable is not None and quesID > 0:
        if chosenTable == "站室题库":
            tablename = "questions"
        elif chosenTable == "公共题库":
            tablename = "commquestions"
        else:
            tablename = ""
        col3, col4, col5, col6 = st.columns(4)
        buttonDisplayQues = col3.button("显示试题", icon=":material/dvr:")
        if buttonDisplayQues:
            if chosenTable == "站室题库":
                sql = f"SELECT Question, qOption, qAnswer, qType, qAnalysis, StationCN, chapterName, SourceType from {tablename} where ID = {quesID}"
            else:
                sql = f"SELECT Question, qOption, qAnswer, qType, qAnalysis, SourceType from {tablename} where ID = {quesID}"
            rows = execute_sql(cur, sql)
            if rows:
                if chosenTable == "站室题库":
                    st.write(f":green[站室: {rows[0][5]} 章节: {rows[0][6]} 试题来源: {rows[0][7]}]")
                else:
                    st.write(f":green[公共题库 试题来源: {rows[0][5]}]")
                col4.button("更新试题", on_click=actionQM, args=(quesID, tablename, rows[0]), icon=":material/published_with_changes:")
                col5.button("删除试题", on_click=actionDelQM, args=(quesID, tablename, rows[0]), icon=":material/delete:")
                if chosenTable == "站室题库":
                    col6.button("移至公共题库", on_click=moveQM, args=(quesID, tablename, rows[0]), icon=":material/move_item:")
                actionQuesModify(rows[0])
            else:
                st.error("未找到该题目, 请检查题库名称及题目ID是否正确")
    else:
        st.error("请选择题库")


def moveQM(quesID, tablename, mRow):
    sql = f"DELETE from {tablename} where ID = {quesID}"
    execute_sql_and_commit(conn, cur, sql)
    sql = f"INSERT INTO commquestions(Question, qOption, qAnswer, qType, qAnalysis, SourceType) VALUES('{mRow[0]}', '{mRow[1]}', '{mRow[2]}', '{mRow[3]}', '{mRow[4]}', '{mRow[7]}')"
    execute_sql_and_commit(conn, cur, sql)
    st.toast("试题移至公共题库成功")


def actionQM(quesID, tablename, mRow):
    mOption, mAnswer, Option = "", "", ["A", "B", "C", "D", "E", "F", "G", "H"]
    mQues = st.session_state.qModifyQues_Question
    mAnalysis = st.session_state.qModifyQues_Answer_Analysis
    if st.session_state.qModifyQues_qType == "单选题" or st.session_state.qModifyQues_qType == "多选题":
        for i in range(st.session_state.qModifyQues_optionCount):
            mOption = mOption + st.session_state[f"qModifyQues_{i}"] + ";"
        if mOption.endswith(";"):
            mOption = mOption[:-1]
        if st.session_state.qModifyQues_qType == "单选题":
            for index, value in enumerate(Option):
                if value == st.session_state.qModifyQues_Answer:
                    mAnswer = index
                    break
        else:
            for i in range(st.session_state.qModifyQues_optionCount):
                if st.session_state[f"qModifyQues_Answer_{i}"]:
                    mAnswer = mAnswer + str(i) + ";"
            if mAnswer.endswith(";"):
                mAnswer = mAnswer[:-1]
    elif st.session_state.qModifyQues_qType == "判断题":
        if "正确" in st.session_state.qModifyQues_Answer:
            mAnswer = 1
        else:
            mAnswer = 0
    elif st.session_state.qModifyQues_qType == "填空题":
        for i in range(st.session_state.qModifyQues_optionCount):
            mAnswer = mAnswer + st.session_state[f"qModifyQues_Answer_{i}"] + ";"
        if mAnswer.endswith(";"):
            mAnswer = mAnswer[:-1]
    sql = f"UPDATE {tablename} set Question = '{mQues}', qOption = '{mOption}', qAnswer = '{mAnswer}', qAnalysis = '{mAnalysis}' where ID = {quesID}"
    execute_sql_and_commit(conn, cur, sql)
    clearModifyQues(quesID, tablename, mRow)
    for key in st.session_state.keys():
        if key.startswith("qModifyQues_"):
            del st.session_state[key]
    st.toast("试题修改成功")


def actionDelQM(quesID, tablename, mRow):
    sql = f"DELETE from {tablename} where ID = {quesID}"
    execute_sql_and_commit(conn, cur, sql)
    clearModifyQues(quesID, tablename, mRow)
    for key in st.session_state.keys():
        if key.startswith("qModifyQues_"):
            del st.session_state[key]
    st.toast("试题删除成功")


def clearModifyQues(quesID, tablename, mRow):
    delTablePack = ["morepractise", "favques"]
    for each in delTablePack:
        sql = f"DELETE from {each} where Question = '{mRow[0]}' and qOption = '{mRow[1]}' and qAnswer = '{mRow[2]}' and qType = '{mRow[3]}'"
        execute_sql_and_commit(conn, cur, sql)
    sql = f"DELETE from studyinfo where cid = {quesID} and quesTable = '{tablename}'"
    execute_sql_and_commit(conn, cur, sql)


def aboutReadme():
    st.markdown(open("./README.md", "r", encoding="utf-8").read())


# noinspection PyUnboundLocalVariable
def training():
    flagProc, failInfo = True, ""
    if st.session_state.examType == "exam":
        quesType = []
        sql = f"SELECT paramName from setup_{st.session_state.StationCN} where paramType = 'questype' and param = 1 order by ID"
        rows = execute_sql(cur, sql)
        for row in rows:
            quesType.append([row[0], getParam(f"{row[0]}数量", st.session_state.StationCN)])
        for each in quesType:
            quesTypeCount = 0
            tmp = each[0].replace("数量", "")
            sql = f"SELECT count(ID) from questions where qType = '{tmp}' and StationCN = '{st.session_state.StationCN}'"
            quesTypeCount += int(execute_sql(cur, sql)[0][0])
            sql = f"SELECT count(ID) from commquestions where qType = '{tmp}'"
            quesTypeCount += int(execute_sql(cur, sql)[0][0])
            if quesTypeCount < each[1]:
                flagProc = False
                failInfo = failInfo + f"{tmp}/"
    elif st.session_state.examType == "training":
        quesType = [["单选题", 30], ["多选题", 10], ["判断题", 10], ["填空题", 0]]
        sql = f"SELECT mcq, mmcq, tfq, fibq from indivquescount where userName = {st.session_state.userName}"
        rows = execute_sql(cur, sql)
        if rows:
            row = rows[0]
            for i in range(4):
                quesType[i][1] = row[i]
        else:
            sql = f"INSERT INTO indivquescount (userName, mcq, mmcq, tfq, fibq) VALUES({st.session_state.userName}, {quesType[0][1]}, {quesType[1][1]}, {quesType[2][1]}, {quesType[3][1]})"
            execute_sql_and_commit(conn, cur, sql)
    if flagProc:
        generPack, examIDPack, chapterPack, genResult = [], [], [], []
        generQues = st.empty()
        with generQues.container():
            if st.session_state.examType == "exam":
                date = int(time.time())
                sql = f"SELECT examName from examidd where StationCN = '{st.session_state.StationCN}' and validDate >= {date} order by validDate"
                rows = execute_sql(cur, sql)
                for row in rows:
                    examIDPack.append(row[0])
                examName = st.selectbox("请选择考试场次", examIDPack, index=None)
                if examName:
                    generButtonQues = st.button("开始考试")
                    if generButtonQues:
                        st.session_state.examName = examName
                        st.spinner("正在生成题库...")
                        reviseQues()
                        sql = "SELECT chapterName from questionaff where chapterName <> '错题集' and chapterName <> '关注题集' and StationCN = '" + st.session_state.StationCN + "'"
                        rows = execute_sql(cur, sql)
                        for row in rows:
                            chapterPack.append(row[0])
                        genResult = GenerExam(chapterPack, st.session_state.StationCN, st.session_state.userName, st.session_state.examName, st.session_state.examType, quesType, st.session_state.examRandom, False)
            elif st.session_state.examType == "training":
                tCol1, tCol2, tCol3 = st.columns(3)
                generButtonQues = tCol1.button("生成题库")
                sql = "SELECT pyLM from verinfo where pyFile = 'chapterChosenType'"
                chapterChosenType = execute_sql(cur, sql)[0][0]
                with tCol2:
                    uCCT = sac.segmented(
                        items=[
                            sac.SegmentedItem(label="默认"),
                            sac.SegmentedItem(label="全选"),
                            sac.SegmentedItem(label="全不选"),
                        ], index=chapterChosenType, align="start", color="orange", return_index=True, size="sm",
                    )
                if uCCT != 0:
                    sql = f"UPDATE verinfo set pyLM = {uCCT} where pyFile = 'chapterChosenType'"
                    execute_sql_and_commit(conn, cur, sql)
                tCol3.checkbox(":red[**仅未学习试题**]", value=False, key="GenerNewOnly", help="仅从未学习试题中生成")
                indivCols = st.columns(4)
                for i in range(4):
                    quesType[i][1] = indivCols[i].number_input(quesType[i][0], min_value=0, max_value=100, value=quesType[i][1], step=1)
                ddCol1, ddCol2 = st.columns(2)
                ddCol1.write("**章节**")
                ddCol2.write("**权重**")
                for each in ["公共题库", "错题集", "关注题集"]:
                    ddCol1, ddCol2 = st.columns(2)
                    sql = f"SELECT chapterName, chapterRatio, ID from questionaff where StationCN = '{st.session_state.StationCN}' and chapterName = '{each}'"
                    row = execute_sql(cur, sql)[0]
                    if uCCT == 0:
                        if each == "公共题库":
                            generPack.append(ddCol1.checkbox(f"**:blue[{row[0]}]**", value=True))
                        else:
                            generPack.append(ddCol1.checkbox(f"**:blue[{row[0]}]**", value=False))
                    elif uCCT == 1:
                        generPack.append(ddCol1.checkbox(f"**:blue[{row[0]}]**", value=True))
                    elif uCCT == 2:
                        generPack.append(ddCol1.checkbox(f"**:blue[{row[0]}]**", value=False))
                    ddCol2.slider("章节权重", min_value=1, max_value=10, value=row[1], step=1, key=f"tempCR_{row[2]}", on_change=updateCRTraining, label_visibility="collapsed")
                sql = "SELECT chapterName, chapterRatio, ID from questionaff where StationCN = '" + st.session_state.StationCN + "' and chapterName <> '公共题库' and chapterName <> '错题集' and chapterName <> '关注题集' order by chapterName"
                rows = execute_sql(cur, sql)
                for row in rows:
                    ddCol1, ddCol2 = st.columns(2)
                    if uCCT == 0 or uCCT == 1:
                        generPack.append(ddCol1.checkbox(f"**:blue[{row[0]}]**", value=True))
                    elif uCCT == 2:
                        generPack.append(ddCol1.checkbox(f"**:blue[{row[0]}]**", value=False))
                    ddCol2.slider("章节权重", min_value=1, max_value=10, value=row[1], step=1, key=f"tempCR_{row[2]}", on_change=updateCRTraining, label_visibility="collapsed")
                if generButtonQues:
                    st.session_state.examName = "练习题库"
                    sql = f"UPDATE indivquescount set mcq = {quesType[0][1]}, mmcq = {quesType[1][1]}, tfq = {quesType[2][1]}, fibq = {quesType[3][1]} where userName = {st.session_state.userName}"
                    execute_sql_and_commit(conn, cur, sql)
                    for index, value in enumerate(generPack):
                        if value:
                            if index == 0:
                                chapterPack.append("公共题库")
                            elif index == 1:
                                chapterPack.append("错题集")
                            elif index == 2:
                                chapterPack.append("关注题集")
                            else:
                                chapterPack.append(rows[index - 3][0])
                    if chapterPack:
                        st.spinner("正在生成题库...")
                        reviseQues()
                        genResult = GenerExam(chapterPack, st.session_state.StationCN, st.session_state.userName, st.session_state.examName, st.session_state.examType, quesType, st.session_state.examRandom, st.session_state.GenerNewOnly)
                    else:
                        st.warning("请至少选择一个章节")
        if genResult:
            if genResult[0]:
                generQues.empty()
                if st.session_state.examType == "exam":
                    st.success(f"题库生成完毕, 总共生成{genResult[1]}道试题, 请在👈左侧边栏选择开始考试")
                else:
                    st.success(f"题库生成完毕, 总共生成{genResult[1]}道试题, 请在👈左侧边栏选择题库练习")
                st.session_state.examTable = genResult[2]
                st.session_state.examFinalTable = genResult[3]
                st.session_state.curQues = 0
                st.session_state.examStartTime = int(time.time())
                st.session_state.confirmSubmit = False
                st.session_state.flagCompleted = False
                st.session_state.goto = False
                st.session_state.radioCompleted = False
                st.session_state.calcScore = False
                st.session_state.delExam = False
                st.session_state.trainingID = str(time.strftime('%Y%m%d%H%M%S', time.localtime(int(time.time()))))
                if st.session_state.examType != "training":
                    st.session_state.examChosen = True
                    updateActionUser(st.session_state.userName, "生成考试试题", st.session_state.loginTime)
                else:
                    st.session_state.examChosen = False
                    updateActionUser(st.session_state.userName, "生成练习试题", st.session_state.loginTime)
            else:
                st.session_state.examChosen = False
                st.error("题库生成试题不满足要求, 请检查考试参数设置, 或个别题型试题候选数量不够, 或请联系管理员")
    else:
        st.error(f":red[⚠️] **{st.session_state.StationCN}试卷生成失败, :red[{failInfo[:-1]}] 试题数量不足, 请检查题库设置或增加以上题型候选试题**")


def reviseQues():
    for each in ["questions", "commquestions"]:
        for each2 in [['（', '('], ['）', ')']]:
            sql = f"UPDATE {each} set Question = replace(Question, '{each2[0]}', '{each2[1]}') where qType = '填空题' and Question like '%{each2[0]}%'"
            execute_sql_and_commit(conn, cur, sql)
        for each2 in ['( )', '(  )', '(   )', '(    )']:
            sql = f"UPDATE {each} set Question = replace(Question, '{each2}', '()') where qType = '填空题' and Question like '%{each2}'"
            execute_sql_and_commit(conn, cur, sql)


@st.fragment
def updateCRTraining():
    for key in st.session_state.keys():
        if key.startswith("tempCR_"):
            upID = key[key.find("_") + 1:]
            sql = f"UPDATE questionaff SET chapterRatio = {st.session_state[key]} WHERE ID = {upID}"
            execute_sql_and_commit(conn, cur, sql)


def updateCRExam():
    for key in st.session_state.keys():
        if key.startswith("crsetup_"):
            upID = key[key.find("_") + 1:]
            sql = f"UPDATE questionaff SET examChapterRatio = {st.session_state[key]} WHERE ID = {upID}"
            execute_sql_and_commit(conn, cur, sql)
    st.success("章节权重更新成功")
    updateKeyAction("考试章节权重更新")


@st.fragment
def updateAnswer(userQuesID):
    sql = f"UPDATE {st.session_state.examFinalTable} set userAnswer = '{st.session_state.answer}', userName = {st.session_state.userName} where ID = {userQuesID}"
    execute_sql_and_commit(conn, cur, sql)
    sql = f"SELECT Question, qAnswer, qType, userAnswer, userName, qOption, qAnalysis, SourceType from {st.session_state.examFinalTable} where ID = {userQuesID}"
    judTable = execute_sql(cur, sql)[0]
    if judTable[1] == judTable[3]:
        sql = f"UPDATE morepractise set WrongTime = WrongTime - 1 where Question = '{judTable[0]}' and qType = '{judTable[2]}' and userName = {judTable[4]}"
        execute_sql_and_commit(conn, cur, sql)
        sql = f"UPDATE users set userRanking = userRanking + 1 where userName = {st.session_state.userName}"
        execute_sql_and_commit(conn, cur, sql)
        #st.session_state.tooltipColor = "#ed872d"
    else:
        sql = f"SELECT ID from morepractise where Question = '{judTable[0]}' and qType = '{judTable[2]}' and userName = {judTable[4]}"
        if execute_sql(cur, sql):
            sql = f"UPDATE morepractise set WrongTime = WrongTime + 1, userAnswer = '{judTable[3]}' where Question = '{judTable[0]}' and qType = '{judTable[2]}' and userName = {judTable[4]} and trainingID <> '{st.session_state.trainingID}'"
            execute_sql_and_commit(conn, cur, sql)
        else:
            sql = f"INSERT INTO morepractise(Question, qOption, qAnswer, qType, qAnalysis, userAnswer, userName, WrongTime, StationCN, SourceType, trainingID) VALUES('{judTable[0]}', '{judTable[5]}', '{judTable[1]}', '{judTable[2]}', '{judTable[6]}', '{judTable[3]}', {judTable[4]}, 1, '{st.session_state.StationCN}', '{judTable[7]}', '{st.session_state.trainingID}')"
            execute_sql_and_commit(conn, cur, sql)
        #st.session_state.tooltipColor = "#8581d9"
    execute_sql_and_commit(conn, cur, sql="DELETE from morepractise where WrongTime < 1")


@st.dialog("考试成绩")
def score_dialog(userScore, passScore):
    examDate = int(time.mktime(time.strptime(time.strftime("%Y-%m-%d %H:%M:%S", time.localtime()), "%Y-%m-%d %H:%M:%S")))
    if userScore >= passScore:
        flagPass = 1
    else:
        flagPass = 0
    st.write(f"考生ID:  {st.session_state.userName}")
    st.write(f"考生姓名: {st.session_state.userCName}")
    st.write(f"考试时间: {time.strftime('%Y-%m-%d %H:%M:%S', time.localtime(examDate))}")
    st.subheader(f"考试成绩: {userScore} 分 / 合格分数线为 {passScore} 分")
    if flagPass == 1:
        st.subheader("考试结果: :blue[通过] 👏")
        st.balloons()
    else:
        st.subheader("考试结果: :red[未通过] 🤪")
        #st.snow()
    if st.session_state.examType == "training":
        st.write("练习模式成绩不计入记录")
    if st.session_state.examType == "exam" and st.session_state.calcScore:
        sql = "INSERT INTO examresult(userName, userCName, examScore, examDate, examPass, examName) VALUES(" + str(st.session_state.userName) + ", '" + st.session_state.userCName + "', " + str(userScore) + ", " + str(examDate) + ", " + str(flagPass) + ", '" + st.session_state.examName + "')"
        execute_sql_and_commit(conn, cur, sql)
    st.session_state.calcScore = False
    buttonScore = st.button("确定")
    if buttonScore:
        st.session_state.delExam = True
        if st.session_state.examType == "exam" and st.session_state.calcScore:
            logout()
        else:
            st.rerun()


def calcScore():
    st.session_state.examStartTime = int(time.time())
    st.session_state.confirmSubmit = True
    st.session_state.curQues = 0
    st.session_state.flagCompleted = False
    flagUseAIFIB = bool(getParam("使用大模型评判错误的填空题答案", st.session_state.StationCN))
    quesScore = getParam("单题分值", st.session_state.StationCN)
    passScore = getParam("合格分数线", st.session_state.StationCN)
    userScore = 0
    sql = f"SELECT qAnswer, qType, userAnswer, Question, qOption, qAnalysis, userName, SourceType from {st.session_state.examFinalTable} where userName = {st.session_state.userName} order by ID"
    rows = execute_sql(cur, sql)
    # 练习模式按照生成试题的80%作为合格分线, 考试模式读取管理员设置的分数线
    if st.session_state.examType == "training":
        passScore = int(len(rows) * 0.8)
    for row in rows:
        # [使用大模型评判错误的填空题答案] 该模块不稳定, 强制不使用
        flagAIScore = False
        if row[0].replace(" ", "").lower() == row[2].replace(" ", "").lower():
            userScore += quesScore
            sql = f"UPDATE users set userRanking = userRanking + 1 where userName = {st.session_state.userName}"
            execute_sql_and_commit(conn, cur, sql)
            sql = f"SELECT ID from morepractise where Question = '{row[3]}' and qType = '{row[1]}' and userName = {row[6]}"
            if execute_sql(cur, sql):
                sql = f"UPDATE morepractise set WrongTime = WrongTime - 1 where Question = '{row[3]}' and qType = '{row[1]}' and userName = {row[6]}"
                execute_sql_and_commit(conn, cur, sql)
            execute_sql_and_commit(conn, cur, sql="DELETE from morepractise where WrongTime < 1")
        else:
            if row[1] == "填空题":
                if flagUseAIFIB:
                    fibQues = row[3]
                    fibQues2 = row[3]
                    userAP = row[2].split(";")
                    quesAP = row[0].split(";")
                    if fibQues.count("()") == len(userAP):
                        #st.toast("正在使用 :red[讯飞星火大模型] 对答案进行分析, 请稍等...")
                        for index, value in enumerate(userAP):
                            b1 = fibQues.find("()")
                            c1 = fibQues2.find("()")
                            if b1 != -1:
                                fibQues = f"{fibQues[:b1]}({value}){fibQues[b1 + 2:]}"
                                fibQues2 = f"{fibQues2[:c1]}({quesAP[index]}){fibQues2[c1 + 2:]}"
                        fibAI = xunfei_xh_AI_fib(fibQues, fibQues2)
                        if fibAI != "" and fibAI.find("无法直接回答") == -1 and fibAI.find("尚未查询") == -1 and fibAI.find("我不确定您想要表达什么意思") == -1 and fibAI.find("由于信息不足，无法给出准确答案") == -1 and fibAI.find("无法确定正确答案") == -1 and fibAI.find("无法提供准确答案") == -1:
                            if st.session_state.debug:
                                print(f"debug: [{row[3]}] [Q:{row[0]} / A:{row[2]}] / A.I.判断: [{fibAI}]")
                            if fibAI == "正确":
                                userScore += quesScore
                                sql = f"UPDATE users set userRanking = userRanking + 1 where userName = {st.session_state.userName}"
                                execute_sql_and_commit(conn, cur, sql)
                                flagAIScore = True
                            else:
                                flagAIScore = False
                    else:
                        st.error("⚠️ 试题或是答案数量不匹配, 请检查")
            if not flagAIScore:
                sql = f"SELECT ID from morepractise where Question = '{row[3]}' and qType = '{row[1]}' and userName = {row[6]}"
                if not execute_sql(cur, sql):
                    sql = f"INSERT INTO morepractise(Question, qOption, qAnswer, qType, qAnalysis, userAnswer, userName, WrongTime, StationCN, SourceType, trainingID) VALUES('{row[3]}', '{row[4]}', '{row[0]}', '{row[1]}', '{row[5]}', '{row[2]}', {row[6]}, 1, '{st.session_state.StationCN}', '{row[7]}', '{st.session_state.trainingID}')"
                    execute_sql_and_commit(conn, cur, sql)
                else:
                    sql = f"UPDATE morepractise set WrongTime = WrongTime + 1, userAnswer = '{row[2]}' where Question = '{row[3]}' and qType = '{row[1]}' and userName = {row[6]} and trainingID <> '{st.session_state.trainingID}'"
                    execute_sql_and_commit(conn, cur, sql)
    if st.session_state.calcScore:
        score_dialog(userScore, passScore)


@st.fragment
def updateOptionAnswer(chosenID, chosen, option):
    for index, value in enumerate(option):
        if chosen == value:
            st.session_state.answer = index
    updateAnswer(chosenID)


@st.fragment
def updateRadioAnswer(chosenID):
    if st.session_state.radioChosen is not None:
        if "正确" in st.session_state.radioChosen:
            st.session_state.answer = 1
        else:
            st.session_state.answer = 0
        st.session_state.radioCompleted = True
        updateAnswer(chosenID)


@st.fragment
def updateRadioAnswer2(chosenID):
    if st.session_state.radioChosen2 is not None:
        if "正确" in st.session_state.radioChosen2:
            st.session_state.answer = 1
        else:
            st.session_state.answer = 0
        updateAnswer(chosenID)


@st.fragment
def updateMOptionAnswer(row):
    mpAnswerPack = []
    for key in st.session_state.keys():
        if key.startswith("moption_"):
            if st.session_state[key]:
                mpAnswerPack.append(int(key.replace("moption_", "")))
    mpAnswerPack.sort()
    st.session_state.answer = ""
    for each in mpAnswerPack:
        st.session_state.answer = st.session_state.answer + str(each) + ";"
    if st.session_state.answer.endswith(";"):
        st.session_state.answer = st.session_state.answer[:-1]
    updateAnswer(row[0])


def delQuestion(delQuesRow):
    delTablePack = ["questions", "commquestions", "morepractise", "favques"]
    for delTable in delTablePack:
        sql = f"DELETE from {delTable} where Question = '{delQuesRow[1]}' and qOption = '{delQuesRow[2]}' and qType = '{delQuesRow[4]}'"
        execute_sql_and_commit(conn, cur, sql)
    updateKeyAction(f"删除试题: {delQuesRow[1]}")


@st.fragment
def updateStudyInfo(studyRow):
    # 遍历题目和公共题目
    for each in ["questions", "commquestions"]:
        # 如果是题目
        if each == "questions":
            # 构建查询SQL语句
            sql = f"SELECT ID, chapterName from {each} where Question = '{studyRow[1]}' and qType = '{studyRow[4]}' and StationCN = '{st.session_state.StationCN}'"
        # 如果是公共题目
        elif each == "commquestions":
            # 构建查询SQL语句
            sql = f"SELECT ID, '公共题库' from {each} where Question = '{studyRow[1]}' and qType = '{studyRow[4]}'"
        # 其他情况
        else:
            # SQL语句为空
            sql = ""
        # 执行SQL语句
        studyResult = execute_sql(cur, sql)
        # 如果查询结果不为空
        if studyResult:
            # 构建查询SQL语句
            sql = f"SELECT ID from studyinfo where cid = {studyResult[0][0]} and questable = '{each}' and userName = {st.session_state.userName} and chapterName = '{studyResult[0][1]}'"
            # 如果查询结果为空
            if not execute_sql(cur, sql):
                # 构建插入SQL语句
                sql = f"INSERT INTO studyinfo(cid, questable, userName, userCName, chapterName, startTime) VALUES({studyResult[0][0]}, '{each}', {st.session_state.userName}, '{st.session_state.userCName}', '{studyResult[0][1]}', {int(time.time())})"
                # 执行SQL语句并提交
                execute_sql_and_commit(conn, cur, sql)


@st.fragment
def delFavQues(favRow):
    sql = f"DELETE from favques where Question = '{favRow[1]}' and userName = {st.session_state.userName} and qType = '{favRow[4]}' and StationCN = '{st.session_state.StationCN}'"
    execute_sql_and_commit(conn, cur, sql)
    st.toast("已从关注题集中删除")


@st.fragment
def addFavQues(favRow):
    sql = f"SELECT ID from favques where Question = '{favRow[1]}' and userName = {st.session_state.userName} and StationCN = '{st.session_state.StationCN}'"
    if not execute_sql(cur, sql):
        sql = f"INSERT INTO favques(Question, qOption, qAnswer, qType, qAnalysis, userName, StationCN, SourceType) VALUES('{favRow[1]}', '{favRow[2]}', '{favRow[3]}', '{favRow[4]}', '{favRow[5]}', {st.session_state.userName}, '{st.session_state.StationCN}', '{favRow[8]}')"
        execute_sql_and_commit(conn, cur, sql)
        st.toast("已添加到关注题集")


# noinspection PyUnboundLocalVariable
@st.fragment
def exam(row):
    option, AIModelName, AIOption, AIOptionIndex = [], "", [], 0
    st.session_state.answer = ""
    flagAIUpdate = bool(getParam("A.I.答案解析更新至题库", st.session_state.StationCN))
    sql = f"SELECT paramName, param from setup_{st.session_state.StationCN} where paramType = 'others' and paramName like '%大模型' order by ID"
    tempTable = execute_sql(cur, sql)
    for index, value in enumerate(tempTable):
        AIOption.append(value[0])
        if value[1] == 1:
            AIOptionIndex = index
    if row[4] == "填空题":
        reviseQues = row[1].replace("(", ":red[ ( _ ]").replace(")", ":red[ _ _ ) ]").strip()
    else:
        reviseQues = row[1].replace("( )", "").strip()
    standardAnswer = getStandardAnswer(row)
    if st.session_state.examType != "exam":
        updateStudyInfo(row)
    st.markdown(f"##### 第{row[0]}题 :green[{reviseQues}]")
    acol = st.columns(2)
    if st.session_state.userType == "admin" and st.session_state.examType != "exam" and st.session_state.debug:
        addFavIndex = 1
        buttonConfirm = acol[0].button("⚠️ 从所有题库中删除此题", type="primary")
        if buttonConfirm:
            st.button("确认删除", type="secondary", on_click=delQuestion, args=(row,))
    else:
        addFavIndex = 0
    if st.session_state.examType == "training":
        sql = f"SELECT ID from favques where Question = '{row[1]}' and userName = {st.session_state.userName} and StationCN = '{st.session_state.StationCN}'"
        if execute_sql(cur, sql):
            acol[addFavIndex].button(label="", icon=":material/heart_minus:", on_click=delFavQues, args=(row,), help="从关注题集中删除")
        else:
            acol[addFavIndex].button(label="", icon=":material/heart_plus:", on_click=addFavQues, args=(row,), help="添加到关注题集")
    st.write(f":red[本题为{row[4]}]:")
    if row[4] == '单选题':
        for index, value in enumerate(row[2].replace("；", ";").split(";")):
            value = value.replace("\n", "").replace("\t", "").strip()
            option.append(f"{chr(65 + index)}. {value}")
        if row[6] == "" or row[6] is None:
            chosen = st.radio(" ", option, index=None, label_visibility="collapsed", horizontal=True)
        else:
            chosen = st.radio(" ", option, index=int(row[6]), label_visibility="collapsed", horizontal=True)
            #st.write(f":red[你已选择: ] :blue[{option[int(row[6])]}]")
        if chosen is not None:
            updateOptionAnswer(row[0], chosen, option)
    elif row[4] == '多选题':
        for index, value in enumerate(row[2].replace("；", ";").split(";")):
            value = value.replace("\n", "").replace("\t", "").strip()
            option.append(f"{chr(65 + index)}. {value}")
        if row[6] != "" and row[6] is not None:
            orgOption = row[6].replace("；", ";").split(";")
        else:
            orgOption = []
        for index, value in enumerate(option):
            if str(index) in orgOption:
                st.checkbox(f"{value}", value=True, key=f"moption_{index}", on_change=updateMOptionAnswer, args=(row,))
            else:
                st.checkbox(f"{value}", value=False, key=f"moption_{index}", on_change=updateMOptionAnswer, args=(row,))
    elif row[4] == '判断题':
        radioArea = st.empty()
        with radioArea.container():
            option = ["A. 正确", "B. 错误"]
            if row[6] != "" and row[6] is not None:
                st.radio(" ", option, index=int(row[6]) ^ 1, key="radioChosen", on_change=updateRadioAnswer, args=(row[0],), label_visibility="collapsed", horizontal=True)
            else:
                st.radio(" ", option, index=None, key="radioChosen", on_change=updateRadioAnswer, args=(row[0],), label_visibility="collapsed", horizontal=True)
            if row[6] != "" and row[6] is not None and st.session_state.radioChosen is None:
                st.write(f":red[**你已选择:** ] :blue[[**{option[int(row[6]) ^ 1][0]}**]]")
            #st.write(st.session_state.radioChosen)
        if st.session_state.radioCompleted:
            radioArea.empty()
            st.session_state.radioCompleted = False
            sql = f"SELECT userAnswer from {st.session_state.examFinalTable} where ID = {row[0]}"
            tempUserAnswer = execute_sql(cur, sql)[0][0]
            if tempUserAnswer != "":
                st.radio(" ", option, index=int(tempUserAnswer) ^ 1, key="radioChosen2", on_change=updateRadioAnswer2, args=(row[0],), label_visibility="collapsed", horizontal=True)
            radioArea.empty()
    elif row[4] == '填空题':
        orgOption = row[6].replace("；", ";").split(";")
        textAnswerArea = st.empty()
        with textAnswerArea.container():
            for i in range(row[1].count("()")):
                if row[6] == "":
                    st.text_input(label=" ", key=f"textAnswer_{i}", placeholder=f"请输入第{i + 1}个括号内的内容", label_visibility="collapsed")
                else:
                    st.text_input(label=" ", value=orgOption[i], key=f"textAnswer_{i}", placeholder=f"请输入第{i + 1}个括号内的内容", label_visibility="collapsed")
            buttonTA = st.button("确定")
            if buttonTA:
                updateTA()
                textAnswerArea.empty()
                st.toast(f"第{row[0]}题答案已更新, 请点击上方按钮继续答题或交卷")
    if st.session_state.examType == "training":
        col1, col2, col3 = st.columns(3)
        with col3:
            AIOptionIndex = sac.segmented(
                items=[
                    sac.SegmentedItem(label="讯飞"),
                    sac.SegmentedItem(label="百度"),
                    sac.SegmentedItem(label="深索"),
                ], label="可选LLM大模型", index=AIOptionIndex, align="start", color="red", return_index=True
            )
        AIModelName = AIOption[AIOptionIndex]
        updateAIModel2(AIOption, AIOptionIndex)
        if row[5] != "":
            with col1:
                buttonAnalysis = st.button("显示答案解析")
            with col2:
                buttonDelAnalysis = st.button("删除本题答案解析")
            if buttonAnalysis:
                st.subheader(f":orange[解析 标准答案: :green[[{standardAnswer}]]]\n{row[5]}", divider="gray")
            if buttonDelAnalysis:
                delAnalysis(row)
        else:
            if AIModelName != "":
                with col1:
                    buttonAnalysis = st.button(f"A.I.答案解析模型 :green[{AIModelName.replace('大模型', '')}]")
                with col2:
                    buttonDelAnalysis = st.button("删除本题答案解析")
                if AIModelName == "文心千帆大模型":
                    AIModelType = st.radio(label="请设置生成内容类型", options=("简洁", "详细"), index=0, horizontal=True, help="返回结果类型, 详细型附加了很多解释内容")
                    if AIModelType == "简洁":
                        AIModel = "ERNIE Speed-AppBuilder"
                    elif AIModelType == "详细":
                        AIModel = "ERNIE-Speed-8K"
                elif AIModelName == "DeepSeek大模型":
                    AIModelType = st.radio(label="请设置模型版本", options=("R1", "V3"), index=0, horizontal=True, help="R1推理好速度慢, V3推理快速结果尚可")
                    if AIModelType == "R1":
                        AIModel = "deepseek-reasoner"
                    elif AIModelType == "V3":
                        AIModel = "deepseek-chat"
                else:
                    AIModelType = ''
                if buttonAnalysis:
                    AIAnswerInfo = st.empty()
                    with AIAnswerInfo.container():
                        if AIModelType != '':
                            st.info(f"正在使用 :red[{AIModelName.replace('大模型', '')} {AIModelType}] 获取答案解析, 内容不能保证正确, 仅供参考! 请稍等...")
                        else:
                            st.info(f"正在使用 :red[{AIModelName.replace('大模型', '')}] 获取答案解析, 内容不能保证正确, 仅供参考! 请稍等...")
                    if AIModelName == "文心千帆大模型":
                        AIAnswer = qianfan_AI(row[1], AIModel, option, row[4])
                    elif AIModelName == "讯飞星火大模型":
                        AIAnswer = xunfei_xh_AI(row[1], option, row[4])
                    elif AIModelName == "DeepSeek大模型":
                        AIAnswer = deepseek_AI(row[1], option, row[4], AIModel)
                    AIAnswerInfo.empty()
                    if AIAnswer != "" and AIAnswer.find("无法直接回答") == -1 and AIAnswer.find("尚未查询") == -1 and AIAnswer.find("我不确定您想要表达什么意思") == -1 and AIAnswer.find("由于信息不足，无法给出准确答案") == -1 and AIAnswer.find("无法确定正确答案") == -1 and AIAnswer.find("无法提供准确答案") == -1:
                        if AIAnswer.startswith(":"):
                            AIAnswer = AIAnswer[1:]
                        AIAnswer = AIAnswer + f"\n\n:red[答案解析来自[{AIModelName}], 非人工解析内容, 仅供参考!]"
                        st.subheader(f":orange[解析 标准答案: :green[[{standardAnswer}]]]\n{AIAnswer}", divider="gray")
                        if flagAIUpdate:
                            AIAnswer = AIAnswer.replace('"', '""').replace("'", "''")
                            for each in ["questions", "commquestions", "morepractise", "favques", st.session_state.examTable, st.session_state.examFinalTable]:
                                sql = f"UPDATE {each} set qAnalysis = '{AIAnswer}' where Question = '{row[1]}' and qType = '{row[4]}'"
                                execute_sql_and_commit(conn, cur, sql)
                            st.toast("A.I.答案解析内容已更新至题库")
                    else:
                        st.info("A.I.获取答案解析失败")
                if buttonDelAnalysis:
                    delAnalysis(row)
            else:
                st.info("没有设置A.I.大模型")
    st.session_state.curQues = row[0]


@st.fragment
def delAnalysis(row):
    for each in ["questions", "commquestions", "morepractise", "favques", st.session_state.examTable, st.session_state.examFinalTable]:
        sql = f"UPDATE {each} set qAnalysis = '' where Question = '{row[1]}' and qType = '{row[4]}'"
        execute_sql_and_commit(conn, cur, sql)
    st.info("本题解析已删除")


@st.fragment
def manualFIB(rowID):
    fibAI = ""
    sql = f"SELECT Question, qAnswer, userAnswer from {st.session_state.examFinalTable} where ID = {rowID}"
    fibRow = execute_sql(cur, sql)[0]
    fibQues = fibRow[0]
    userAP = fibRow[2].split(";")
    qAP = fibRow[1].split(";")
    if len(qAP) == len(userAP):
        for each in qAP:
            b1 = fibQues.find("()")
            if b1 != -1:
                fibQues = f"{fibQues[:b1]}({each}){fibQues[b1 + 2:]}"
        fibAI = xunfei_xh_AI_fib(userAP, fibQues)

    return fibAI


@st.fragment
def getStandardAnswer(qRow):
    radioOption, standardAnswer = ["A", "B", "C", "D", "E", "F", "G", "H"], ""
    if qRow[4] == "单选题" or qRow[4] == "多选题":
        orgOption = qRow[3].replace("；", ";").split(";")
        for value in orgOption:
            standardAnswer = standardAnswer + radioOption[int(value)] + ", "
    elif qRow[4] == "判断题":
        if qRow[3] == "1":
            standardAnswer = "正确"
        else:
            standardAnswer = "错误"
    elif qRow[4] == "填空题":
        standardAnswer = qRow[3].replace("；", ";").replace(";", ", ")
    if standardAnswer.endswith(", "):
        standardAnswer = standardAnswer[:-2]

    return standardAnswer


# noinspection PyTypeChecker
@st.fragment
def updateTA():
    textAnswerPack = []
    for key in st.session_state.keys():
        if "textAnswer_" in key:
            textAnswerPack.append([int(key.replace("textAnswer_", "")), st.session_state[key]])
    textAnswerPack.sort()
    st.session_state.answer = ""
    for each in textAnswerPack:
        st.session_state.answer += each[1] + ";"
    if st.session_state.answer.endswith(";"):
        st.session_state.answer = st.session_state.answer[:-1]
    updateAnswer(st.session_state.curQues)


def changeCurQues(step, cQuesCount):
    st.session_state.curQues += step
    if st.session_state.curQues < 1:
        st.session_state.curQues = 1
    elif st.session_state.curQues > cQuesCount:
        st.session_state.curQues = cQuesCount


@st.fragment
def quesGoto():
    if st.session_state.chosenID is not None:
        st.session_state.goto = True
        cop = re.compile('[^0-9^.]')
        st.session_state.curQues = int(cop.sub('', st.session_state.chosenID))


@st.fragment
def displayTimeCountdown():
    countdownType = "Circle"
    if countdownType == "Flip":
        remindTimeText = open("./MyComponentsScript/Countdown-NoFlip.txt", "r", encoding="utf-8").read()
    elif countdownType == "Circle":
        remindTimeText = open("./MyComponentsScript/Countdown-Circle.txt", "r", encoding="utf-8").read()
    else:
        remindTimeText = ""
    timeArea = st.empty()
    with timeArea.container():
        #st.write(f"### :red[{st.session_state.examName}]")
        #st.markdown(f"<font face='微软雅黑' color=red size=16><center>**{st.session_state.examName}**</center></font>", unsafe_allow_html=True)
        #st.markdown(f"### <font face='微软雅黑' color=red><center>{st.session_state.examName}</center></font>", unsafe_allow_html=True)
        flagTime = bool(getParam("显示考试时间", st.session_state.StationCN))
        if flagTime:
            info1, info2, info3, info4 = st.columns(4)
        else:
            info1, info2, info3 = st.columns(3)
        if st.session_state.examType == "exam" or flagTime:
            examTimeLimit = int(getParam("考试时间", st.session_state.StationCN) * 60)
            examEndTime = st.session_state.examStartTime + examTimeLimit
            if countdownType == "Flip":
                examEndTimeText = time.strftime("%Y-%m-%dT%H:%M:%S", time.localtime(examEndTime))
            elif countdownType == "Circle":
                examEndTimeText = str(examTimeLimit - int(time.time() - st.session_state.examStartTime))
            else:
                examEndTimeText = ""
            remindTimeText = remindTimeText.replace("remindTime", f'"{examEndTimeText}"')
            remainingTime = examTimeLimit - (int(time.time() - st.session_state.examStartTime))
            if remainingTime < 0:
                if st.session_state.examType == "exam":
                    st.warning("⚠️ 考试已结束, 将强制交卷!")
                    st.session_state.calcScore = True
                    calcScore()
                else:
                    st.session_state.examStartTime = int(time.time())
            elif remainingTime < 900:
                st.warning(f"⚠️ :red[考试剩余时间已不足{int(remainingTime / 60) + 1}分钟, 请抓紧时间完成考试!]")
            with info1:
                if remindTimeText != "":
                    components.html(remindTimeText, height=94)
        sql = f"SELECT count(ID) from {st.session_state.examFinalTable} where userAnswer <> ''"
        acAnswer1 = execute_sql(cur, sql)[0][0]
        sql = f"SELECT count(ID) from {st.session_state.examFinalTable} where userAnswer = ''"
        acAnswer2 = execute_sql(cur, sql)[0][0]
        if flagTime:
            info2.metric(label="已答题", value=acAnswer1)
            info3.metric(label="未答题", value=acAnswer2)
            info4.metric(label="总题数", value=acAnswer1 + acAnswer2)
        else:
            info1.metric(label="已答题", value=acAnswer1)
            info2.metric(label="未答题", value=acAnswer2)
            info3.metric(label="总题数", value=acAnswer1 + acAnswer2)
        style_metric_cards(border_left_color=st.session_state.tooltipColor)


@st.fragment
def displayBigTime():
    components.html(open("./MyComponentsScript/Clock-Big.txt", "r", encoding="utf-8").read(), height=140)


@st.fragment
def displaySmallTime():
    components.html(open("./MyComponentsScript/Clock-Small.txt", "r", encoding="utf-8").read(), height=34)


@st.fragment
def displaySmallClock():
    components.html(open("./MyComponentsScript/Clock-Number.txt", "r", encoding="utf-8").read(), height=30)


@st.fragment
def displayBigTimeCircle():
    components.html(open("./MyComponentsScript/Clock-Big-Circle.txt", "r", encoding="utf-8").read(), height=260)


@st.fragment
def displayVisitCounter():
    sql = "SELECT pyLM from verinfo where pyFile = 'visitcounter'"
    visitcount = execute_sql(cur, sql)[0][0]
    countScript = (open("./MyComponentsScript/FlipNumber.txt", "r", encoding="utf-8").read()).replace("visitcount", str(visitcount))
    components.html(countScript, height=40)


@st.fragment
def displayAppInfo():
    infoStr = open("./MyComponentsScript/glowintext.txt", "r", encoding="utf-8").read()
    infoStr = infoStr.replace("软件名称", APPNAME)
    verinfo, verLM, likeCM = getVerInfo()
    infoStr = infoStr.replace("软件版本", f"软件版本: {int(verinfo / 10000)}.{int((verinfo % 10000) / 100)}.{int(verinfo / 10)} building {verinfo}")
    infoStr = infoStr.replace("更新时间", f"更新时间: {time.strftime('%Y-%m-%d %H:%M', time.localtime(verLM))}")
    #infoStr = infoStr.replace("用户评价", f"用户评价: {EMOJI[int(likeCM) - 1][0]} {likeCM} I feel {EMOJI[int(likeCM) - 1][1]}")
    infoStr = infoStr.replace("更新内容", f"更新内容: {UPDATETYPE['New']} 鉴于sqlite3数据库文件被多个用户同时访问时，可能会出现错误，现将数据库改为MySQL")
    components.html(infoStr, height=340)


@st.fragment
def changelog():
    changelogInfo = open("./CHANGELOG.md", "r", encoding="utf-8").read()
    st.markdown(changelogInfo)


@st.dialog("交卷")
def submit_dialog(prompt):
    st.write(f":red[**{prompt}**]")
    buttonSubmit = st.button("确定")
    buttonCancel = st.button("取消")
    if buttonSubmit:
        st.session_state.calcScore = True
        st.rerun()
    elif buttonCancel:
        st.session_state.calcScore = False
        st.rerun()


def ClearStr(strValue):
    strValue = strValue.replace("\n", "").replace("\t", "").strip()

    return strValue


@st.fragment
def addExamIDD():
    flagSuccess, examDateStr = False, ""
    # 创建一个空的itemArea
    itemArea = st.empty()
    with itemArea.container():
        # 获取考试名称输入
        examName = st.text_input("考试名称", value="", help="名称不能设置为练习题库(此为保留题库)")
        # 清理考试名称字符串
        examName = ClearStr(examName)
        # 获取考试有效期输入
        examDate = st.date_input("请设置考试有效期", min_value=datetime.date.today() + datetime.timedelta(days=1), max_value=datetime.date.today() + datetime.timedelta(days=180), value=datetime.date.today() + datetime.timedelta(days=3), help="考试有效期最短1天, 最长180天, 默认3天")
        # 检查考试名称和有效期是否有效
        if examName and examDate and examName != "练习题库":
            # 创建添加考试场次的按钮
            buttonSubmit = st.button("添加考试场次")
            if buttonSubmit:
                # 将考试日期转换为字符串
                examDateStr = examDate
                # 将考试日期转换为时间戳
                examDate = int(time.mktime(time.strptime(f"{examDate} 23:59:59", "%Y-%m-%d %H:%M:%S")))
                # 查询数据库中是否已经存在该考试名称和站点
                sql = f"SELECT ID from examidd where examName = '{examName}' and StationCN = '{st.session_state.StationCN}'"
                # 如果不存在，则插入新的考试场次
                if not execute_sql(cur, sql):
                    sql = f"INSERT INTO examidd(examName, validDate, StationCN) VALUES('{examName}', {examDate}, '{st.session_state.StationCN}')"
                    execute_sql_and_commit(conn, cur, sql)
                    # 设置操作成功标志
                    flagSuccess = True
                    # 清空itemArea
                    itemArea.empty()
                else:
                    # 如果考试场次已存在，则显示错误信息
                    st.error(f"[{examName}] 考试场次已存在")
        else:
            # 如果考试名称为空，则显示警告信息
            if not examName:
                st.warning("请输入考试名称")
    # 如果操作成功，则进行后续处理
    if flagSuccess:
        # 查询数据库中是否成功添加了新的考试场次
        sql = f"SELECT ID from examidd where examName = '{examName}' and StationCN = '{st.session_state.StationCN}'"
        if execute_sql(cur, sql):
            # 显示成功信息
            st.success(f"考试场次: [{examName}] 有效期: [{examDateStr} 23:59:59] 添加成功")
            # 更新关键操作日志
            updateKeyAction(f"新建考试场次{examName}")
            # 清空itemArea
            itemArea.empty()
        else:
            # 如果添加失败，则显示错误信息
            st.error(f"考试场次 [{examName}] 添加失败")


@st.fragment
def addStation():
    flagSuccess = False
    # 创建一个空的区域用于放置界面元素
    itemArea = st.empty()
    with itemArea.container():
        # 输入站室名称
        sn = st.text_input("站室名称", value="")
        # 清除输入字符串中的多余空格
        sn = ClearStr(sn)
        if sn:
            # 添加按钮
            buttonSubmit = st.button("添加站室名称")
            if buttonSubmit:
                # 查询站室名称是否已存在
                sql = "SELECT ID from stations where Station = '" + sn + "'"
                if not execute_sql(cur, sql):
                    sql = f"INSERT INTO stations(Station) VALUES('{sn}')"
                    execute_sql_and_commit(conn, cur, sql)
                    flagSuccess = True
                    # 清空区域
                    itemArea.empty()
                else:
                    # 如果站室名称已存在，显示错误信息
                    st.error(f"[{sn}] 已存在")
        else:
            # 如果站室名称为空，显示警告信息
            if not sn:
                st.warning("请输入站室名称")

    if flagSuccess:
        # 查询新添加的站室记录
        sql = "SELECT ID from stations where Station = '" + sn + "'"
        if execute_sql(cur, sql):
            # 查询是否存在以站室名称命名的临时表
            sql = f"SELECT * from sqlite_master where type = 'table' and name = 'setup_{sn}'"
            tempTable = execute_sql(cur, sql)
            if not tempTable:
                # 如果不存在，则创建临时表
                sql = """CREATE TABLE exampleTable (
                            ID integer not null primary key autoincrement,
                            paramName text not null,
                            param integer,
                            paramType text not null
                        );"""
                sql = sql.replace("exampleTable", f"setup_{sn}")
                cur.execute(sql)
                conn.commit()
                # 将默认表的数据插入到新创建的临时表中
                sql = f"INSERT INTO setup_{sn}(paramName, param, paramType) SELECT paramName, param, paramType from setup_默认"
                execute_sql_and_commit(conn, cur, sql)

            # 遍历每个章节名称，并检查是否存在对应的章节记录
            for each in ["公共题库", "错题集", "关注题集"]:
                sql = f"SELECT ID from questionaff where chapterName = '{each}' and StationCN = '{sn}'"
                if not execute_sql(cur, sql):
                    # 如果不存在，则插入新的章节记录
                    sql = f"INSERT INTO questionaff(chapterName, StationCN, chapterRatio, examChapterRatio) VALUES('{each}', '{sn}', 10, 10)"
                    execute_sql_and_commit(conn, cur, sql)

            # 显示添加成功的消息
            st.success(f"[{sn}] 站室添加成功")
            # 更新关键操作记录
            updateKeyAction(f"新建站室{sn}")
            # 清空区域
            itemArea.empty()
        else:
            # 如果添加站室失败，显示错误信息
            st.error(f"[{sn}] 添加站室失败")


@st.fragment
def addUser():
    flagSuccess = False
    # 获取所有站室的中文名称
    stationCName = getStationCNALL()
    # 创建一个空的容器用于放置元素
    itemArea = st.empty()
    with itemArea.container():
        # 创建两列布局
        col1, col2 = st.columns(2)
        # 在第一列中创建一个数字输入框，用于输入用户编码
        userName = col1.number_input("用户编码", min_value=1, max_value=999999, value=1, help="建议使用员工编码, 姓名和站室可以有重复, 但是编码必须具有唯一性")
        # 在第二列中创建一个文本输入框，用于输入用户姓名
        userCName = col2.text_input("用户姓名", max_chars=10, autocomplete="name", help="请输入用户中文姓名")
        # 创建一个滑动选择器，用于选择站室
        station = st.select_slider("站室", stationCName, value=st.session_state.StationCN)
        # 创建一个密码输入框，用于输入密码
        userPassword1 = st.text_input("设置密码", max_chars=8, type="password", autocomplete="off", help="设置用户密码")
        # 再次创建一个密码输入框，用于确认密码
        userPassword2 = st.text_input("请再次输入密码", max_chars=8, type="password", placeholder="请与上一步输入的密码一致", autocomplete="off")
        # 创建一个开关，用于选择用户类型（管理员/普通用户）
        userType = sac.switch(label="管理员", on_label="On", align='start', size='md', value=False)
        # 清除用户姓名中的多余空格
        userCName = ClearStr(userCName)
        # 检查所有必填项是否都已填写
        if userName and userCName and userPassword1 and userPassword2 and userPassword1 != "" and userPassword2 != "":
            # 创建一个提交按钮
            buttonSubmit = st.button("添加用户")
            if buttonSubmit:
                # 检查两次输入的密码是否一致
                if userPassword1 == userPassword2:
                    # 将用户编码转换为整数
                    un = int(userName)
                    # 根据用户类型设置用户类型字符串
                    if userType:
                        ut = "admin"
                    else:
                        ut = "user"
                    # 输出选择的站室
                    st.write(station)
                    # 检查用户是否已存在
                    sql = "SELECT ID from users where userName = " + str(un)
                    if not execute_sql(cur, sql):
                        # 加密密码
                        userPassword1 = getUserEDKeys(userPassword1, "enc")
                        # 插入新用户数据
                        sql = f"INSERT INTO users(userName, userCName, userType, StationCN, userPassword) VALUES({un}, '{userCName}', '{ut}', '{station}', '{userPassword1}')"
                        execute_sql_and_commit(conn, cur, sql)
                        # 设置操作成功标志
                        flagSuccess = True
                        # 清空容器
                        itemArea.empty()
                    else:
                        # 输出错误信息
                        st.error(f"ID: [{userName}] 姓名: [{userCName}] 用户已存在或用户编码重复")
                else:
                    # 输出密码不一致的错误信息
                    st.error("两次输入密码不一致")
        else:
            # 检查哪些必填项未填写，并输出相应的警告信息
            if not userCName:
                st.warning("请输入用户姓名")
            elif not userPassword1:
                st.warning("请输入密码")
            elif not userPassword2:
                st.warning("请确认密码")

    # 如果操作成功，则更新用户信息
    if flagSuccess:
        # 查询新用户数据
        sql = "SELECT ID from users where userName = " + str(un) + " and StationCN = '" + station + "' and userCName = '" + userCName + "'"
        if execute_sql(cur, sql):
            # 输出操作成功的提示信息
            st.success(f"ID: [{userName}] 姓名: [{userCName}] 类型: [{ut}] 站室: [{station}] 用户添加成功")
            # 记录操作日志
            updateKeyAction(f"新建用户: {userName} 姓名: {userCName} 类型: {ut} 站室: {station}")
            # 清空容器
            itemArea.empty()
        else:
            # 输出操作失败的错误信息
            st.error(f"ID: [{userName}] 姓名: [{userCName}] 类型: [{ut}] 站室: [{station}] 用户添加失败")


def getStationCNALL(flagALL=False):
    StationCNamePack = []
    if flagALL:
        StationCNamePack.append("全站")
    sql = "SELECT Station from stations order by ID"
    rows = execute_sql(cur, sql)
    for row in rows:
        StationCNamePack.append(row[0])

    return StationCNamePack


def updateDAParam(updateParamType):
    for key in st.session_state.keys():
        if key.startswith("dasetup_"):
            upID = key[key.find("_") + 1:]
            sql = f"UPDATE setup_{st.session_state.StationCN} SET param = {int(st.session_state[key])} WHERE ID = {upID}"
            execute_sql_and_commit(conn, cur, sql)
    st.success(f"{updateParamType} 参数更新成功")
    updateKeyAction("考试参数更新")


def updateSwitchOption(quesType):
    if st.session_state[quesType]:
        sql = f"UPDATE setup_{st.session_state.StationCN} set param = 1 where paramName = '{quesType}'"
    else:
        sql = f"UPDATE setup_{st.session_state.StationCN} set param = 0 where paramName = '{quesType}'"
    execute_sql_and_commit(conn, cur, sql)
    if quesType == "测试模式":
        st.session_state.debug = bool(st.session_state[quesType])
    if quesType == "时钟样式":
        st.session_state.clockType = bool(st.session_state[quesType])
    #st.success(f"{quesType} 设置更新成功")


def setupReset():
    execute_sql_and_commit(conn, cur, sql=f"DELETE from setup_{st.session_state.StationCN} where ID > 0")
    sql = f"INSERT INTO setup_{st.session_state.StationCN}(paramName, param, paramType) SELECT paramName, param, paramType from setup_默认"
    execute_sql_and_commit(conn, cur, sql)
    sql = f"UPDATE questionaff set chapterRatio = 10, examChapterRatio = 10 where StationCN = '{st.session_state.StationCN}' and (chapterName = '公共题库' or chapterName = '错题集')"
    execute_sql_and_commit(conn, cur, sql)
    sql = f"UPDATE questionaff set chapterRatio = 5, examChapterRatio = 5 where StationCN = '{st.session_state.StationCN}' and chapterName <> '公共题库' and chapterName <> '错题集'"
    execute_sql_and_commit(conn, cur, sql)
    st.success("所有设置已重置")
    updateKeyAction("重置所有设置")


def updateAIModel():
    sql = f"UPDATE setup_{st.session_state.StationCN} set param = 0 where paramType = 'others' and paramName like '%大模型'"
    execute_sql_and_commit(conn, cur, sql)
    sql = f"UPDATE setup_{st.session_state.StationCN} set param = 1 where paramType = 'others' and paramName = '{st.session_state.AIModel}'"
    execute_sql_and_commit(conn, cur, sql)
    st.success(f"LLM大模型已设置为{st.session_state.AIModel}")


@st.fragment
def updateAIModel2(AIOption, AIOptionIndex):
    sql = f"UPDATE setup_{st.session_state.StationCN} set param = 0 where paramType = 'others' and paramName like '%大模型'"
    execute_sql_and_commit(conn, cur, sql)
    sql = f"UPDATE setup_{st.session_state.StationCN} set param = 1 where paramType = 'others' and paramName = '{AIOption[AIOptionIndex]}'"
    execute_sql_and_commit(conn, cur, sql)


# noinspection PyTypeChecker
def highlight_max(x, forecolor='black', backcolor='yellow'):
    is_max = x == x.max()

    return [f'color: {forecolor}; background-color: {backcolor}' if v else '' for v in is_max]


def queryExamAnswer(tablename):
    if tablename == "morepractise":
        chosenType = ["错题"]
    else:
        chosenType = ["对题", "错题"]
    options = st.multiselect(
        "查询类型",
        chosenType,
        ["错题"],
    )
    if options:
        searchButton = st.button("查询")
        if searchButton:
            if len(options) == 2:
                sql = "SELECT Question, qOption, qAnswer, qType, qAnalysis, userAnswer, ID from " + tablename + " where userAnswer <> '' and userName = " + str(st.session_state.userName) + " order by ID"
            elif len(options) == 1:
                if options[0] == "对题":
                    sql = "SELECT Question, qOption, qAnswer, qType, qAnalysis, userAnswer, ID from " + tablename + " where userAnswer <> '' and qAnswer = userAnswer and userName = " + str(st.session_state.userName) + " order by ID"
                elif options[0] == "错题":
                    sql = "SELECT Question, qOption, qAnswer, qType, qAnalysis, userAnswer, ID from " + tablename + " where userAnswer <> '' and qAnswer <> userAnswer and userName = " + str(st.session_state.userName) + " order by ID"
                else:
                    sql = ""
            else:
                sql = ""
            rows = execute_sql(cur, sql)
            if rows:
                for row in rows:
                    if row[2] != row[5]:
                        flagAnswer = "错误"
                        st.subheader("", divider="red")
                    else:
                        flagAnswer = "正确"
                        st.subheader("", divider="green")
                    st.subheader(f"题目: :grey[{row[0]}]")
                    if row[3] == "单选题":
                        st.write(":red[标准答案:]")
                        option, userAnswer = [], ["A", "B", "C", "D"]
                        tmp = row[1].replace("；", ";").split(";")
                        for index, each in enumerate(tmp):
                            each = each.replace("\n", "").replace("\t", "").strip()
                            option.append(f"{userAnswer[index]}. {each}")
                        st.radio(" ", option, key=f"compare_{row[6]}", index=int(row[2]), horizontal=True, label_visibility="collapsed", disabled=True)
                        st.write(f"你的答案: :red[{userAnswer[int(row[5])]}] 你的选择为: :blue[[{flagAnswer}]]")
                    elif row[3] == "多选题":
                        userOption = ["A", "B", "C", "D", "E", "F", "G", "H"]
                        st.write(":red[标准答案:]")
                        option = row[1].replace("；", ";").split(";")
                        orgOption = row[2].replace("；", ";").split(";")
                        for index, value in enumerate(option):
                            value = value.replace("\n", "").replace("\t", "").strip()
                            if str(index) in orgOption:
                                st.checkbox(f"{userOption[index]}. {value}:", value=True, disabled=True)
                            else:
                                st.checkbox(f"{userOption[index]}. {value}:", value=False, disabled=True)
                        userAnswer = row[5].replace("；", ";").split(";")
                        tmp = ""
                        for each in userAnswer:
                            tmp = tmp + userOption[int(each)] + ", "
                        st.write(f"你的答案: :red[{tmp[:-2]}] 你的选择为: :blue[[{flagAnswer}]]")
                    elif row[3] == "判断题":
                        st.write(":red[标准答案:]")
                        option = ["A. 正确", "B. 错误"]
                        tmp = int(row[2]) ^ 1
                        st.radio(" ", option, key=f"compare_{row[6]}", index=tmp, horizontal=True, label_visibility="collapsed", disabled=True)
                        tmp = int(row[5]) ^ 1
                        st.write(f"你的答案: :red[{option[tmp]}] 你的选择为: :blue[[{flagAnswer}]]")
                    elif row[3] == "填空题":
                        option = row[2].replace("；", ";").split(";")
                        userAnswer = row[5].replace("；", ";").split(";")
                        st.write(":red[标准答案:]")
                        for index, value in enumerate(option):
                            st.write(f"第{index + 1}个填空: :green[{value}]")
                        st.write("你的答案:")
                        for index, value in enumerate(userAnswer):
                            st.write(f"第{index + 1}个填空: :red[{value}]")
                        st.write(f"你的填写为: :blue[[{flagAnswer}]]")
                    if row[4] != "":
                        st.markdown(f"答案解析: :green[{row[4]}]")
            else:
                st.info("暂无数据")
    else:
        st.warning("请设置查询类型")


# noinspection PyTypeChecker
def queryExamResult():
    # 初始化查询选项列表
    searchOption = []

    # 构建SQL查询语句
    sql = f"SELECT ID, examName from examidd where StationCN = '{st.session_state.StationCN}' order by ID"
    # 执行SQL查询
    rows = execute_sql(cur, sql)

    # 遍历查询结果，将考试名称添加到查询选项列表中
    for row in rows:
        searchOption.append(row[1])

    # 创建下拉选择框，供用户选择考试场次
    searchExamName = st.selectbox("请选择考试场次", searchOption, index=None)

    # 创建多选框，供用户选择查询类型
    options = st.multiselect(
        "查询类型",
        ["通过", "未通过"],
        ["未通过"],
    )

    # 根据用户是否选择了考试场次，设置查询按钮的可用状态
    if searchExamName:
        searchButton = st.button("查询")
    else:
        searchButton = st.button("查询", disabled=True)

    # 当用户点击查询按钮且选择了考试场次时，执行查询逻辑
    if searchButton and searchExamName:
        if options:
            # 创建标签页
            tab1, tab2 = st.tabs(["简报", "详情"])

            # 构建SQL查询语句
            sql = f"SELECT userName, userCName, examScore, examDate, examPass from examresult where examName = '{searchExamName}' and ("
            for each in options:
                if each == "通过":
                    sql = sql + " examPass = 1 or "
                elif each == "未通过":
                    sql = sql + " examPass = 0 or "
            if sql.endswith(" or "):
                sql = sql[:-4] + ") order by ID DESC"

            # 执行SQL查询
            rows = execute_sql(cur, sql)

            # 处理查询结果
            if rows:
                df = pd.DataFrame(rows, dtype=str)
                df.columns = ["编号", "姓名", "成绩", "考试日期", "考试结果"]
                for index, value in enumerate(rows):
                    df.loc[index, "考试日期"] = time.strftime("%Y-%m-%d %H:%M:%S", time.localtime(int(df["考试日期"][index])))
                    df.loc[index, "考试结果"] = "通过" if int(df["考试结果"][index]) == 1 else "未通过"

                # 在详情标签页中显示查询结果
                tab2.dataframe(df.style.apply(highlight_max, backcolor='yellow', subset=["成绩", "考试结果"]))

            # 在简报标签页中显示查询结果的摘要信息
            if rows:
                for row in rows:
                    tab1.markdown(f"考生ID:  :red[{row[0]}] 考生姓名: :red[{row[1]}] 考试时间: :red[{time.strftime('%Y-%m-%d %H:%M:%S', time.localtime(row[3]))}]")
                    tab1.subheader(f"考试成绩: {row[2]} 分")
                    if row[4] == 1:
                        tab1.subheader("考试结果: :blue[通过] 👏")
                        tab1.subheader("", divider="orange")
                    else:
                        tab1.subheader("考试结果: :red[未通过] 😝")
                        tab1.subheader("", divider="red")
            else:
                st.info("暂无数据")
        else:
            st.warning("请设置查询类型")


def queryExamResultUsers():
    # 初始化一个空列表，用于存储考试名称
    ExamNamePack = []

    # 构建SQL查询语句，获取指定站室的考试名称
    sql = f"SELECT ID, examName from examidd where StationCN = '{st.session_state.StationCN}' order by ID"
    # 执行SQL查询，获取考试名称
    rows = execute_sql(cur, sql)

    # 遍历查询结果，将考试名称添加到ExamNamePack列表中
    for row in rows:
        ExamNamePack.append(row[1])

    # 使用streamlit的selectbox组件，让用户选择考试场次
    searchExamName = st.selectbox("请选择考试场次", ExamNamePack, index=None)

    # 使用streamlit的multiselect组件，让用户选择查询类型
    options = st.multiselect(
        "查询类型",
        ["已参加考试", "未参加考试"],
        ["未参加考试"],
    )

    # 使用streamlit的button组件，让用户点击查询按钮
    searchButton = st.button("查询")

    # 如果用户点击了查询按钮并选择了考试场次
    if searchButton and searchExamName:
        if options:
            # 使用streamlit的tabs组件，创建两个标签页
            tab1, tab2 = st.tabs(["简报", "详情"])

            # 根据用户选择的查询类型构建不同的SQL查询语句
            if len(options) == 2:
                sql = "SELECT userName, userCName, StationCN from users where StationCN = '" + st.session_state.StationCN + "' and userType <> 'supervisor' order by ID"
            elif len(options) == 1:
                if options[0] == "已参加考试":
                    sql = "SELECT users.userName, users.userCName, users.StationCN from users, examresult where users.userType <> 'supervisor' and examresult.examName = '" + searchExamName + "' and examresult.userName = users.userName and users.StationCN = '" + st.session_state.StationCN + "'"
                elif options[0] == "未参加考试":
                    sql = "SELECT userName, userCName, StationCN from users where userType <> 'supervisor' and userName not in (SELECT users.userName from users, examresult where examresult.examName = '" + searchExamName + "' and examresult.userName = users.userName) and StationCN = '" + st.session_state.StationCN + "'"

            # 执行SQL查询，获取用户信息
            rows = execute_sql(cur, sql)

            # 如果查询结果不为空
            if rows:
                # 将查询结果转换为DataFrame，并设置列名
                df = pd.DataFrame(rows)
                df.columns = ["编号", "姓名", "站室"]
                # 在第二个标签页中显示查询结果
                tab2.dataframe(df)

            # 再次遍历查询结果，获取每个考生的详细信息
            if rows:
                for row in rows:
                    # 构建SQL查询语句，获取每个考生的考试成绩
                    sql = "SELECT userName, userCName, examScore, examDate, examPass from examresult where examName = '" + searchExamName + "' and userName = " + str(row[0])
                    rows2 = execute_sql(cur, sql)

                    # 如果查询到考试成绩
                    if rows2:
                        # 在第一个标签页中显示考生的考试成绩和考试结果
                        tab1.markdown(f"考生ID:  :red[{rows2[0][0]}] 考生姓名: :red[{rows2[0][1]}] 考试时间: :red[{time.strftime('%Y-%m-%d %H:%M:%S', time.localtime(rows2[0][3]))}]")
                        tab1.subheader(f"考试成绩: {rows2[0][2]} 分")
                        if rows2[0][4] == 1:
                            tab1.subheader("考试结果: :blue[通过] 👏")
                            tab1.subheader("", divider="orange")
                        else:
                            tab1.subheader("考试结果: :red[未通过] 🤪")
                            tab1.subheader("", divider="red")
                    else:
                        # 如果未查询到考试成绩，显示未参加考试的信息
                        tab1.subheader("未参加考试", divider="red")
                        tab1.markdown(f"考生ID:  :red[{row[0]}] 考生姓名: :red[{row[1]}] 站室: :red[{row[2]}]")
            else:
                # 如果没有查询到任何数据，显示提示信息
                st.info("暂无数据")
        else:
            # 如果用户未选择任何查询类型，显示警告信息
            st.warning("请设置查询类型")


def verifyUserPW(vUserName, vUserPW):
    st.session_state.userPwRecheck = False
    vUserEncPW = ""
    sql = f"SELECT userPassword from users where userName = {vUserName}"
    pwTable = execute_sql(cur, sql)
    if pwTable:
        vUserEncPW = pwTable[0][0]
        vUserDecPW = getUserEDKeys(vUserEncPW, "dec")
        if vUserPW == vUserDecPW:
            st.session_state.userPwRecheck = True

    return st.session_state.userPwRecheck, vUserEncPW


def resetPassword():
    # 显示副标题和分隔线
    st.subheader(":orange[密码重置及更改账户类型]", divider="red")

    # 检查是否需要重置用户信息
    if st.session_state.userPwRecheck:
        # 显示重置用户信息提示
        st.write(":red[**重置用户信息**]")

        # 创建三列布局
        rCol1, rCol2, rCol3 = st.columns(3)

        # 获取用户编码
        rUserName = rCol1.number_input("用户编码", value=0)

        # 检查用户编码是否不为0
        if rUserName != 0:
            # 执行SQL查询用户信息
            sql = f"SELECT userCName, userType from users where userName = {rUserName}"
            rows = execute_sql(cur, sql)

            # 检查是否查询到用户信息
            if rows:
                # 显示用户姓名
                rCol2.write(f"用户姓名: **{rows[0][0]}**")

                # 在第三列创建布局
                with rCol3:
                    rUserType = False

                    # 根据用户类型设置开关
                    if rows[0][1] == "admin" or rows[0][1] == "supervisor":
                        rUserType = sac.switch(label="管理员", value=True, on_label="On", align='start', size='md')
                    elif rows[0][1] == "user":
                        rUserType = sac.switch(label="管理员", value=False, on_label="On", align='start', size='md')

                # 显示重置类型提示
                st.write("重置类型")

                # 创建重置类型的复选框
                rOption1 = st.checkbox("密码", value=False)
                rOption2 = st.checkbox("账户类型", value=False)

                # 创建重置按钮
                btnResetUserPW = st.button("重置", type="primary")

                # 检查是否点击了重置按钮并选择了重置类型
                if btnResetUserPW and (rOption1 or rOption2):
                    st.button("确认", type="secondary", on_click=actionResetUserPW, args=(rUserName, rOption1, rOption2, rUserType,))
                    st.session_state.userPwRecheck = False
                # 如果未选择任何重置类型，显示警告
                elif not rOption1 and not rOption2:
                    st.warning("请选择重置类型")
            # 如果未查询到用户信息，显示错误
            else:
                st.error("用户不存在")
    # 如果不需要重置用户信息，显示密码输入框
    else:
        vUserPW = st.text_input("请输入密码", max_chars=8, placeholder="请输入管理员密码, 以验证身份", type="password", autocomplete="off")

        # 检查是否输入了密码
        if vUserPW:
            # 验证密码
            if verifyUserPW(st.session_state.userName, vUserPW)[0]:
                st.rerun()
            # 如果密码错误，显示错误提示
            else:
                st.error("密码错误, 请重新输入")


def actionResetUserPW(rUserName, rOption1, rOption2, rUserType):
    rInfo = ""

    # 如果 rOption1 为真
    if rOption1:
        # 获取用户加密密钥
        resetPW = getUserEDKeys("1234", "enc")
        # 构建 SQL 更新语句
        sql = f"UPDATE users SET userPassword = '{resetPW}' where userName = {rUserName}"
        # 执行 SQL 并提交
        execute_sql_and_commit(conn, cur, sql)
        # 更新信息，表示密码已重置
        rInfo += "密码已重置为: 1234 / "
        # 更新操作日志
        updateKeyAction("密码重置")

    # 如果 rOption2 为真
    if rOption2:
        # 如果 rUserType 有值
        if rUserType:
            # 构建 SQL 更新语句，将用户类型更改为管理员
            sql = f"UPDATE users SET userType = 'admin' where userName = {rUserName}"
            # 更新信息，表示账户类型已更改为管理员
            rInfo += "账户类型已更改为: 管理员 / "
            # 更新操作日志
            updateKeyAction("更改账户类型为管理员")
        else:
            # 构建 SQL 更新语句，将用户类型更改为普通用户
            sql = f"UPDATE users SET userType = 'user' where userName = {rUserName}"
            # 更新信息，表示账户类型已更改为用户
            rInfo += "账户类型已更改为: 用户 / "
            # 更新操作日志
            updateKeyAction("更改账户类型为用户")
        # 执行 SQL 并提交
        execute_sql_and_commit(conn, cur, sql)

    # 显示操作结果
    st.success(f"**{rInfo[:-3]}**")


def displayKeyAction():
    # 显示标题和操作日志分隔线
    st.subheader(":violet[操作日志]", divider="red")

    # 检查会话状态中的用户密码是否已重新检查
    if st.session_state.userPwRecheck:
        # 构造SQL查询语句，从keyactionlog表中查询操作日志
        sql = "SELECT userName, userCName, StationCN, userAction, datetime(actionDate, 'unixepoch', 'localtime') from keyactionlog order by actionDate DESC"
        # 执行SQL查询
        rows = execute_sql(cur, sql)
        # 如果查询结果不为空
        if rows:
            # 将查询结果转换为DataFrame，并设置列名
            df = pd.DataFrame(rows, columns=["用户编码", "用户姓名", "所属站室", "操作内容", "操作时间"])
            # 显示DataFrame
            st.write(df)
    else:
        # 提示用户输入密码
        vUserPW = st.text_input("请输入密码", max_chars=8, placeholder="请输入管理员密码, 以验证身份", type="password", autocomplete="off")
        # 如果用户输入密码
        if vUserPW:
            # 验证用户密码
            if verifyUserPW(st.session_state.userName, vUserPW)[0]:
                # 密码验证成功，重新运行函数
                st.rerun()
            else:
                # 密码验证失败，显示错误信息
                st.error("密码错误, 请重新输入")


def ls_get(key):

    return st_javascript(f"localStorage.getItem('{key}');")


def ls_set(key, value):
    value = json.dumps(value, ensure_ascii=False)

    return st_javascript(f"localStorage.setItem('{key}', JSON.stringify('{value}');")


def getAllStations():
    STATIONPACK, stationIndex = [], 0
    sql = "SELECT Station from stations where Station <> '调控中心' order by ID"
    rows = execute_sql(cur, sql)
    for row in rows:
        STATIONPACK.append(row[0])
        if st.session_state.StationCN == row[0]:
            stationIndex = rows.index(row)

    return STATIONPACK, stationIndex


def displayUserManual():
    pdfFile = "./Demo/ETest使用手册.pdf"
    with open(pdfFile, "rb") as f:
        base64_pdf = base64.b64encode(f.read()).decode('utf-8')
    pdf_display = f'<embed src="data:application/pdf;base64,{base64_pdf}" width="800" height="1000" type="application/pdf">'
    st.markdown(pdf_display, unsafe_allow_html=True)


def aiGenerate_Image():
    st.subheader(":green[A.I.文字生图]", divider="rainbow")
    st.markdown("严禁使用敏感词汇, 包括但不限于： \n\t:red[**涉及国家安全的信息；\n\t涉及政治与宗教类的信息；\n\t涉及暴力与恐怖主义的信息；\n\t涉及黄赌毒类的信息；\n\t涉及不文明的信息等**]")
    genImageMode = sac.segmented(
        items=[
            sac.SegmentedItem(label="通义万相"),
            sac.SegmentedItem(label="讯飞星火"),
        ], label="可选LLM大模型", index=0, align="start", color="red"
    )
    txt_generate_image = st.text_input("输入文字，点击按钮即可生成图片", placeholder="一只坐着的橘黄色的猫，表情愉悦，活泼可爱，逼真准确，请勿包含敏感词汇")
    if genImageMode == "通义万相":
        txt_generate_image_neg = st.text_input("用来描述不希望在画面中看到的内容", placeholder="低分辨率、错误、最差质量、低质量、残缺、多余的手指、比例不良等")
    else:
        txt_generate_image_neg = ""
    btn_generate_image = st.button("生成图片")
    if btn_generate_image and txt_generate_image != "":
        result = [False, ""]
        AIGMInfo = st.empty()
        with AIGMInfo.container():
            st.info(f"正在使用 :green[{genImageMode}] 生成图片, 请稍等...")
        if genImageMode == "通义万相":
            result = tywx_generate_image(txt_generate_image.strip(), txt_generate_image_neg.strip())

        elif genImageMode == "讯飞星火":
            result = xfxh_generate_image(txt_generate_image.strip())
        if result[0]:
            st.image(result[1])
        else:
            st.error(f"生成失败: {result[1]}")
        AIGMInfo.empty()


global APPNAME, EMOJI, UPDATETYPE, STATIONPACK

conn = pymysql.connect(
host='localhost',
port=3001,
user='root',
password='7745',
database='etest-mysql',
charset='utf8mb4',
autocommit=True
)
cur = conn.cursor()

st.logo("./Images/etest-logo2.png", icon_image="./Images/exam2.png", size="medium")

# noinspection PyRedeclaration
APPNAME = "调控中心安全生产业务考试系统"
# noinspection PyRedeclaration
EMOJI = [["🥺", "very sad!"], ["😣", "bad!"], ["😋", "not bad!"], ["😊", "happy!"], ["🥳", "fab, thank u so much!"]]
# noinspection PyRedeclaration
UPDATETYPE = {"New": "✨", "Optimize": "🚀", "Fix": "🐞"}
selected = None
if "logged_in" not in st.session_state:
    st.session_state.logged_in = False
    login()

if st.session_state.logged_in:
    with st.sidebar:
        if st.session_state.clockType:
            displaySmallTime()
        else:
            displaySmallClock()
        if st.session_state.examType == "exam":
            selected = sac.menu([
                sac.MenuItem('主页', icon='house'),
                sac.MenuItem('功能', icon='grid-3x3-gap', children=[
                    sac.MenuItem('选择考试', icon='list-task'),
                    sac.MenuItem('开始考试', icon='pencil-square'),
                ]),
                sac.MenuItem('账户', icon='person-gear', children=[
                    sac.MenuItem('密码修改', icon='key', disabled=True),
                    sac.MenuItem('登出', icon='box-arrow-right'),
                ]),
                sac.MenuItem('关于', icon='layout-wtf', children=[
                    #sac.MenuItem('Readme', icon='github'),
                    sac.MenuItem('使用手册', icon='question-diamond'),
                    sac.MenuItem('关于...', icon='link-45deg'),
                ]),
            ], open_all=True)
        elif st.session_state.examType == "training":
            if st.session_state.userType == "admin" or st.session_state.userType == 'supervisor':
                selected = sac.menu([
                    sac.MenuItem('主页', icon='house'),
                    sac.MenuItem('功能', icon='grid-3x3-gap', children=[
                        sac.MenuItem('生成题库', icon='list-task'),
                        sac.MenuItem('题库练习', icon='pencil-square'),
                        sac.MenuItem('数据录入', icon='database-add'),
                        sac.MenuItem('试题修改', icon='clipboard-check'),
                        sac.MenuItem('文件导出', icon='journal-arrow-down'),
                        sac.MenuItem('题库功能', icon='database-gear'),
                        sac.MenuItem('参数设置', icon='gear'),
                    ]),
                    sac.MenuItem('信息', icon='info-circle', children=[
                        sac.MenuItem('学习信息', icon='book'),
                        sac.MenuItem('证书及榜单', icon='bookmark-star'),
                    ]),
                    sac.MenuItem('查询', icon='search', children=[
                        sac.MenuItem('信息查询', icon='info-lg'),
                        sac.MenuItem('用户状态', icon='people'),
                        sac.MenuItem('操作日志', icon='incognito'),
                    ]),
                    sac.MenuItem('账户', icon='person-gear', children=[
                        sac.MenuItem('密码修改', icon='key'),
                        sac.MenuItem('密码重置', icon='bootstrap-reboot'),
                        sac.MenuItem('登出', icon='box-arrow-right'),
                    ]),
                    sac.MenuItem('关于', icon='layout-wtf', children=[
                        sac.MenuItem('Changelog', icon='view-list'),
                        sac.MenuItem('Readme', icon='github'),
                        sac.MenuItem('使用手册', icon='question-diamond'),
                        sac.MenuItem('彩蛋', icon='images'),
                        sac.MenuItem('关于...', icon='link-45deg'),
                    ]),
                ], open_index=[1], open_all=False)
            elif st.session_state.userType == "user":
                selected = sac.menu([
                    sac.MenuItem('主页', icon='house'),
                    sac.MenuItem('功能', icon='grid-3x3-gap', children=[
                        sac.MenuItem('生成题库', icon='list-task'),
                        sac.MenuItem('题库练习', icon='pencil-square'),
                    ]),
                    sac.MenuItem('信息', icon='info-circle', children=[
                        sac.MenuItem('学习信息', icon='book'),
                        sac.MenuItem('证书及榜单', icon='bookmark-star'),
                    ]),
                    sac.MenuItem('账户', icon='person-gear', children=[
                        sac.MenuItem('密码修改', icon='key'),
                        sac.MenuItem('登出', icon='box-arrow-right'),
                    ]),
                    sac.MenuItem('关于', icon='layout-wtf', children=[
                        sac.MenuItem('Changelog', icon='view-list'),
                        sac.MenuItem('Readme', icon='github'),
                        sac.MenuItem('使用手册', icon='question-diamond'),
                        sac.MenuItem('彩蛋', icon='images'),
                        sac.MenuItem('关于...', icon='link-45deg'),
                    ]),
                ], open_index=[1, 2, 3, 4, 5, 6], open_all=False)
        if st.session_state.userType == "supervisor":
            spv = getAllStations()
            st.session_state.StationCN = st.selectbox("请选择站室", options=spv[0], index=spv[1])
            sql = f"UPDATE users set StationCN = '{st.session_state.StationCN}' where userName = {st.session_state.userName}"
            execute_sql_and_commit(conn, cur, sql)
            preExamTypeIndex = 0
            if st.session_state.examType == "training":
                preExamTypeIndex = 0
            elif st.session_state.examType == "exam":
                preExamTypeIndex = 1
            tmpExamType = st.selectbox("请选择模式类型", options=["练习", "考试"], index=preExamTypeIndex)
            if tmpExamType == "练习":
                st.session_state.examType = "training"
                st.session_state.examName = "练习题库"
                st.session_state.examRandom = True
            elif tmpExamType == "考试":
                st.session_state.examType = "exam"
                st.session_state.examRandom = bool(getParam("考试题库每次随机生成", st.session_state.StationCN))
        st.write(f"### 姓名: :orange[{st.session_state.userCName}] 站室: :orange[{st.session_state.StationCN}]")
        st.caption("📢:red[**不要刷新页面, 否则会登出**]")
        #st.caption("**请使用 :red[[登出]] 功能退出页面, 否则会影响下次登录**")
    updatePyFileinfo()
    if selected != "密码重置" and selected != "用户状态" and selected != "操作日志":
        st.session_state.userPwRecheck = False
    if selected == "主页":
        displayBigTimeCircle()
        displayAppInfo()
        displayVisitCounter()

    elif selected == "生成题库" or selected == "选择考试":
        if st.session_state.examType == "training":
            #st.write("### :red[生成练习题库]")
            #st.markdown("<font face='微软雅黑' color=blue size=20><center>**生成练习题库**</center></font>", unsafe_allow_html=True)
            st.markdown("### <font face='微软雅黑' color=teal><center>生成练习题库</center></font>", unsafe_allow_html=True)
        elif st.session_state.examType == "exam":
            #st.markdown("<font face='微软雅黑' color=red size=20><center>**选择考试**</center></font>", unsafe_allow_html=True)
            st.markdown("### <font face='微软雅黑' color=red><center>选择考试</center></font>", unsafe_allow_html=True)
        if not st.session_state.examChosen or not st.session_state.calcScore:
            sql = "UPDATE verinfo set pyLM = 0 where pyFile = 'chapterChosenType'"
            execute_sql_and_commit(conn, cur, sql)
            training()
        else:
            st.error("你不能重复选择考试场次")
    elif selected == "题库练习" or selected == "开始考试":
        if st.session_state.examType == "exam":
            updateActionUser(st.session_state.userName, "考试", st.session_state.loginTime)
        elif st.session_state.examType == "training":
            updateActionUser(st.session_state.userName, "练习", st.session_state.loginTime)
        if "confirmSubmit" not in st.session_state:
            st.session_state.confirmSubmit = False
        if "examFinalTable" in st.session_state and "examName" in st.session_state and not st.session_state.confirmSubmit:
            sql = f"SELECT userName, examName from examresult GROUP BY userName, examName HAVING Count(examName) >= {st.session_state.examLimit} and userName = {st.session_state.userName} and examName = '{st.session_state.examName}'"
            if not execute_sql(cur, sql) or st.session_state.examType == "training":
                if st.session_state.calcScore:
                    calcScore()
                for key in st.session_state.keys():
                    if key.startswith("moption_") or key.startswith("textAnswer_"):
                        del st.session_state[key]
                displayTimeCountdown()
                qcol1, qcol2, qcol3, qcol4 = st.columns(4)
                sql = "SELECT * from " + st.session_state.examFinalTable + " order by ID"
                rows = execute_sql(cur, sql)
                quesCount = len(rows)
                preButton, nextButton, submitButton = False, False, False
                #st.write(f"Cur:{st.session_state.curQues} Comp:{st.session_state.flagCompleted}")
                if st.session_state.flagCompleted:
                    if st.session_state.curQues == 1:
                        preButton = qcol3.button("上题", icon=":material/arrow_back_ios:", disabled=True)
                    else:
                        preButton = qcol3.button("上题", icon=":material/arrow_back_ios:", on_click=changeCurQues, args=(-1, quesCount,))
                    if st.session_state.curQues == quesCount:
                        nextButton = qcol4.button("下题", icon=":material/arrow_forward_ios:", disabled=True)
                    else:
                        nextButton = qcol4.button("下题", icon=":material/arrow_forward_ios:", on_click=changeCurQues, args=(1, quesCount,))
                    submitButton = qcol1.button("交卷", icon=":material/publish:")
                elif st.session_state.confirmSubmit:
                    preButton = qcol3.button("上题", icon=":material/arrow_back_ios:", disabled=True)
                    nextButton = qcol4.button("下题", icon=":material/arrow_forward_ios:", disabled=True)
                    submitButton = qcol1.button("交卷", icon=":material/publish:", disabled=True)
                elif st.session_state.curQues == 0:
                    preButton = qcol3.button("上题", icon=":material/arrow_back_ios:", disabled=True)
                    nextButton = qcol4.button("下题", icon=":material/arrow_forward_ios:", on_click=changeCurQues, args=(1, quesCount,))
                    submitButton = qcol1.button("交卷", icon=":material/publish:", disabled=True)
                    exam(rows[0])
                elif st.session_state.curQues == 1:
                    preButton = qcol3.button("上题", icon=":material/arrow_back_ios:", disabled=True)
                    nextButton = qcol4.button("下题", icon=":material/arrow_forward_ios:", on_click=changeCurQues, args=(1, quesCount,))
                    submitButton = qcol1.button("交卷", icon=":material/publish:", disabled=True)
                elif st.session_state.curQues == quesCount:
                    preButton = qcol3.button("上题", icon=":material/arrow_back_ios:", on_click=changeCurQues, args=(-1, quesCount,))
                    nextButton = qcol4.button("下题", icon=":material/arrow_forward_ios:", disabled=True)
                    submitButton = qcol1.button("交卷", icon=":material/publish:")
                    st.session_state.flagCompleted = True
                elif 1 < st.session_state.curQues < quesCount:
                    preButton = qcol3.button("上题", icon=":material/arrow_back_ios:", on_click=changeCurQues, args=(-1, quesCount,))
                    nextButton = qcol4.button("下题", icon=":material/arrow_forward_ios:", on_click=changeCurQues, args=(1, quesCount,))
                    submitButton = qcol1.button("交卷", icon=":material/publish:", disabled=True)
                iCol1, iCol2 = st.columns(2)
                completedPack, cpStr, cpCount = [], "", 0
                sql = f"SELECT ID, qType from {st.session_state.examFinalTable} where userAnswer = '' order by ID"
                rows3 = execute_sql(cur, sql)
                for row3 in rows3:
                    completedPack.append(f"第{row3[0]}题 [{row3[1]}] 未作答")
                    cpStr = cpStr + str(row3[0]) + "/"
                sql = f"SELECT ID, qType from {st.session_state.examFinalTable} where userAnswer <> '' order by ID"
                rows3 = execute_sql(cur, sql)
                for row3 in rows3:
                    completedPack.append(f"第{row3[0]}题 [{row3[1]}] 已作答")
                cpCount = len(rows3)
                if cpCount == quesCount:
                    iCol1.caption(":orange[作答提示: 全部题目已作答]")
                elif quesCount - cpCount > 40:
                    iCol1.caption(f":blue[作答提示:] :red[你还有{quesCount - cpCount}道题未作答, 请尽快完成]")
                elif quesCount - cpCount > 0:
                    iCol1.caption(f":blue[作答提示:] :red[{cpStr[:-1]}] :blue[题还未作答, 可以在👉右测下拉列表中跳转]")
                else:
                    iCol1.caption(":red[你还未开始答题]")
                iCol2.selectbox(":green[答题卡] :red[[未答题前置排序]]", completedPack, index=None, on_change=quesGoto, key="chosenID")
                st.divider()
                if (preButton or nextButton or submitButton or st.session_state.goto) and not st.session_state.confirmSubmit:
                    sql = f"SELECT * from {st.session_state.examFinalTable} where ID = {st.session_state.curQues}"
                    row = execute_sql(cur, sql)[0]
                    if preButton or nextButton or st.session_state.goto:
                        if st.session_state.goto:
                            st.session_state.goto = False
                            st.write("#### :blue[跳转到指定题号: ]")
                        exam(row)
                    if submitButton:
                        emptyAnswer = "你没有作答的题为:第["
                        sql = f"SELECT ID from {st.session_state.examFinalTable} where userAnswer = '' order by ID"
                        rows2 = execute_sql(cur, sql)
                        for row2 in rows2:
                            emptyAnswer = emptyAnswer + str(row2[0]) + ", "
                        if emptyAnswer.endswith(", "):
                            emptyAnswer = emptyAnswer[:-2] + "]题, 请检查或直接交卷!"
                        else:
                            emptyAnswer = "你的所有题目均已作答, 确认交卷吗?"
                        submit_dialog(emptyAnswer)
                    preButton, nextButton, submitButton = False, False, False
            elif st.session_state.examType == "exam":
                st.info("你本场考试已达到次数限制, 无法再次进行, 如有疑问请联系管理员", icon="ℹ️")
        else:
            if st.session_state.examType == "training":
                st.info("请先生成新的题库", icon="ℹ️")
            elif st.session_state.examType == "exam":
                st.info("请先选择考试场次并点击开始考试", icon="ℹ️")
    elif selected == "数据录入":
        st.subheader(":orange[基础数据录入]", divider="violet")
        selectFunc = st.selectbox("请选择数据表", ["考试场次", "站室", "用户"], index=None, help="请选择数据表")
        stationCName = getStationCNALL()
        if selectFunc == "考试场次":
            buttonAdd = st.button("新增")
            if buttonAdd:
                addExamIDD()
        elif selectFunc == "站室":
            buttonAdd = st.button("新增")
            if buttonAdd:
                addStation()
        elif selectFunc == "用户":
            buttonAdd = st.button("新增")
            if buttonAdd:
                addUser()
        if selectFunc is not None:
            updateActionUser(st.session_state.userName, f"添加{selectFunc}", st.session_state.loginTime)
    elif selected == "试题修改":
        quesModify()
    elif selected == "文件导出":
        dboutput()
    elif selected == "题库功能":
        dbfunc()
    elif selected == "参数设置":
        st.subheader(":green[系统参数设置]")
        updateActionUser(st.session_state.userName, "设置系统参数", st.session_state.loginTime)
        with st.expander("# :blue[考试参数设置]"):
            col1, col2, col3, col4 = st.columns(4)
            col5, col6, col7 = st.columns(3)
            sql = f"SELECT paramName, param, ID from setup_{st.session_state.StationCN} where paramType = 'exam' order by ID"
            rows = execute_sql(cur, sql)
            for row in rows:
                if row[0] == "单题分值":
                    quesScore = row[1]
                if row[0] == "考题总数":
                    quesTotal = row[1]
                if row[0] == "单选题数量":
                    col1.slider(row[0], min_value=0, max_value=100, value=row[1], key=f"dasetup_{row[2]}")
                elif row[0] == "多选题数量":
                    col2.slider(row[0], min_value=0, max_value=100, value=row[1], key=f"dasetup_{row[2]}")
                elif row[0] == "判断题数量":
                    col3.slider(row[0], min_value=0, max_value=100, value=row[1], key=f"dasetup_{row[2]}")
                elif row[0] == "填空题数量":
                    col4.slider(row[0], min_value=0, max_value=100, value=row[1], key=f"dasetup_{row[2]}")
                elif row[0] == "单题分值":
                    col5.number_input(row[0], min_value=1, max_value=5, value=row[1], key=f"dasetup_{row[2]}", help="所有题型统一分值")
                elif row[0] == "考题总数":
                    col6.number_input(row[0], min_value=10, max_value=120, value=row[1], key=f"dasetup_{row[2]}", help="仅对考试有效, 练习模式不受限制")
                elif row[0] == "合格分数线":
                    st.slider(row[0], min_value=60, max_value=120, value=row[1], step=10, key=f"dasetup_{row[2]}", help=f"建议为{int(quesScore * quesTotal * 0.8)}分")
                elif row[0] == "同场考试次数限制":
                    col7.number_input(row[0], min_value=1, max_value=5, value=row[1], key=f"dasetup_{row[2]}", help="最多5次")
                elif row[0] == "考试题库每次随机生成":
                    #st.toggle(row[0], value=row[1], key=f"dasetup_{row[2]}", help="开启有效, 关闭无效")
                    sac.switch(label=row[0], value=row[1], key=row[0], on_label="On", align='start', size='md')
                    updateSwitchOption(row[0])
                elif row[0] == "考试时间":
                    st.slider(row[0], min_value=30, max_value=150, value=row[1], step=15, key=f"dasetup_{row[2]}", help="建议为60-90分钟")
                elif row[0] == "使用大模型评判错误的填空题答案":
                    sac.switch(label=row[0], value=row[1], key=row[0], on_label="On", align='start', size='md')
                    updateSwitchOption(row[0])
                else:
                    st.slider(row[0], min_value=1, max_value=150, value=row[1], key=f"dasetup_{row[2]}")
            updateDA = st.button("考试参数更新", on_click=updateDAParam, args=("考试",))
        with st.expander("# :red[章节权重设置]"):
            sql = "SELECT chapterName, examChapterRatio, ID from questionaff where chapterName <> '公共题库' and chapterName <> '错题集' and StationCN = '" + st.session_state.StationCN + "'"
            rows = execute_sql(cur, sql)
            if rows:
                sql = "SELECT chapterName, examChapterRatio, ID from questionaff where chapterName = '公共题库' and StationCN = '" + st.session_state.StationCN + "'"
                row = execute_sql(cur, sql)[0]
                st.slider(row[0], min_value=0, max_value=10, value=row[1], key=f"crsetup_{row[2]}", help="权重越大的章节占比越高")
                sql = "SELECT chapterName, examChapterRatio, ID from questionaff where chapterName = '错题集' and StationCN = '" + st.session_state.StationCN + "'"
                row = execute_sql(cur, sql)[0]
                st.slider(row[0], min_value=0, max_value=10, value=row[1], key=f"crsetup_{row[2]}", help="仅在练习题库中有效")
                for row in rows:
                    st.slider(row[0], min_value=0, max_value=10, value=row[1], key=f"crsetup_{row[2]}", help="权重越大的章节占比越高")
                st.button("章节权重更新", on_click=updateCRExam)
            else:
                st.info("该站室没有可设置章节")
        with st.expander("# :green[题型设置]"):
            sql = f"SELECT paramName, param from setup_{st.session_state.StationCN} where paramType = 'questype' order by ID"
            rows = execute_sql(cur, sql)
            for row in rows:
                sac.switch(label=row[0], value=row[1], key=row[0], on_label="On", align='start', size='md')
                updateSwitchOption(row[0])
        with st.expander("# :violet[导出文件字体设置]"):
            col20, col21, col22 = st.columns(3)
            sql = f"SELECT paramName, param, ID from setup_{st.session_state.StationCN} where paramType = 'fontsize' order by ID"
            rows = execute_sql(cur, sql)
            for row in rows:
                if row[0] == "抬头字体大小":
                    col20.number_input(row[0], min_value=8, max_value=32, value=row[1], key=f"dasetup_{row[2]}", help="题库导出至Word文件中的字体大小")
                elif row[0] == "题型字体大小":
                    col21.number_input(row[0], min_value=8, max_value=32, value=row[1], key=f"dasetup_{row[2]}")
                elif row[0] == "题目字体大小":
                    col22.number_input(row[0], min_value=8, max_value=32, value=row[1], key=f"dasetup_{row[2]}")
                elif row[0] == "选项字体大小":
                    col20.number_input(row[0], min_value=8, max_value=32, value=row[1], key=f"dasetup_{row[2]}")
                elif row[0] == "复核信息字体大小":
                    col21.number_input(row[0], min_value=8, max_value=32, value=row[1], key=f"dasetup_{row[2]}")
            updateDA = st.button("字体设置更新", on_click=updateDAParam, args=("字体设置",))
        with st.expander("# :orange[其他设置]"):
            sql = f"SELECT paramName, param, ID from setup_{st.session_state.StationCN} where paramType = 'others' order by ID"
            rows = execute_sql(cur, sql)
            for row in rows:
                if row[0] == "显示考试时间" or row[0] == "A.I.答案解析更新至题库" or row[0] == "测试模式":
                    sac.switch(label=row[0], value=row[1], key=row[0], on_label="On", align='start', size='md')
                    updateSwitchOption(row[0])
                elif row[0] == "时钟样式":
                    sac.switch(label=row[0], value=row[1], key=row[0], on_label="翻牌", off_label="数字", align='start', size='md')
                    updateSwitchOption(row[0])
            AIModel, AIModelIndex = [], 0
            sql = f"SELECT paramName, param, ID from setup_{st.session_state.StationCN} where paramName like '%大模型' and paramType = 'others' order by ID"
            rows = execute_sql(cur, sql)
            for index, value in enumerate(rows):
                AIModel.append(value[0])
                if value[1] == 1:
                    AIModelIndex = index
            st.radio("选择LLM大模型", options=AIModel, index=AIModelIndex, key="AIModel", horizontal=True, on_change=updateAIModel, help="讯飞输出质量高, 规范引用准确, 建议选用;文心千帆输出速度快, 内容可用;DeepSeek内容准确性相对高一些")
        st.divider()
        buttonReset = st.button("重置所有设置", type="primary")
        if buttonReset:
            buttonConfirm = st.button("确认重置", type="secondary", on_click=setupReset)
            updateActionUser(st.session_state.userName, "重置所有设置", st.session_state.loginTime)
    elif selected == "信息查询":
        st.subheader(":violet[信息查询]", divider="orange")
        selectFunc = st.selectbox("查询项目", ["考试信息", "未参加考试人员", "答题解析"], index=None)
        if selectFunc == "考试信息":
            queryExamResult()
        elif selectFunc == "未参加考试人员":
            queryExamResultUsers()
        elif selectFunc == "答题解析":
            queryExamName = st.selectbox("请选择考试场次", ["练习题库", "错题集"], index=0)
            if queryExamName:
                if queryExamName == "错题集":
                    tablename = "morepractise"
                else:
                    tablename = f"exam_final_{st.session_state.StationCN}_{st.session_state.userName}_{queryExamName}"
                sql = "SELECT * from sqlite_master where type = 'table' and name = '" + tablename + "'"
                tempTable = execute_sql(cur, sql)
                if tempTable:
                    queryExamAnswer(tablename)
                else:
                    st.info("暂无数据")
        if selectFunc is not None:
            updateActionUser(st.session_state.userName, f"查询{selectFunc}", st.session_state.loginTime)
    elif selected == "用户状态":
        userStatus()
    elif selected == "操作日志":
        displayKeyAction()
    elif selected == "学习信息":
        studyinfo()
    elif selected == "证书及榜单":
        userRanking()
    elif selected == "密码修改":
        changePassword()
    elif selected == "密码重置":
        resetPassword()
    elif selected == "登出":
        logout()
    elif selected == "Changelog":
        changelog()
    elif selected == "Readme":
        aboutReadme()
    elif selected == "使用手册":
        displayUserManual()
    elif selected == "彩蛋":
        aiGenerate_Image()
    elif selected == "关于...":
        aboutInfo()
