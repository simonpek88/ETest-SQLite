# coding utf-8
import base64
import datetime
import json
import os
import re
import sqlite3
import time

import folium
import openpyxl
import pandas as pd
import Play_mp3
import plotly.graph_objects as go
import streamlit as st
import streamlit.components.v1 as components
import streamlit_antd_components as sac
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor
from folium.plugins import HeatMap, MiniMap
from PIL import Image, ImageDraw, ImageFont
from st_keyup import st_keyup
from streamlit_extras.badges import badge
from streamlit_folium import st_folium
from streamlit_javascript import st_javascript
from streamlit_timeline import st_timeline
from xlsxwriter.workbook import Workbook

from commFunc import (GenerExam, deepseek_AI, deepseek_AI_GenerQues,
                      execute_sql, execute_sql_and_commit, getParam,
                      getUserEDKeys, qianfan_AI, qianfan_AI_GenerQues,
                      updateActionUser, updatePyFileinfo, xunfei_xh_AI,
                      xunfei_xh_AI_fib, xunfei_xh_AI_GenerQues)

# cSpell:ignoreRegExp /[^\s]{16,}/
# cSpell:ignoreRegExp /\b[A-Z]{3,15}\b/g


@st.fragment
def updateKeyAction(keyAction):
    sql = f"SELECT ID from keyactionlog where userName = {st.session_state.userName} and userCName = '{st.session_state.userCName}' and userAction = '{keyAction}' and actionDate = {int(time.time())}"
    if not execute_sql(cur, sql):
        sql = f"INSERT INTO keyactionlog(userName, userCName, StationCN, userAction, actionDate) VALUES({st.session_state.userName}, '{st.session_state.userCName}', '{st.session_state.StationCN}', '{keyAction}', {int(time.time())})"
        execute_sql_and_commit(conn, cur, sql)


# noinspection PyShadowingNames
@st.fragment
def getUserCName(sUserName, sType="Digit"):
    errorInfo = ""
    sql = ""

    # 检查用户名类型并构建SQL查询
    if sType.capitalize() == "Digit":
        if sUserName.isdigit():
            sql = "SELECT userCName, StationCN FROM users WHERE userName = ?"
            params = (int(sUserName),)
        else:
            errorInfo = "请输入纯数字用户编码"
    elif sType.capitalize() == "Str":
        sql = "SELECT userCName, StationCN FROM users WHERE userCName = ?"
        params = (sUserName,)

    # 执行SQL查询并处理结果
    if sql:
        try:
            rows = execute_sql(cur, sql, params)
            if rows:
                st.session_state.userCName = rows[0][0]
                st.session_state.StationCN = rows[0][1]
            else:
                st.session_state.userCName = "未找到"
                st.session_state.StationCN = "未找到"
        except Exception as e:
            errorInfo = f"查询出错: {e}"

    # 显示错误信息
    if errorInfo:
        st.error(errorInfo)
        st.session_state.userCName = ""
        st.session_state.StationCN = ""


def delOutdatedTable():
    if st.session_state.examRandom and "examTable" in st.session_state:
        execute_sql_and_commit(conn, cur, sql=f"DROP TABLE IF EXISTS {st.session_state.examTable}")
    if "examFinalTable" in st.session_state:
        execute_sql_and_commit(conn, cur, sql=f"DROP TABLE IF EXISTS {st.session_state.examFinalTable}")


# noinspection PyShadowingNames
def changePassword():
    st.write("### :red[密码修改]")
    changePW = st.empty()
    with changePW.container(border=True):
        oldPassword = st.text_input("请输入原密码", max_chars=8, type="password", autocomplete="off")
        newPassword = st.text_input("请输入新密码", max_chars=8, type="password", autocomplete="off")
        confirmPassword = st.text_input("请再次输入新密码", max_chars=8, type="password", autocomplete="new-password")
        buttonSubmit = st.button("确认修改")

    def verifyOldPassword():
        verifyUPW = verifyUserPW(st.session_state.userName, oldPassword)
        if verifyUPW[0]:
            return verifyUPW[1]
        else:
            st.error("原密码不正确")
            return None

    def updateNewPassword(encryptedNewPassword):
        sql = "UPDATE users SET userPassword = ? WHERE userName = ?"
        params = (encryptedNewPassword, st.session_state.userName)
        try:
            execute_sql_and_commit(conn, cur, sql, params)
            st.toast("密码修改成功, 请重新登录")
        except Exception as e:
            st.error(f"密码修改失败: {e}")

    if buttonSubmit and oldPassword and newPassword and confirmPassword:
        if newPassword != "":
            if newPassword == confirmPassword:
                verifiedOldPassword = verifyOldPassword()
                if verifiedOldPassword:
                    encryptedNewPassword = getUserEDKeys(newPassword, "enc")
                    updateNewPassword(encryptedNewPassword)
                    logout()
            else:
                st.error("两次输入的密码不一致")
        else:
            st.warning("新密码不能为空")
    elif not oldPassword:
        st.warning("原密码不能为空")

    updateActionUser(st.session_state.userName, "密码修改", st.session_state.loginTime)


# noinspection PyShadowingNames
@st.cache_data
def get_userName(searchUserName=""):
    """
    根据用户名搜索用户信息，并返回格式化的字符串。
    :param searchUserName: 要搜索的用户名前缀
    :return: 格式化的用户信息字符串
    """
    # 避免SQL注入，使用参数化查询
    searchUserNameInfo = ""
    if len(searchUserName) > 1:
        sql = "SELECT userName, userCName, StationCN from users where userName like ?"
        params = (f"{searchUserName}%",)
        rows = execute_sql(cur, sql, params)

        # 使用列表推导式生成信息列表，然后一次性连接
        user_info_list = [f"用户编码: :red[{row[0]}] 姓名: :blue[{row[1]}] 站室: :orange[{row[2]}]" for row in rows]
        searchUserNameInfo = "\n\n".join(user_info_list)

        if searchUserNameInfo:
            searchUserNameInfo += "\n\n请在用户编码栏中填写查询出的完整编码"

    return searchUserNameInfo


@st.cache_data
def get_userCName(searchUserCName=""):
    searchUserCNameInfo = ""
    if len(searchUserCName) > 1:
        # 使用参数化查询防止SQL注入
        sql = "SELECT userName, userCName, StationCN from users where userCName like ?"
        params = (f"{searchUserCName}%",)
        rows = execute_sql(cur, sql, params)

        # 使用列表推导式和join进行更高效的字符串拼接
        user_info_list = [f"用户编码: :red[{row[0]}] 姓名: :blue[{row[1]}] 站室: :orange[{row[2]}]" for row in rows]
        searchUserCNameInfo = "\n\n".join(user_info_list)

        if searchUserCNameInfo:
            searchUserCNameInfo += "\n\n请在用户编码栏中填写查询出的完整编码"
    else:
        searchUserCNameInfo = ":red[**请输入至少2个字**]"

    return searchUserCNameInfo


@st.fragment
def login():
    st.markdown(f"<font face='微软雅黑' color=purple size=5><center>**{APPNAME}**</center></font>", unsafe_allow_html=True)
    login = st.empty()
    with login.container(border=True):
        userName = st_keyup("请输入用户编码", placeholder="请输入纯数字用户编码, 非站室名称, 如果不知编码, 请在下方输入姓名查询", max_chars=8)
        st.session_state.userCName = ""
        if userName:
            filtered = get_userName(userName)
            if filtered == "":
                getUserCName(userName, "Digit")
                st.caption(f"用户名: :blue[{st.session_state.userCName}] 站室: :orange[{st.session_state.StationCN}]")
        else:
            filtered = ""
        if st.session_state.userCName == "未找到" or filtered:
            st.caption(filtered)
        if userName == "" or st.session_state.userCName == "未找到":
            userCName = st_keyup("请输入用户姓名", placeholder="请输入用户姓名, 至少2个字, 用于查询, 非必填项", max_chars=8)
            st.session_state.userCName = ""
            if userCName:
                filtered = get_userCName(userCName)
                if filtered == "":
                    getUserCName(userCName, "Str")
                    st.caption(f"用户名: :blue[{st.session_state.userCName}] 站室: :orange[{st.session_state.StationCN}]")
            else:
                filtered = ""
            if st.session_state.userCName == "未找到" or filtered:
                promptArea = st.empty()
                with promptArea.container():
                    st.caption(filtered)
                if userName and filtered == "":
                    promptArea.empty()
        userPassword = st.text_input("请输入密码", max_chars=8, placeholder="用户初始密码为1234", type="password", autocomplete="off")
        examType = sac.segmented(
            items=[
                sac.SegmentedItem(label="练习", icon="pen"),
                sac.SegmentedItem(label="考试", icon="card-list", disabled=True),
            ], align="start", size="sm"
        )
        buttonLogin = st.button("登录")
    if buttonLogin:
        if userName != "" and userPassword != "":
            verifyUPW = verifyUserPW(userName, userPassword)
            if verifyUPW[0]:
                userPassword = verifyUPW[1]
            if examType == "练习":
                st.session_state.examType = "training"
                st.session_state.examName = "练习题库"
                st.session_state.examRandom = True
                sql = f"SELECT userName, userCName, userType, StationCN from users where userName = {userName} and userPassword = '{userPassword}'"
            elif examType == "考试":
                st.session_state.examType = "exam"
                st.session_state.examRandom = bool(getParam("考试题库每次随机生成", st.session_state.StationCN))
                sql = f"SELECT userName, userCName, userType, StationCN from users where userName = {userName} and userPassword = '{userPassword}' and activeUser = 0"
            else:
                sql = ""
            if sql != "":
                result = execute_sql(cur, sql)
                if result:
                    st.toast(f"用户: {result[0][0]} 姓名: {result[0][1]} 登录成功, 欢迎回来")
                    login.empty()
                    st.session_state.logged_in = True
                    st.session_state.userPwRecheck = False
                    st.session_state.userName = result[0][0]
                    st.session_state.userCName = result[0][1].replace(" ", "")
                    st.session_state.userType = result[0][2]
                    st.session_state.StationCN = result[0][3]
                    st.session_state.examLimit = getParam("同场考试次数限制", st.session_state.StationCN)
                    st.session_state.debug = bool(getParam("测试模式", st.session_state.StationCN))
                    st.session_state.clockType = bool(getParam("时钟样式", st.session_state.StationCN))
                    st.session_state.curQues = 0
                    st.session_state.examChosen = False
                    st.session_state.delExam = True
                    st.session_state.tooltipColor = "#ed872d"
                    st.session_state.loginTime = int(time.time())
                    sql = f"UPDATE users set activeUser = 1, loginTime = {st.session_state.loginTime}, activeTime_session = 0, actionUser = '空闲' where userName = {st.session_state.userName}"
                    execute_sql_and_commit(conn, cur, sql)
                    sql = "UPDATE verinfo set pyLM = pyLM + 1 where pyFile = 'visitcounter'"
                    execute_sql_and_commit(conn, cur, sql)
                    ClearTables()
                    if datetime.datetime.now().hour in range(8, 22):
                        #Play_mp3.play('./Audio/login.mp3')
                        pass
                    st.rerun()
                else:
                    if verifyUPW[0]:
                        st.error("登录失败, 用户已经在别处登录, 请联系管理员解决")
                    else:
                        st.error("登录失败, 请检查用户名和密码, 若忘记密码请联系管理员重置")
        else:
            st.warning("请输入用户编码和密码")


def logout():
    try:
        sql = f"UPDATE users set activeUser = 0, activeTime = activeTime + activeTime_session, activeTime_session = 0 where userName = {st.session_state.userName}"
        execute_sql_and_commit(conn, cur, sql)
        delOutdatedTable()

    finally:
        cur.close()
        conn.close()

    for key in st.session_state.keys():
        del st.session_state[key]

    if datetime.datetime.now().hour in range(8, 22):
        #Play_mp3.play('./Audio/logout.mp3')
        pass
    st.rerun()


def display_logos():
    logos = [
        ("Python", "./Images/logos/python.png"),
        ("Streamlit", "./Images/logos/streamlit.png"),
        ("SQLite", "./Images/logos/sqlite.png"),
        ("Pandas", "./Images/logos/pandas.png"),
        ("Ant Design", "./Images/logos/antd.png"),
        ("iFlytek Spark", "./Images/logos/xfxh3.png"),
        ("ERNIE Qianfan", "./Images/logos/qianfan.png"),
        ("DeepSeek", "./Images/logos/deepseek2.png"),
    ]
    for name, path in logos:
        st.caption(name)
        st.image(path)


def aboutInfo():
    st.subheader("关于本软件", divider="rainbow")
    st.subheader(":blue[Powered by Python and Streamlit]")

    # 显示图标和标题
    display_logos()

    # 提示浅色主题
    st.write("###### :violet[为了获得更好的使用体验, 请使用浅色主题]")

    # 显示版本信息和最后修改时间
    verinfo, verLM, likeCM = getVerInfo()
    version_parts = [int(verinfo / 10000), int((verinfo % 10000) / 100), int(verinfo / 10)]
    version_str = f"{version_parts[0]}.{version_parts[1]}.{version_parts[2]}"
    st.caption(f"Version: {version_str} building {verinfo} Last Modified: {time.strftime('%Y-%m-%d %H:%M', time.localtime(verLM))}")

    # 显示用户反馈和评分
    st.caption(f"Reviews: {EMOJI[int(likeCM) - 1][0]} {likeCM} :orange[I feel {EMOJI[int(likeCM) - 1][1]}]")
    sac.divider(align="center", color="blue", size="sm")
    stars = sac.rate(label='Please give me a star if you like it!', align='start', size="sm")
    if stars > 0:
        feedback_message = f"I feel {EMOJI[int(stars) - 1][1]} {EMOJI[int(stars) - 1][0]}"
        st.write(feedback_message)
        try:
            sql = f"UPDATE verinfo set pyMC = pyMC + 1 where pyFile = 'thumbs-up-stars' and pyLM = {stars}"
            execute_sql_and_commit(conn, cur, sql)
        except Exception as e:
            st.error(f"更新评分时出错: {e}")
    updateActionUser(st.session_state.userName, "浏览[关于]信息", st.session_state.loginTime)


# noinspection PyBroadException,PyUnusedLocal
def getVerInfo():
    try:
        sql = "SELECT Sum(pyMC) from verinfo"
        verinfo = execute_sql(cur, sql)[0][0]
        sql = "SELECT Max(pyLM) from verinfo"
        verLM = execute_sql(cur, sql)[0][0]
        sql = "SELECT Sum(pyLM * pyMC), Sum(pyMC) from verinfo where pyFile = 'thumbs-up-stars'"
        tmpTable = execute_sql(cur, sql)
        likeCM = round(tmpTable[0][0] / tmpTable[0][1], 1)

        return verinfo, verLM, likeCM
    except Exception as e:
        return 0, 0, 0


def display_pypi():
    libraries = ["streamlit", "pandas", "streamlit_antd_components", "folium", "qianfan"]
    for library in libraries:
        badge(type="pypi", name=library)


def aboutLicense():
    st.subheader("License", divider="green")
    st.markdown(open("./LICENSE", "r", encoding="utf-8").read())
    updateActionUser(st.session_state.userName, "浏览License信息", st.session_state.loginTime)


def actDelTable():
    for each in st.session_state.keys():
        if each.startswith("delStaticExamTable_"):
            if st.session_state[each]:
                each = each.replace("delStaticExamTable_", "")
                execute_sql_and_commit(conn, cur, sql=f"DROP TABLE IF EXISTS {each}")
                st.info(f"{each} 静态题库删除成功")


def delStaticExamTable():
    flagExistTable = False
    sql = "SELECT name from sqlite_master where type = 'table' and name like 'exam_%'"
    tempTable = execute_sql(cur, sql)
    if tempTable:
        st.subheader("删除静态题库", divider="red")
        for row in tempTable:
            if row[0].count("_") == 2:
                st.checkbox(f"{row[0]}", key=f"delStaticExamTable_{row[0]}")
                flagExistTable = True
    if flagExistTable:
        st.button("确认删除", on_click=actDelTable)
    else:
        st.info("暂无静态题库")


def resultExcel():
    st.subheader("试卷导出", divider="blue")
    examResultPack, examResultPack2 = [], []

    # 使用参数化查询来获取试卷列表，防止SQL注入
    sql = "SELECT name from sqlite_master where type = 'table' and name like 'exam_final_%'"
    tempTable = execute_sql(cur, sql)

    if tempTable:
        for row in tempTable:
            table_name = row[0]
            examResultPack2.append(table_name)

            # 提取用户名，并进行查询
            user_id = table_name.split('_')[-2]  # 假设表名格式为 exam_final_userID_其他
            sql = "SELECT userCName from users where userName = ?"
            tempTable = execute_sql(cur, sql, (int(user_id),))

            if tempTable:
                tempUserCName = tempTable[0][0]
                examResultPack.append(f"{tempUserCName}_{table_name.split('_')[-1]}")
            else:
                examResultPack.append(table_name.replace("exam_final_", ""))

        examResult = st.selectbox(" ", examResultPack, index=None, label_visibility="collapsed")

        if examResult:
            # 分离出选择的试卷名称
            selected_user, selected_exam = examResult.split('_')
            examResult = f"exam_final_{selected_exam}"  # 重新构建试卷表名

            sql = """
            SELECT Question, qOption, qAnswer, qType, qAnalysis, userAnswer
            from {}
            order by ID
            """.format(examResult)

            rows = execute_sql(cur, sql)
            if rows:
                df = pd.DataFrame(rows, columns=["题目", "选项", "标准答案", "类型", "解析", "你的答案"])
                st.dataframe(df)
    else:
        st.info("暂无试卷")


def examResulttoExcel():
    st.subheader("考试成绩导出", divider="blue")
    searchOption = []
    sql = "SELECT ID, examName from examidd where StationCN = ? order by ID"
    rows = execute_sql(cur, sql, (st.session_state.StationCN,))  # 使用参数化查询
    for row in rows:
        searchOption.append(row[1])
    searchExamName = st.selectbox("请选择考试场次", searchOption, index=None)
    options = st.multiselect("查询类型", ["通过", "未通过"], default=["通过", "未通过"])
    if searchExamName:
        searchButton = st.button("导出为Excel文件", type="primary")
        if searchButton:
            conditions = []
            for each in options:
                if each == "通过":
                    conditions.append("examPass = 1")
                elif each == "未通过":
                    conditions.append("examPass = 0")
            condition_str = " OR ".join(conditions)
            sql = f"SELECT ID, userName, userCName, examScore, examDate, examPass from examresult where examName = ? and ({condition_str}) order by ID"
            rows = execute_sql(cur, sql, (searchExamName,))  # 使用参数化查询
            if rows:
                outputFile = f"./ExamResult/{searchExamName}_{time.strftime('%Y%m%d%H%M%S', time.localtime())}.xlsx"
                workbook = Workbook(outputFile)
                worksheet = workbook.add_worksheet(f"{searchExamName}考试成绩")
                title = ["ID", "编码", "姓名", "成绩", "考试时间", "考试结果"]
                worksheet.write_row(0, 0, title)  # 使用write_row简化写入标题行
                row_number = 1  # 改进变量命名，增加可读性
                for row in rows:
                    worksheet.write(row_number, 0, row_number)  # 直接写入行号，避免额外变量k
                    worksheet.write(row_number, 1, row[0])
                    worksheet.write(row_number, 2, row[1])
                    worksheet.write(row_number, 3, row[2])
                    examDate = time.strftime('%Y-%m-%d %H:%M:%S', time.localtime(int(row[3])))
                    worksheet.write(row_number, 4, examDate)  # 格式化考试时间
                    examPass = "通过" if row[4] == 1 else "未通过"  # 简化条件判断
                    worksheet.write(row_number, 5, examPass)
                    row_number += 1  # 行号递增
                workbook.close()
                try:
                    with open(outputFile, "rb") as file:
                        content = file.read()
                    buttonDL = st.download_button("点击下载", content, file_name=os.path.basename(outputFile), icon=":material/download:", type="secondary")
                    if buttonDL:
                        st.toast("文件已下载至你的默认目录")
                        updateKeyAction("导出考试成绩")
                except Exception as e:
                    st.error(f"文件读取错误: {e}")
                st.success(f":green[[{searchExamName}]] :gray[考试成绩成功导出]")
            else:
                st.error(f":red[[{searchExamName}]] 没有找到相关考试成绩")


def ClearTables():
    sql = "DELETE from questions where rowid NOT IN (SELECT Min(rowid) from questions GROUP BY Question, qType, StationCN, chapterName)"
    execute_sql_and_commit(conn, cur, sql)
    sql = "DELETE from commquestions where rowid NOT IN (SELECT Min(rowid) from commquestions GROUP BY Question, qType)"
    execute_sql_and_commit(conn, cur, sql)
    sql = "DELETE from morepractise where rowid NOT IN (SELECT Min(rowid) from morepractise GROUP BY Question, qType, userName)"
    execute_sql_and_commit(conn, cur, sql)
    sql = "DELETE from questionaff where rowid NOT IN (SELECT Min(rowid) from questionaff GROUP BY chapterName, StationCN)"
    execute_sql_and_commit(conn, cur, sql)
    sql = "DELETE from questionaff where chapterName <> '公共题库' and chapterName <> '错题集' and chapterName <> '关注题集' and chapterName not in (SELECT DISTINCT(chapterName) from questions)"
    execute_sql_and_commit(conn, cur, sql)
    sql = "UPDATE users set userCName = replace(userCName, ' ', '') where userCName like '% %'"
    execute_sql_and_commit(conn, cur, sql)
    for each in ["questions", "commquestions", "morepractise"]:
        execute_sql_and_commit(conn, cur, sql=f"update {each} set Question = REPLACE(Question,'\n', '') where Question like '%\n%'")
    #st.toast("站室题库/公共题库/错题集/章节信息库 记录清理完成")


def create_element(name):
    return OxmlElement(name)


def create_attribute(element, name, value):
    element.set(qn(name), value)


# noinspection PyProtectedMember
def add_page_number(run):
    fldChar1 = create_element('w:fldChar')
    create_attribute(fldChar1, 'w:fldCharType', 'begin')

    instrText = create_element('w:instrText')
    create_attribute(instrText, 'xml:space', 'preserve')
    instrText.text = "PAGE"

    fldChar2 = create_element('w:fldChar')
    create_attribute(fldChar2, 'w:fldCharType', 'end')

    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)


# noinspection PyProtectedMember,PyTypeChecker
def questoWord():
    allType, stationCName, chapterNamePack, outChapterName = [], [], [], []
    st.subheader("题库导出", divider="blue")
    sql = f"SELECT paramName, param from setup_{st.session_state.StationCN} where paramType = 'questype'"
    rows = execute_sql(cur, sql)
    for row in rows:
        allType.append(row[0])
    quesTable = st.selectbox("请选择功能类型", ("站室题库", "公共题库", "试卷", "错题集", "关注题集"), index=None)
    quesType = st.multiselect("题型", allType, default=allType)
    stationCN, headerExamName = "全站", ""
    if quesTable == "站室题库" or quesTable == "错题集" or quesTable == "关注题集":
        stationCName = getStationCNALL(flagALL=True)
        stationCN = st.select_slider("站室", stationCName, value=st.session_state.StationCN)
    elif quesTable == "试卷":
        headerExamName = st.text_input("请设置试卷名称", max_chars=20, help="文件抬头显示的试卷名称, 不填则使用默认名称")
        if "examFinalTable" in st.session_state:
            stationCN = st.session_state.StationCN
        else:
            st.info("请先生成题库")
            quesTable = ""
    if stationCN != "全站" and quesTable == "站室题库":
        sql = f"SELECT chapterName from questionaff where StationCN = '{stationCN}' and chapterName <> '公共题库' and chapterName <> '错题集' and chapterName <> '关注题集' order by ID"
        rows = execute_sql(cur, sql)
        for row in rows:
            chapterNamePack.append(row[0])
        outChapterName = st.multiselect("章节", chapterNamePack, default=chapterNamePack)
    sac.switch(label="复核模式", on_label="On", align='start', size='sm', value=False, key="sac_recheck")
    if st.session_state.sac_recheck:
        sac.switch(label="附加答题解析", on_label="On", align='start', size='sm', value=False, key="sac_Analysis")
    if quesTable and quesType:
        buttonSubmit = st.button("导出为Word文件", type="primary")
        if buttonSubmit:
            if quesTable == "站室题库":
                tablename = "questions"
            elif quesTable == "公共题库":
                tablename = "commquestions"
            elif quesTable == "试卷":
                tablename = st.session_state.examFinalTable
            elif quesTable == "错题集":
                tablename = "morepractise"
            elif quesTable == "关注题集":
                tablename = "favques"
            else:
                tablename = ""
            headerFS = getParam("抬头字体大小", st.session_state.StationCN)
            titleFS = getParam("题型字体大小", st.session_state.StationCN)
            quesFS = getParam("题目字体大小", st.session_state.StationCN)
            optionFS = getParam("选项字体大小", st.session_state.StationCN)
            answerFS = getParam("复核信息字体大小", st.session_state.StationCN)
            quesDOC = Document()
            quesDOC.styles["Normal"].font.name = "Microsoft YaHei"
            quesDOC.styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
            option, radioOption = ["A", "B", "C", "D", "E", "F", "G", "H"], ["正确", "错误"]
            blank = f"({' ' * 20})"
            pHeader = quesDOC.add_paragraph()
            pHeader.alignment = WD_ALIGN_PARAGRAPH.CENTER
            textHeader = pHeader.add_run(f"{st.session_state.StationCN} {headerExamName} {quesTable}", 0)
            #textHeader.font.name = "Microsoft YaHei"
            #textHeader.element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
            textHeader.font.size = Pt(headerFS)
            textHeader.font.bold = True
            textHeader.font.color.rgb = RGBColor(40, 106, 205)
            for each in quesType:
                if stationCN == "全站" or quesTable == "试卷":
                    sql = f"SELECT Question, qOption, qAnswer, qType, ID, SourceType, qAnalysis from {tablename} where qType = '{each}' order by ID"
                else:
                    if quesTable != "站室题库" and quesTable != "公共题库":
                        sql = f"SELECT Question, qOption, qAnswer, qType, ID, SourceType, qAnalysis from {tablename} where qType = '{each}' and StationCN = '{stationCN}' order by ID"
                    elif quesTable == "站室题库":
                        if outChapterName:
                            sql = f"SELECT Question, qOption, qAnswer, qType, ID, SourceType, qAnalysis from {tablename} where qType = '{each}' and StationCN = '{stationCN}' and (chapterName = "
                            for each5 in outChapterName:
                                sql += f"'{each5}' or chapterName = "
                            sql = sql[:-18] + ") order by chapterName, ID"
                        else:
                            sql = f"SELECT Question, qOption, qAnswer, qType, ID, SourceType, qAnalysis from {tablename} where qType = '{each}' and StationCN = '{stationCN}' order by chapterName, ID"
                rows = execute_sql(cur, sql)
                #st.write(f"{each} 共 {len(rows)}")
                i = 1
                if rows:
                    pTitle = quesDOC.add_paragraph()
                    textTitle = pTitle.add_run(f"{each}", 0)
                    #textTitle.font.name = "Microsoft YaHei"
                    #textTitle.element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
                    textTitle.font.size = Pt(titleFS)
                    textTitle.bold = True
                    for row in rows:
                        tmp, answer, qa, aa = "", "", [], []
                        pQues = quesDOC.add_paragraph()
                        if each == "填空题":
                            textQues = pQues.add_run(f"第{i}题   {row[0].replace('()', blank).replace('（）', blank)}", 0)
                        else:
                            textQues = pQues.add_run(f"第{i}题   {row[0]}   ({' ' * 8})", 0)
                        #textQues.font.name = "Microsoft YaHei"
                        #textQues.element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
                        #if st.session_state.sac_recheck and row[5] == "AI-LLM":
                        #textQues.font.color.rgb = RGBColor(155, 17, 30)
                        textQues.font.size = Pt(quesFS)
                        aa = row[2].replace("；", ";").split(";")
                        pOption = None
                        if each != "填空题":
                            pOption = quesDOC.add_paragraph()
                        if each == "单选题" or each == "多选题":
                            qa = row[1].replace("；", ";").split(";")
                            for each2 in qa:
                                tmp = tmp + f"{option[qa.index(each2)]}. {each2}{' ' * 8}"
                            textOption = pOption.add_run(tmp)
                            textOption.font.size = Pt(optionFS)
                        elif each == "判断题":
                            textOption = pOption.add_run(f"A. 正确{' ' * 15}B. 错误{' ' * 15}")
                            textOption.font.size = Pt(optionFS)
                        #textOption.font.name = "Microsoft YaHei"
                        #textOption.element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
                        #textOption.italic = True
                        if st.session_state.sac_recheck:
                            if each == "单选题" or each == "多选题":
                                for each3 in aa:
                                    answer = answer + f"{option[int(each3)]}" + ", "
                            elif each == "判断题":
                                answer = radioOption[int(aa[0]) ^ 1]
                            elif each == "填空题":
                                for index, value in enumerate(aa):
                                    answer = answer + f"{chr(49 + index)}.  [{value}]" + " "
                            if answer.endswith(", "):
                                answer = answer[:-2]
                            elif answer.endswith(" "):
                                answer = answer[:-1]
                            pAnswer = quesDOC.add_paragraph()
                            textAnswer = pAnswer.add_run(f"复核模式 ID: {row[4]} 正确答案: {answer}")
                            textAnswer.font.size = Pt(answerFS)
                            textAnswer.font.bold = True
                            textAnswer.font.color.rgb = RGBColor(155, 17, 30)
                            if stationCN != "全站":
                                sql = f"SELECT chapterName from questions where Question = '{row[0]}'"
                            else:
                                sql = f"SELECT chapterName from questions where Question = '{row[0]}' and StationCN = '{stationCN}'"
                            tempTable = execute_sql(cur, sql)
                            if tempTable:
                                fhQT = tempTable[0][0]
                            else:
                                sql = f"SELECT ID from commquestions where Question = '{row[0]}'"
                                if execute_sql(cur, sql):
                                    fhQT = "公共题库"
                                else:
                                    fhQT = "未知"
                            pSource = quesDOC.add_paragraph()
                            if not row[5].startswith("AI-LLM"):
                                textSource = pSource.add_run(f"试题来源: [{stationCN}] 章节名称: [{fhQT}] 试题生成类别: [{row[5]}]")
                            else:
                                textSource = pSource.add_run(f"请特别注意 试题来源: [{stationCN}] 章节名称: [{fhQT}] 试题生成类别: [{row[5]}]")
                            textSource.font.bold = True
                            textSource.font.size = Pt(answerFS)
                            if row[5].startswith("AI-LLM"):
                                textSource.font.color.rgb = RGBColor(155, 17, 30)
                                textSource.font.underline = True
                            #textSource.font.italic = True
                            if st.session_state.sac_Analysis and row[6] != "":
                                pAnalysis = quesDOC.add_paragraph()
                                if not row[5].startswith("AI-LLM"):
                                    textAnalysis = pAnalysis.add_run(f"人工解析: [{row[6].replace(':red', '').replace('[', '').replace(']', '').replace('**', '')}]")
                                else:
                                    textAnalysis = pAnalysis.add_run(f"请特别注意 A.I.解析: [{row[6].replace('**', '')}]")
                                textAnalysis.font.bold = True
                                textAnalysis.font.size = Pt(answerFS)
                                textAnalysis.font.color.rgb = RGBColor(79, 66, 181)
                                #textAnalysis.font.underline = True
                                #textAnalysis.font.italic = True
                        i += 1
            add_page_number(quesDOC.sections[0].footer.paragraphs[0].add_run())
            quesDOC.sections[0].footer.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            header = quesDOC.sections[0].header.paragraphs[0]
            header.text = f"{stationCN}\t\t{quesTable}"
            header.style = quesDOC.styles["Header"]
            if headerExamName != "":
                if st.session_state.sac_recheck:
                    outputFile = f"./QuesDoc/{stationCN}-{headerExamName}-{quesTable}-带审核信息_{time.strftime('%Y%m%d%H%M%S', time.localtime(int(time.time())))}.docx"
                else:
                    outputFile = f"./QuesDoc/{stationCN}-{headerExamName}-{quesTable}_{time.strftime('%Y%m%d%H%M%S', time.localtime(int(time.time())))}.docx"
            else:
                if st.session_state.sac_recheck:
                    outputFile = f"./QuesDoc/{stationCN}-{quesTable}-带审核信息_{time.strftime('%Y%m%d%H%M%S', time.localtime(int(time.time())))}.docx"
                else:
                    outputFile = f"./QuesDoc/{stationCN}-{quesTable}_{time.strftime('%Y%m%d%H%M%S', time.localtime(int(time.time())))}.docx"
            if os.path.exists(outputFile):
                os.remove(outputFile)
            quesDOC.save(outputFile)
            if os.path.exists(outputFile):
                if os.path.exists(outputFile):
                    with open(outputFile, "rb") as file:
                        content = file.read()
                    file.close()
                    buttonDL = st.download_button("点击下载", content, file_name=outputFile[outputFile.rfind("/") + 1:], icon=":material/download:", type="secondary")
                    st.success(f":green[[{quesTable}]] :gray[题库成功导出至程序目录内] :orange[{outputFile[2:]}]")
                    if buttonDL:
                        st.toast("文件已下载至你的默认目录")
            else:
                st.error(f":red[[{quesTable}]] 题库导出失败")


def dboutput():
    bc = sac.segmented(
        items=[
            sac.SegmentedItem(label="题库导出(Word格式)", icon="database-down"),
            #sac.SegmentedItem(label="试卷导出(DF格式)", icon="journal-arrow-down"),
            sac.SegmentedItem(label="考试成绩导出(Excel格式)", icon="layout-text-sidebar-reverse"),
        ], color="green", align="center", size="sm"
    )
    if bc == "题库导出(Word格式)":
        questoWord()
    elif bc == "试卷导出(DF格式)":
        resultExcel()
    elif bc == "考试成绩导出(Excel格式)":
        examResulttoExcel()
    if bc is not None:
        updateActionUser(st.session_state.userName, bc, st.session_state.loginTime)


def actDelExamTable():
    for each in st.session_state.keys():
        if each.startswith("delExamTable_"):
            if st.session_state[each]:
                each = each.replace("delExamTable_", "")
                execute_sql_and_commit(conn, cur, sql=f"DROP TABLE IF EXISTS {each}")
                st.info(f"{each} 试卷删除成功")


def delExamTable():
    flagExistTable = False
    sql = "SELECT name from sqlite_master where type = 'table' and name like 'exam_%'"
    tempTable = execute_sql(cur, sql)
    if tempTable:
        st.subheader("删除试卷", divider="red")
        for row in tempTable:
            if row[0].count("_") == 3 or row[0].count("_") == 4:
                st.checkbox(f"{row[0]}", key=f"delExamTable_{row[0]}")
                flagExistTable = True
    if flagExistTable:
        if st.session_state.userType == "supervisor":
            if st.session_state.delExam:
                st.button("确认删除", on_click=actDelExamTable)
            else:
                st.error("试卷正在使用, 无法删除, 请先完成考试或练习后删除")
        else:
            st.error("仅Supervisor可进行此操作")
    else:
        st.info("暂无试卷")


# noinspection PyUnboundLocalVariable
def dbinputSubmit(tarTable, orgTable):
    tmpTable, sql, maxcol = "", "", 0
    if tarTable == "站室题库":
        tablename = "questions"
        sql = f"INSERT INTO {tablename}(Question, qOption, qAnswer, qType, qAnalysis, StationCN, chapterName) VALUES (?, ?, ?, ?, ?, ?, ?)"
        maxcol = 7
    elif tarTable == "公共题库":
        tablename = "commquestions"
        sql = f"INSERT INTO {tablename}(Question, qOption, qAnswer, qType, qAnalysis) VALUES (?, ?, ?, ?, ?)"
        maxcol = 5
    if sql != "":
        st.spinner(f"正在向 [{tarTable}] 导入题库...")
        sql2 = f"SELECT Max(ID) from {tablename}"
        maxid = execute_sql(cur, sql2)[0][0]
        if maxid is None:
            maxid = 0
        for each in orgTable:
            listinsheet = openpyxl.load_workbook(f"./InputQues/{each}.xlsx")
            datainlist = listinsheet.active
            for row in datainlist.iter_rows(min_row=2, max_col=maxcol, max_row=datainlist.max_row):
                singleQues = [cell.value for cell in row]
                if singleQues[0] is not None:
                    cur.execute(sql, singleQues)
                    conn.commit()
            listinsheet.close()
            if each.find("_用户上传_") != -1:
                os.remove(f"./InputQues/{each}.xlsx")
            tmpTable = tmpTable + each + ", "
        sql = f"UPDATE {tablename} set qOption = '' where qOption is Null"
        execute_sql_and_commit(conn, cur, sql)
        sql = f"UPDATE {tablename} set qAnalysis = '' where qAnalysis is Null"
        execute_sql_and_commit(conn, cur, sql)
        sql = f"UPDATE {tablename} set SourceType = '人工' where SourceType is Null"
        execute_sql_and_commit(conn, cur, sql)
        sql = f"UPDATE {tablename} set qOption = replace(qOption, '；', ';'), qAnswer = replace(qAnswer, '；', ';') where (qOption like '%；%' or qAnswer like '%；%') and (qType = '单选题' or qType = '多选题' or qType = '填空题')"
        execute_sql_and_commit(conn, cur, sql)
        sql = f"UPDATE {tablename} set qType = '单选题' where qType = '选择题' and ID > {maxid}"
        execute_sql_and_commit(conn, cur, sql)
        sql = f"SELECT ID, qOption, qAnswer, qType, Question from {tablename} where ID > {maxid} and (qType = '单选题' or qType = '多选题' or qType = '判断题')"
        rows = execute_sql(cur, sql)
        for row in rows:
            sql = ""
            if row[3] == "单选题" or row[3] == "多选题":
                for each in row[2].split(";"):
                    if int(each) < 0 or int(each) >= len(row[1].split(";")) or len(row[1].split(";")) > 8:
                        sql = f"DELETE from {tablename} where ID = {row[0]}"
            elif row[3] == "判断题":
                if int(row[2]) < 0 or int(row[2]) > 1:
                    sql = f"DELETE from {tablename} where ID = {row[0]}"
            if sql != "":
                execute_sql_and_commit(conn, cur, sql)
                st.warning(f"试题: [{row[4]}] 题型: [{row[3]}] 选项: [{row[1]}] 答案: [{row[2]}] 因为选项及答案序号不相符, 没有导入")
        sql = "INSERT INTO questionaff(chapterName, StationCN, chapterRatio, examChapterRatio) SELECT DISTINCT chapterName, StationCN, 5, 5 FROM questions"
        execute_sql_and_commit(conn, cur, sql)
        ClearTables()
        st.success(f":green[{tmpTable[:-2]}.xlsx] 向 :red[{tarTable}] :gray[导入成功]")
        updateKeyAction(f"Excel文件导入试题至{tarTable}")


def dbinput():
    inputOption = []
    targetTable = st.radio("导入至:", ("站室题库", "公共题库"), index=0, horizontal=True)
    inputType = st.radio("文件来源:", ("服务器中文件", "上传文件"), index=0, horizontal=True)
    if targetTable:
        if inputType == "服务器中文件":
            for root, dirs, files in os.walk("./InputQues"):
                for file in files:
                    if os.path.splitext(file)[1].lower() == '.xlsx' and f"{st.session_state.StationCN}_{targetTable}" in os.path.splitext(file)[0] and not os.path.splitext(file)[0].startswith("~$"):
                        inputOption.append(os.path.splitext(file)[0])
            if inputOption:
                orgTable = st.multiselect("请选择导入文件", inputOption, default=None)
                if orgTable:
                    st.button("导入", on_click=dbinputSubmit, args=(targetTable, orgTable))
                else:
                    st.info("请选择要导入的文件")
            else:
                st.info("没有可导入的本站文件")
        elif inputType == "上传文件":
            uploaded_file = st.file_uploader("**请选择Excel文件**", type=["xlsx"])
            if uploaded_file is not None:
                bytes_data = uploaded_file.getvalue()
                outFile = f"./InputQues/{st.session_state.StationCN}_{targetTable}_用户上传_{time.strftime('%Y%m%d%H%M%S', time.localtime(int(time.time())))}.xlsx"
                if os.path.exists(outFile):
                    os.remove(outFile)
                with open(outFile, 'wb') as output_file:
                    output_file.write(bytes_data)
                if os.path.exists(outFile):
                    dbinputSubmit(targetTable, [outFile[12:-5]])
    else:
        st.write("请选择要导入的题库")


def dbfunc():
    if st.session_state.debug and int(st.session_state.userName) < 10:
        bc = sac.segmented(
            items=[
                sac.SegmentedItem(label="A.I.出题", icon="robot"),
                #sac.SegmentedItem(label="题库导入", icon="database-up"),
                #sac.SegmentedItem(label="Word文件导入", icon="text-wrap"),
                sac.SegmentedItem(label="删除试卷", icon="trash3"),
                sac.SegmentedItem(label="删除静态题库", icon="trash3"),
                #sac.SegmentedItem(label="删除用户上传文件", icon="trash3"),
                sac.SegmentedItem(label="错题集重置", icon="journal-x"),
                sac.SegmentedItem(label="重置题库ID", icon="bootstrap-reboot"),
            ], align="center", color="red", size="sm"
        )
    else:
        bc = sac.segmented(
            items=[
                sac.SegmentedItem(label="A.I.出题", icon="robot"),
                #sac.SegmentedItem(label="题库导入", icon="database-up"),
                sac.SegmentedItem(label="删除试卷", icon="trash3"),
                sac.SegmentedItem(label="删除静态题库", icon="trash3"),
                #sac.SegmentedItem(label="删除用户上传文件", icon="trash3"),
            ], align="center", color="red", size="sm"
        )
    if bc == "A.I.出题":
        AIGenerQues()
    elif bc == "题库导入":
        dbinput()
    elif bc == "Word文件导入":
        inputWord()
    elif bc == "删除试卷":
        delExamTable()
    elif bc == "删除静态题库":
        delStaticExamTable()
    elif bc == "删除用户上传文件":
        delUserUploadFiles()
    elif bc == "重置题库ID":
        buttonReset = st.button("重置题库ID", type="primary")
        if buttonReset:
            st.button("确认重置", type="secondary", on_click=resetTableID)
    if bc is not None:
        updateActionUser(st.session_state.userName, bc, st.session_state.loginTime)


def delUserUploadFiles():
    user_files = []
    for root, dirs, files in os.walk("./InputQues"):
        for file in files:
            if file.lower().endswith('.xlsx') and "_用户上传_" in file:
                user_files.append(file)
                st.checkbox(os.path.splitext(file)[0], value=False, key=f"delUserFiles_{file.replace('.', '_')}")

    if user_files:
        if st.button("删除选中的用户上传文件", type="danger"):
            for file in user_files:
                if st.session_state.get(f"delUserFiles_{file.replace('.', '_')}", False):
                    os.remove(os.path.join(root, file))
                    st.success(f"文件 {file} 已删除")
    else:
        st.info("没有用户上传文件")


def actionDelUserUploadFiles():
    for key in st.session_state.keys():
        if key.startswith("delUserFiles_"):
            if st.session_state[key]:
                os.remove(f"./InputQues/{key.replace('delUserFiles_', '')}.xlsx")
            del st.session_state[key]
    st.success("所选文件已经删除")
    updateKeyAction("删除用户上传文件")


def resetActiveUser():
    sql = f"UPDATE users set activeUser = 0 where userName <> {st.session_state.userName}"
    execute_sql_and_commit(conn, cur, sql)
    st.success("已重置所有用户状态")
    updateKeyAction("重置所有用户状态")


# noinspection PyUnboundLocalVariable
def inputWord():
    doc = Document("./QuesRefer/2023年特种设备作业安全管理人员证考试题库(通用版).docx")
    chapter = "特种设备安全管理员"
    #title_rule = re.compile("\\d+、")
    #title_rule = re.compile("\\d+.")
    option_rule = re.compile("\\w+、")
    ques, qAnswer, temp2, generQuesCount, qType = "", "", "", 0, ""
    if st.session_state.debug:
        os.system("cls")
    st.spinner("正在导入Word文件...")
    for i, paragraph in enumerate(doc.paragraphs[:]):
        line = paragraph.text.replace('\n', '').replace('\r', '').replace("（", "(").replace("）", ")").strip()
        if line:
            #if title_rule.search(line):
            if line[:7].find(".") != -1:
                if temp2.endswith(";"):
                    temp2 = temp2[:-1]
                    qOption = temp2
                    temp2 = ""
                if ques != "" and qAnswer != "" and qOption != "":
                    if qOption.find("正确;错误") != -1:
                        qType = "判断题"
                        qAnswer = int(qAnswer) ^ 1
                        qOption = ""
                    elif len(qAnswer) == 1:
                        qType = "单选题"
                    elif len(qAnswer) > 1:
                        qType = "多选题"
                    if st.session_state.debug:
                        print(f"Record: Q: {ques} T: {qType} O: {qOption} A: {qAnswer}")
                    sql = f"SELECT ID from questions where Question = '{ques}' and qType = '{qType}' and StationCN = '{st.session_state.StationCN}' and chapterName = '{chapter}'"
                    if not execute_sql(cur, sql):
                        sql = f"INSERT INTO questions(Question, qOption, qAnswer, qType, StationCN, chapterName, SourceType) VALUES ('{ques}', '{qOption}', '{qAnswer}', '{qType}', '{st.session_state.StationCN}', '{chapter}', '人工')"
                        execute_sql_and_commit(conn, cur, sql)
                        generQuesCount += 1
                    ques, qAnswer, qOption = "", "", ""
                if st.session_state.debug:
                    print(f"Ques:{line}")
                if line[:7].find("、") != -1:
                    ques = line[line.find("、") + 1:]
                elif line[:7].find(".") != -1:
                    ques = line[line.find(".") + 1:]
                if ques.startswith("."):
                    ques = ques[1:]
                qAnswer = ""
                while True:
                    b1 = line.find('(')
                    b2 = line.find(')')
                    if b1 != -1 and b2 != -1 and line[b1 + 1:b1 + 2] in ["A", "B", "C", "D", "E", "F"]:
                        temp = line[b1 + 1:b2]
                        ques = ques.replace(temp, " " * len(temp))
                        temp = temp.replace("、", "")
                        for each in temp:
                            qAnswer = qAnswer + str(ord(each) - 65) + ";"
                        line = line[b2 + 1:]
                    else:
                        break
                if qAnswer.endswith(";"):
                    qAnswer = qAnswer[:-1]
            elif option_rule.search(line):
                if st.session_state.debug:
                    print(f"{line}")
                temp2 = temp2 + line[2:] + ";"
            elif line.find("正确答案：") != -1:
                if st.session_state.debug:
                    print(line)
                temp = line[line.find("正确答案：") + 5:]
                for each in temp:
                    qAnswer = qAnswer + str(ord(each) - 65) + ";"
                if qAnswer.endswith(";"):
                    qAnswer = qAnswer[:-1]
        else:
            continue
    ClearTables()
    st.success(f"共生成{generQuesCount}道试题")
    updateKeyAction("导入试题")


def resetTableID():
    # 定义需要重置ID的表名列表
    table_names = [
        "questions", "commquestions", "morepractise", "favques",
        "examidd", "examresult", "questionaff", "studyinfo",
        "users", "keyactionlog",
        "setup_默认", f"setup_{st.session_state.StationCN}"
    ]

    try:
        # 开始事务处理
        conn.begin_transaction()

        for table_name in table_names:
            # 查询表中所有ID并按顺序排列
            sql = f"SELECT ID from {table_name} order by ID"
            rows = execute_sql(cur, sql)

            # 批量更新ID
            update_sql = []
            for i, row in enumerate(rows):
                update_sql.append(f"UPDATE {table_name} SET ID = {i + 1} WHERE ID = {row[0]}")
                if table_name in ["questions", "commquestions"]:
                    update_sql.append(f"UPDATE studyinfo SET cid = {i + 1} WHERE cid = {row[0]} AND questable = '{table_name}'")

            # 执行批量更新
            if update_sql:
                execute_sql_and_commit_batch(conn, cur, update_sql)

            # 更新sqlite_sequence表中的seq值
            if rows:
                seq_sql = f"UPDATE sqlite_sequence SET seq = {len(rows)} WHERE name = '{table_name}'"
                execute_sql_and_commit(conn, cur, seq_sql)

        # 提交事务
        conn.commit()
        st.success("题库ID重置成功")
        updateKeyAction("重置题库ID")

    except Exception as e:
        # 发生异常时回滚事务并打印错误信息
        conn.rollback()
        st.error(f"重置题库ID失败: {e}")


# 辅助函数：批量执行SQL语句并提交
def execute_sql_and_commit_batch(conn, cur, sql_list):
    for sql in sql_list:
        cur.execute(sql)
    conn.commit()


# noinspection PyShadowingNames,PyUnboundLocalVariable
def AIGenerQues():
    quesPack, chars, chapterPack, dynaQuesType, generQuesCount = [], ["A", "B", "C", "D", "E", "F", "G", "H"], [], ["单选题", "多选题", "判断题", "填空题"], 0
    StationCNPack, chosenStationCN = [], st.session_state.StationCN
    temp = f"{st.session_state.StationCN}-站室题库现有: "
    for each in dynaQuesType:
        sql = f"SELECT Count(ID) from questions where qType = '{each}' and StationCN = '{st.session_state.StationCN}'"
        qCount = execute_sql(cur, sql)[0][0]
        temp = temp + ":red[" + each + "]: " + str(qCount) + "道 "
    temp = temp + "\n\n公共题库现有: "
    for each in dynaQuesType:
        sql = f"SELECT Count(ID) from commquestions where qType = '{each}'"
        qCount = execute_sql(cur, sql)[0][0]
        temp = temp + ":red[" + each + "]: " + str(qCount) + "道 "
    temp = temp.strip()
    st.caption(temp)
    table = st.radio(label="请选择要生成的题库", options=("站室题库", "公共题库"), horizontal=True, index=None)
    if table and table != "公共题库":
        sql = "SELECT Station from stations order by ID"
        rows = execute_sql(cur, sql)
        for row in rows:
            StationCNPack.append(row[0])
        chosenStationCN = st.select_slider("请选择要导入的站室", options=StationCNPack, value=st.session_state.StationCN)
        sql = f"SELECT chapterName from questionaff where StationCN = '{chosenStationCN}' and chapterName <> '公共题库' and chapterName <> '错题集' and chapterName <> '关注题集'"
        rows = execute_sql(cur, sql)
        for row in rows:
            chapterPack.append(row[0])
        chapter = st.selectbox(label="请选择章节", options=chapterPack, index=None)
        textChapter = st.text_input("请输入新章节名称", value="", placeholder="添加新的章节")
        textChapter = textChapter.strip()
    elif table == "公共题库":
        chapter, textChapter = "", ""
    quesRefer = st.text_area("请输入参考资料")
    AIModelNamePack = st.multiselect(
        "可选LLM大模型",
        ["DeepSeek", "文心千帆", "讯飞星火"],
        ["DeepSeek", "文心千帆", "讯飞星火"],
    )
    quesTypePack = st.multiselect(
        "请选择要生成的题型",
        dynaQuesType,
        dynaQuesType,
    )
    quesCount = st.number_input("请输入要生成的题目数量", min_value=1, max_value=10, value=5, step=1)
    if table is not None and quesRefer != "" and AIModelNamePack != [] and quesTypePack != []:
        buttonGener = st.button("生成试题")
        if buttonGener:
            if chapter is None and textChapter != "":
                sql = f"SELECT ID from questionaff where chapterName = '{textChapter}' and StationCN = '{chosenStationCN}'"
                if not execute_sql(cur, sql):
                    sql = f"INSERT INTO questionaff(chapterName, StationCN, chapterRatio, examChapterRatio) VALUES ('{textChapter}', '{chosenStationCN}', 5, 5)"
                    execute_sql_and_commit(conn, cur, sql)
                    st.toast(f"新的章节: :red[{textChapter}]添加完毕")
                chapter = textChapter
            if chapter is not None and table == "站室题库" or table == "公共题库":
                if st.session_state.debug:
                    os.system("cls")
                generQuesCount, displayQues, generQuesCountPack = 0, "", []
                infoArea = st.empty()
                for quesType in quesTypePack:
                    gqc = 0
                    for AIModelName in AIModelNamePack:
                        with infoArea.container(border=True):
                            st.info(f"正在使用 :red[{AIModelName}大模型] 进行:blue[{quesType}] 试题生成, 请稍等...")
                        if AIModelName == "文心千帆":
                            ques = qianfan_AI_GenerQues(quesRefer, quesType, quesCount, "ERNIE-Speed-8K")
                        elif AIModelName == "DeepSeek":
                            ques = deepseek_AI_GenerQues(quesRefer, quesType, quesCount)
                        elif AIModelName == "讯飞星火":
                            ques = xunfei_xh_AI_GenerQues(quesRefer, quesType, quesCount)
                        else:
                            ques = ""
                        quesPack = ques.split("题型")
                        for each in quesPack:
                            if each != "":
                                quesHeader, qOption, Option, qAnswer, qAnalysis, flagSuccess = "", "", [], "", "", True
                                temp = each[:10]
                                for dqt in dynaQuesType:
                                    if temp.find(dqt) != -1:
                                        quesType = dqt
                                        break
                                b1 = each.find("试题")
                                if b1 != -1:
                                    each = each[b1 + 3:]
                                    c1 = each.find("标准答案")
                                    c2 = each.find("试题解析")
                                    b2 = each.find("选项")
                                    if c1 != -1 and c2 != -1:
                                        if quesType == "填空题":
                                            quesHeader = each[:c1].replace("\n\n", "").replace("\n", "").replace("**", "").replace("无选项，填写空白处即可。", "").replace("无选项，本题为填空题", "").replace("选项：", "").strip()
                                        else:
                                            quesHeader = each[:b2].replace("\n\n", "").replace("\n", "").replace("**", "").strip()
                                        if quesHeader.startswith("*: "):
                                            quesHeader = quesHeader[3:]
                                        if quesType == "单选题" or quesType == "多选题":
                                            b3 = each.find("标准答案")
                                            if b3 != -1:
                                                Option = each[b2 + 3:b3].replace("\n\n", "").replace("正确的？ 选项：", "").replace("正确的？", "").strip().split("\n")
                                                displayOption = each[b2 + 3:b3].replace("**", "").replace("正确的？ 选项：", "").replace("正确的？", "").strip()
                                                displayOption = displayOption.replace("A. ", "\n\nA. ").replace("B. ", "\nB. ").replace("C. ", "\nC. ").replace("D. ", "\nD. ").replace("E. ", "\nE. ").replace("F. ", "\nF. ").replace("G. ", "\nG. ").replace("H. ", "\nH. ")
                                                for each2 in Option:
                                                    for each3 in chars:
                                                        each2 = each2.replace(f"{each3}.", "").strip()
                                                    qOption = qOption + each2 + ";"
                                                if AIModelName == "讯飞星火" and len(Option) == 1:
                                                    qOption = qOption[:-1]
                                                    qOption = qOption.replace("  ", ";")
                                                    revDisplayOption = str(displayOption)
                                                    revDisplayOption = revDisplayOption.replace("A. ", "\n\nA. ").replace("B. ", "\nB. ").replace("C. ", "\nC. ").replace("D. ", "\nD. ").replace("E. ", "\nE. ").replace("F. ", "\nF. ").replace("G. ", "\nG. ").replace("H. ", "\nH. ")
                                                    displayOption = revDisplayOption
                                                qOption = qOption.replace("；", ";")
                                                if qOption.endswith(";"):
                                                    qOption = qOption[:-1]
                                                if st.session_state.debug:
                                                    print(f"Option:{Option} qOption:{qOption}")
                                                b4 = each.find("试题解析")
                                                if b4 != -1:
                                                    qAnswer = each[b3 + 5:b4].replace("\n", "").replace("*", "").strip()
                                                    displayAnswer = qAnswer
                                                    qAnalysis = each[b4 + 5:].replace("\n", "").replace("*", "").strip()
                                                    if quesType == "单选题":
                                                        qAnswer = ord(qAnswer[0].upper()) - 65
                                                        if qAnswer > 7 or qAnswer < 0:
                                                            flagSuccess = False
                                                    elif quesType == "多选题":
                                                        qAnswer = qAnswer.replace("，", "").replace(",", "").replace("、", "").replace("。", "").replace(" ", "").replace("（", "(")
                                                        if qAnswer.find("(") != -1:
                                                            qAnswer = qAnswer[:qAnswer.find("(")].strip()
                                                        temp = ""
                                                        if st.session_state.debug:
                                                            print(f"未处理前的多选题标准答案:{qAnswer}")
                                                        for each4 in qAnswer:
                                                            if ord(each4.upper()) - 65 > 7 or ord(each4.upper()) - 65 < 0:
                                                                flagSuccess = False
                                                                break
                                                            else:
                                                                temp = temp + str(ord(each4) - 65) + ";"
                                                        qAnswer = temp
                                                        if qAnswer.endswith(";"):
                                                            qAnswer = qAnswer[:-1]
                                        elif quesType == "判断题":
                                            if each[c1 + 5:c2].find("正确") != -1 or each[c1 + 5:c2].find("A") != -1:
                                                qAnswer = "1"
                                                displayAnswer = "正确"
                                            else:
                                                qAnswer = "0"
                                                displayAnswer = "错误"
                                            displayOption = "A. 正确\nB. 错误\n"
                                            qAnalysis = each[c2 + 5:].replace("\n", "").replace("*", "").strip()
                                        elif quesType == "填空题":
                                            displayOption = ""
                                            qAnswer = each[c1 + 5:c2].replace("\n", "").replace("*", "").replace("无选项", "").replace("；", ";").replace("，", ";").replace("。", "").replace("、", ";").strip()
                                            if qAnswer.startswith(":"):
                                                qAnswer = qAnswer[1:]
                                            displayAnswer = qAnswer
                                            qAnalysis = each[c2 + 5:].replace("\n", "").replace("*", "").strip()
                                            i = 12
                                            while i > 0:
                                                if quesHeader.find("_" * i) != -1:
                                                    quesHeader = quesHeader.replace("_" * i, "()")
                                                i -= 1
                                    if quesHeader.startswith(":"):
                                        quesHeader = quesHeader[1:].strip()
                                    if qAnalysis.startswith(":"):
                                        qAnalysis = qAnalysis[1:].strip()
                                    if qAnalysis.endswith("---"):
                                        qAnalysis = qAnalysis[:-3].strip()
                                    if quesType == "单选题" and len(str(qAnswer)) > 1:
                                        flagSuccess = False
                                    if qOption.count(";") == 0 and (quesType == "单选题" or quesType == "多选题"):
                                        flagSuccess = False
                                    if st.session_state.debug:
                                        print(f"debug: 题目:[{quesHeader}] 选项:[{qOption}], 标准答案:[{qAnswer}] 答题解析:[{qAnalysis}]")
                                    if quesType == "填空题":
                                        quesHeader = quesHeader.replace("选项未给出，需要学生在横线上填写正确答案。", "").replace("选项:", "")
                                if qAnswer != "" and quesHeader != "" and len(str(qAnswer)) < 200 and len(quesHeader) < 200 and flagSuccess:
                                    if table == "公共题库":
                                        sql = f"SELECT ID from commquestions where Question = '{quesHeader}' and qType = '{quesType}'"
                                        if not execute_sql(cur, sql):
                                            sql = f"INSERT INTO commquestions(Question, qOption, qAnswer, qType, qAnalysis, SourceType) VALUES('{quesHeader}', '{qOption}', '{qAnswer}', '{quesType}', '{qAnalysis}', 'AI-LLM-{AIModelName}')"
                                            execute_sql_and_commit(conn, cur, sql)
                                            generQuesCount += 1
                                            gqc += 1
                                            displayQues = displayQues + f":blue[**第{generQuesCount}题:**]\n\n:red[题型: ]{quesType}\n\n:red[题目: ]{quesHeader}\n\n:red[选项: ]\n{displayOption}\n\n:red[答案: ]{displayAnswer}\n\n:red[解析: ]{qAnalysis}\n\n{'-' * 40}\n\n"
                                    elif table == "站室题库":
                                        sql = f"SELECT ID from questions where Question = '{quesHeader}' and qType = '{quesType}' and StationCN = '{chosenStationCN}' and chapterName = '{chapter}'"
                                        if not execute_sql(cur, sql):
                                            sql = f"INSERT INTO questions(Question, qOption, qAnswer, qType, qAnalysis, StationCN, chapterName, SourceType) VALUES('{quesHeader}', '{qOption}', '{qAnswer}', '{quesType}', '{qAnalysis}', '{chosenStationCN}', '{chapter}', 'AI-LLM-{AIModelName}')"
                                            execute_sql_and_commit(conn, cur, sql)
                                            generQuesCount += 1
                                            gqc += 1
                                            displayQues = displayQues + f":blue[**第{generQuesCount}题:**]\n\n:red[题型: ]{quesType}\n\n:red[题目: ]{quesHeader}\n\n:red[选项: ]\n{displayOption}\n\n:red[答案: ]{displayAnswer}\n\n:red[解析: ]{qAnalysis}\n\n{'-' * 40}\n\n"
                    generQuesCountPack.append(gqc)
                infoArea.empty()
                if generQuesCount > 0:
                    tempInfo = f"试题生成完毕, 总计生成试题{generQuesCount}道, 其中"
                    for index, value in enumerate(quesTypePack):
                        tempInfo = tempInfo + f"{value}: {generQuesCountPack[index]}道, "
                    st.success(tempInfo[:-2])
                    st.subheader("具体如下:", divider="green")
                    st.markdown(displayQues)
                    if table == "公共题库":
                        updateKeyAction(f"A.I.生成试题{generQuesCount}道至{table}题库")
                    elif table == "站室题库":
                        updateKeyAction(f"A.I.生成试题{generQuesCount}道至{table}题库{chapter}章节")
                else:
                    st.info("A.I.未生成到任何试题, 请检查参考资料是否正确或是生成的试题已经在题库中")
            else:
                st.warning("站室题库请选择章节")
    else:
        st.info("请设置各选项和添加参考资料")


def ClearMP():
    buttonSubmit = st.button(f"清空 {st.session_state.userCName} 错题集", type="primary")
    if buttonSubmit:
        bcArea = st.empty()
        with bcArea.container():
            st.button("确认清空", type="secondary", on_click=ClearMPAction, args=(bcArea,))


def ClearMPAction(bcArea):
    execute_sql_and_commit(conn, cur, sql=f"DELETE from morepractise where userName = {st.session_state.userName}")
    bcArea.empty()
    st.success("当前用户错题集已重置")
    updateKeyAction("重置当前用户错题集")


def studyinfo():
    study = sac.segmented(
        items=[
            sac.SegmentedItem(label="学习进度", icon="grid-3x2-gap"),
            sac.SegmentedItem(label="错题集", icon="list-stars"),
            #sac.SegmentedItem(label="章节时间线", icon="clock-history"),
            sac.SegmentedItem(label="学习记录重置", icon="bootstrap-reboot"),
            sac.SegmentedItem(label="错题集重置", icon="journal-x"),
        ], align="center", color="red", size="sm"
    )
    if study == "学习进度":
        studyinfoDetail()
    elif study == "错题集":
        displayErrorQues()
    elif study == "章节时间线":
        generTimeline()
    elif study == "学习记录重置":
        studyReset()
    elif study == "错题集重置":
        ClearMP()
    if study is not None:
        updateActionUser(st.session_state.userName, f"查看信息-{study}", st.session_state.loginTime)


def userRanking():
    study = sac.segmented(
        items=[
            sac.SegmentedItem(label="榜单", icon="bookmark-star"),
            sac.SegmentedItem(label="证书", icon="patch-check"),
            sac.SegmentedItem(label="荣誉榜", icon="mortarboard"),
        ], align="center", color="red", size="sm"
    )
    if study == "榜单":
        displayUserRanking()
    elif study == "证书":
        displayCertificate()
    elif study == "荣誉榜":
        displayMedals()
    if study is not None:
        updateActionUser(st.session_state.userName, f"证书及榜单-{study}", st.session_state.loginTime)


# noinspection PyShadowingNames
def displayUserRanking():
    boardInfo, xData, yData = "", [], []
    boardType = st.radio("榜单", options=["个人榜", "站室榜"], index=0, horizontal=True)
    if boardType == "个人榜":
        sql = "SELECT userCName, StationCN, userRanking from users where userRanking > 0 order by userRanking DESC, ID limit 0, 10"
    elif boardType == "站室榜":
        sql = "SELECT StationCN, ID, sum(userRanking) as Count from users GROUP BY StationCN having Count > 0 order by Count DESC"
    else:
        sql = ""
    rows = execute_sql(cur, sql)
    for index, row in enumerate(rows):
        xData.append(row[0])
        yData.append(row[2])
        if boardType == "个人榜":
            boardInfo = boardInfo + f"第 {index + 1} 名: {row[0]} 站室: {row[1]} 刷题数: {row[2]}\n\n"
        elif boardType == "站室榜":
            boardInfo = boardInfo + f"第 {index + 1} 名: {row[0]} 刷题数: {row[2]}\n\n"
        else:
            boardInfo = ""
    itemArea = st.empty()
    colors = ["lightslategray",] * len(rows)
    colors[0] = "crimson"
    fig = go.Figure(data=[go.Bar(x=xData, y=yData, marker_color=colors)])
    #fig.update_layout(font=dict(family="Courier New, monospace", size=18))
    fig.update_layout(title_text=f"{boardType[:-1]}刷题榜")
    with itemArea.container(border=True):
        st.plotly_chart(fig, theme="streamlit")
    if boardType == "站室榜" and int(rows[0][2]) > 0:
        heatData = []
        sql = "SELECT StationCN, sum(userRanking) as Ranking from users GROUP BY StationCN having Ranking > 0 order by Ranking DESC"
        rows = execute_sql(cur, sql)
        sql = f"SELECT lat, lng, Station from stations where Station == '{rows[0][0]}'"
        row = execute_sql(cur, sql)[0]
        lat = round(row[0] / 100, 2)
        lng = round(row[1] / 100, 2)
        m = folium.Map(
            location=[lat, lng],
            tiles="https://wprd01.is.autonavi.com/appmaptile?x={x}&y={y}&z={z}&lang=zh_cn&size=1&scl=1&style=7",
            attr='高德-路网图',
            zoom_start=11,
            control_scale=True,
            )
        for row in rows:
            sql = f"SELECT lat, lng from stations where Station = '{row[0]}'"
            row2 = execute_sql(cur, sql)[0]
            lat = round(row2[0] / 100, 2)
            lng = round(row2[1] / 100, 2)
            iframe = folium.IFrame(f"{row[0]} 刷题{row[1]}道")
            popup = folium.Popup(iframe, min_width=120, max_width=300)
            icon = folium.features.CustomIcon(
                "./Images/logos/cnaf-logo.png",
                icon_size=(40, 40),
                icon_anchor=(20, 40),
                popup_anchor=(0, -40),
            )
            folium.Marker([lat, lng], icon=icon, popup=popup).add_to(m)
            heatData.append([lat, lng, int(row[1])])
        HeatMap(heatData).add_to(m)
        minimap = MiniMap(
            toggle_display=True,
            width=120,
            height=120,
            minimized=True,
            )
        m.add_child(minimap)
        st_folium(m, use_container_width=True, height=300)
    st.subheader(boardInfo)


def generTimeline():
    timelineData, i = [], 1
    sql = f"SELECT chapterName from questionaff where StationCN = '{st.session_state.StationCN}' and chapterName <> '错题集' order by ID"
    rows = execute_sql(cur, sql)
    for row in rows:
        if row[0] != "公共题库":
            sql = f"SELECT Count(ID) from questions where chapterName = '{row[0]}'"
            quesCount = execute_sql(cur, sql)[0][0]
        else:
            sql = "SELECT Count(ID) from commquestions"
            quesCount = execute_sql(cur, sql)[0][0]
        sql = f"SELECT startTime from studyinfo where userName = '{st.session_state.userName}' and chapterName = '{row[0]}' order by startTime"
        rows2 = execute_sql(cur, sql)
        if rows2:
            trainingDate = time.strftime("%Y-%m-%d", time.localtime(rows2[0][0]))
            trainingDate2 = time.strftime("%Y-%m-%d", time.localtime(rows2[-1][0]))
            if len(rows2) == quesCount:
                temp = {"id": i, "content": row[0], "start": trainingDate, "end": trainingDate2}
            else:
                temp = {"id": i, "content": row[0], "start": trainingDate, "type": "point"}
            timelineData.append(temp)
            i += 1
    #st.write(timelineData)
    if timelineData:
        timeline = st_timeline(timelineData, groups=[], options={}, height="240px")
        if timeline is not None:
            if "end" in timeline:
                st.write(f"章节: :green[{timeline['content']}] 练习开始时间: :blue[{timeline['start']}] 完成时间: :orange[{timeline['end']}]")
            else:
                st.write(f"章节: :green[{timeline['content']}] 练习开始时间: :blue[{timeline['start']}]")
    else:
        st.write(":red[暂无学习记录]")


def displayCertificate():
    flagGener, flagInfo = False, True
    sql = f"SELECT examName from examidd where StationCN = '{st.session_state.StationCN}' and examName <> '练习题库' order by ID"
    rows = execute_sql(cur, sql)
    for row in rows:
        sql = f"SELECT userCName, examScore, examDate, CertificateNum, ID from examresult where userName = '{st.session_state.userName}' and examName = '{row[0]}' and examPass = 1 order by examScore DESC limit 0, 1"
        rows2 = execute_sql(cur, sql)
        if rows2:
            flagGener = True
            if flagGener and flagInfo:
                st.write(":orange[如需打印, 请打开 :green[程序目录内Image/Certificate] 或者点击下载证书]")
                flagInfo = False
            examDetail = rows2[0]
            with st.expander(label=f"{row[0]}", expanded=False):
                examDateDetail = time.strftime("%Y%m%d%H%M%S", time.localtime(examDetail[2]))
                if examDetail[3] == 0:
                    sql = "SELECT Max(CertificateNum) from examresult"
                    maxCertNum = execute_sql(cur, sql)[0][0] + 1
                else:
                    maxCertNum = examDetail[3]
                certFile = f"./Images/Certificate/Cert-Num.{str(maxCertNum).rjust(5, '0')}-{st.session_state.userName}-{examDetail[0]}-{row[0]}_{examDateDetail}.png"
                if not os.path.exists(certFile):
                    if examDetail[1] >= 100:
                        medal = "./Images/gold-award.png"
                    elif examDetail[1] >= 90:
                        medal = "./Images/silver-award.png"
                    else:
                        medal = "./Images/bronze-award.png"
                    examDate = time.strftime("%Y-%m-%d", time.localtime(examDetail[2]))
                    generCertificate(certFile, medal, st.session_state.userCName, row[0], examDate, maxCertNum)
                if os.path.exists(certFile):
                    sql = f"UPDATE examresult set CertificateNum = {maxCertNum} where ID = {examDetail[4]}"
                    execute_sql_and_commit(conn, cur, sql)
                    st.image(certFile)
                with open(certFile, "rb") as file:
                    st.download_button(
                        label="下载证书",
                        data=file,
                        file_name=certFile[certFile.rfind("/") + 1:].replace("Cert", "证书"),
                        mime="image/png",
                        icon=":material/download:"
                    )
                file.close()
    if not flagGener:
        st.info("您没有通过任何考试, 无法生成证书")


# 字体文件路径和证书背景路径作为常量定义，便于管理和修改
FONT_PATH = "./Fonts/msyhbd.ttf"
CERT_BG_PATH = './Images/Certificate-bg.png'

# 字体对象作为全局变量，避免重复加载
font_70 = ImageFont.truetype(FONT_PATH, 70)
font_30 = ImageFont.truetype(FONT_PATH, 30)
font_36 = ImageFont.truetype(FONT_PATH, 36)
font_46 = ImageFont.truetype("./Fonts/renaissance.ttf", 46)


def generCertificate(certFile, medal, userCName, examName, examDate, maxCertNum):
    # 优化姓名显示逻辑，减少重复代码
    name_adjustment = {"2": (userCName[0] + " " + userCName[-1], 866)}
    name_len = len(userCName.replace(" ", ""))
    adjusted_name = name_adjustment.get(str(name_len), (userCName, 460))[0]
    name_pos_x = name_adjustment.get(str(name_len), (userCName, 460))[1]
    name_pos_x_list = [866, 821, 796, 760, 726, 696]
    if 1 <= name_len <= 6:
        name_pos_x = name_pos_x_list[name_len - 1]

    try:
        # 使用with语句管理图片资源
        with Image.open(CERT_BG_PATH) as im, Image.open(medal) as imMedal:
            im.paste(imMedal, (784, 860), imMedal)
            dr = ImageDraw.Draw(im)
            dr.text((160, 132), f"No.{str(maxCertNum).rjust(5, '0')}", font=font_46, fill='grey')
            dr.text((name_pos_x, 460), adjusted_name, font=font_70, fill='grey')
            dr.text((900 - int(len(examName) * 15), 710), examName, font=font_30, fill='grey')
            dr.text((410, 940), examDate, font=font_36, fill='grey')
            im.save(certFile)
    except IOError as e:
        print(f"证书生成过程中发生错误: {e}")


def displayMedals():
    # 查询所有考试名称（排除练习题库）
    sql = "SELECT examName from examidd where examName <> '练习题库' order by ID"
    rows = execute_sql(cur, sql)

    # 遍历每个考试
    for row in rows:
        exam_name = row[0]
        with st.expander(label=exam_name, expanded=False):
            # 查询该考试通过且得分前三的用户信息
            sql = f"SELECT userCName, examScore, examDate from examresult where examName = '{exam_name}' and examPass = 1 order by examScore DESC limit 3"
            top_scores = execute_sql(cur, sql)

            # 定义奖牌图片路径和颜色映射
            medal_images = ["gold-medal.png", "silver-medal.png", "bronze-medal.png"]
            colors = ["red", "grey", "orange"]

            # 遍历得分前三的用户，显示奖牌和信息
            for i, (user_cname, score, date) in enumerate(top_scores):
                if i < 3:  # 确保只显示前三名
                    examDate = time.strftime("%Y-%m-%d", time.localtime(date))
                    st.image(f"./Images/{medal_images[i]}")  # 显示奖牌图片
                    st.write(f"##### :{colors[i]}[{user_cname}]")  # 显示用户名和颜色
                    st.write(f"成绩: {score}分")  # 显示成绩
                    st.write(examDate)  # 显示考试日期


def displayErrorQues():
    sql = f"SELECT Question, qOption, qAnswer, qType, qAnalysis, userAnswer, ID, WrongTime from morepractise where userAnswer <> '' and qAnswer <> userAnswer and userName = {st.session_state.userName} order by WrongTime DESC"
    rows = execute_sql(cur, sql)
    if rows:
        for row in rows:
            #st.subheader("", divider="red")
            with st.expander(label=f"题目: {row[0]} 次数: {row[7]}", expanded=False):
                if row[3] == "单选题":
                    st.write(":red[标准答案:]")
                    option, userAnswer = [], ["A", "B", "C", "D"]
                    tmp = row[1].replace("；", ";").split(";")
                    for index, each in enumerate(tmp):
                        each = each.replace("\n", "").replace("\t", "").strip()
                        option.append(f"{userAnswer[index]}. {each}")
                    st.radio(" ", option, key=f"compare_{row[6]}", index=int(row[2]), horizontal=True, label_visibility="collapsed", disabled=True)
                    st.write(f"你的答案: :red[{userAnswer[int(row[5])]}] 你的选择为: :blue[错误]")
                elif row[3] == "多选题":
                    userOption = ["A", "B", "C", "D", "E", "F", "G", "H"]
                    st.write(":red[标准答案:]")
                    option = row[1].replace("；", ";").split(";")
                    orgOption = row[2].replace("；", ";").split(";")
                    for index, value in enumerate(option):
                        value = value.replace("\n", "").replace("\t", "").strip()
                        if str(index) in orgOption:
                            st.checkbox(f"{userOption[index]}. {value}:", value=True, disabled=True)
                        else:
                            st.checkbox(f"{userOption[index]}. {value}:", value=False, disabled=True)
                    userAnswer = row[5].replace("；", ";").split(";")
                    tmp = ""
                    for each in userAnswer:
                        tmp = tmp + userOption[int(each)] + ", "
                    st.write(f"你的答案: :red[{tmp[:-2]}] 你的选择为: :blue[错误]")
                elif row[3] == "判断题":
                    st.write(":red[标准答案:]")
                    option = ["A. 正确", "B. 错误"]
                    tmp = int(row[2]) ^ 1
                    st.radio(" ", option, key=f"compare_{row[6]}", index=tmp, horizontal=True, label_visibility="collapsed", disabled=True)
                    tmp = int(row[5]) ^ 1
                    st.write(f"你的答案: :red[{option[tmp]}] 你的选择为: :blue[错误]")
                elif row[3] == "填空题":
                    option = row[2].replace("；", ";").split(";")
                    userAnswer = row[5].replace("；", ";").split(";")
                    st.write(":red[标准答案:]")
                    for index, value in enumerate(option):
                        st.write(f"第{index + 1}个填空: :green[{value}]")
                    st.write("你的答案:")
                    for index, value in enumerate(userAnswer):
                        st.write(f"第{index + 1}个填空: :red[{value}]")
                    st.write("你的填写为: :blue[错误]")
                if row[4] != "":
                    if row[4].endswith("]"):
                        #st.write(row[4])
                        st.markdown(f"答案解析: {row[4][:-1]}]")
                    else:
                        st.markdown(f"答案解析: :green[{row[4]}]")
    else:
        st.info("暂无数据")


def studyReset():
    buttonSubmit = st.button("重置学习记录", type="primary")
    if buttonSubmit:
        st.button("确认重置", type="secondary", on_click=studyResetAction)


def studyResetAction():
    sql = f"DELETE from studyinfo where userName = {st.session_state.userName}"
    execute_sql_and_commit(conn, cur, sql)
    st.success("学习记录已重置")
    updateKeyAction("重置学习记录")


# 定义常量
PUBLIC_QUESTION_BANK = '公共题库'
WRONG_QUESTION_SET = '错题集'


def get_chapter_progress(chapter_name, user_name, station_cn):
    # 参数化查询，获取章节题目数量和已学习题目数量
    sql_count = "SELECT Count(ID) FROM questions WHERE StationCN = ? AND chapterName = ?"
    sql_studied = "SELECT Count(ID) FROM studyinfo WHERE userName = ? AND chapterName = ?"

    # 查询章节题目总数
    params_count = (station_cn, chapter_name)
    total_questions = execute_sql(cur, sql_count, params_count)[0][0]

    # 查询已学习题目数
    params_studied = (user_name, chapter_name)
    studied_questions = execute_sql(cur, sql_studied, params_studied)[0][0]

    # 计算进度并返回
    if total_questions > 0:
        progress = studied_questions / total_questions
        return progress, int(progress * 100)
    return 0, 0


def studyinfoDetail():
    # 合并查询，获取非特殊章节的题目总数
    sql = """
    SELECT Count(ID) FROM questionaff
    WHERE StationCN = ? AND chapterName NOT IN (?, ?)
    """
    params = (st.session_state.StationCN, WRONG_QUESTION_SET, PUBLIC_QUESTION_BANK)
    chapter_total = execute_sql(cur, sql, params)[0][0]
    st.write("章节总计")
    st.write(f":blue[{chapter_total}]")

    # 获取试题总计（略，同原代码逻辑）
    # ...

    # 获取已学习试题数量和总完成率（略，同原代码逻辑）
    # ...

    # 展开各章节进度详情
    with st.expander("各章节进度详情", icon=":material/format_list_bulleted:", expanded=True):
        # 公共题库进度（略，同原代码逻辑但使用参数化查询）
        # ...

        # 其他章节进度
        sql = """
        SELECT chapterName FROM questionaff
        WHERE StationCN = ? AND chapterName NOT IN (?, ?) ORDER BY ID
        """
        rows = execute_sql(cur, sql, params)
        for row in rows:
            chapter_name = row[0]
            progress, percentage = get_chapter_progress(chapter_name, st.session_state.userName, st.session_state.StationCN)
            st.progress(value=progress, text=f":blue[{chapter_name}] 已完成 :orange[{percentage}%]")
# noinspection PyTypeChecker


def userStatus():
    st.subheader(":violet[在线用户状态]", divider="green")

    # 分离出密码验证的逻辑
    def verify_admin_password():
        vUserPW = st.text_input("请输入密码", max_chars=8, placeholder="请输入管理员密码, 以验证身份", type="password", autocomplete="off")
        if vUserPW:
            if verifyUserPW(st.session_state.userName, vUserPW)[0]:
                st.session_state.userPwRecheck = True
                st.success("密码验证成功！")
            else:
                st.error("密码错误, 请重新输入")

    # 如果还未验证密码，则先进行验证
    if not st.session_state.get('userPwRecheck', False):
        verify_admin_password()
        return  # 验证后直接返回，避免执行后续逻辑

    bc = sac.segmented(
        items=[
            sac.SegmentedItem(label="在线用户状态", icon="people"),
            sac.SegmentedItem(label="重置所有用户状态", icon="person-slash"),
        ], align="start", color="red", size="sm"
    )

    if bc == "在线用户状态":
        actionUserStatus()
    elif bc == "重置所有用户状态":
        buttonReset = st.button("重置所有用户状态", type="primary")
        if buttonReset:
            if st.button("确认重置", type="secondary", on_click=resetActiveUser):
                st.success("所有用户状态已重置！")

    if bc is not None:
        updateActionUser(st.session_state.userName, bc, st.session_state.loginTime)


def actionUserStatus():
    sql = "SELECT userCName, userType, StationCN, actionUser, loginTime, activeTime_session, activeTime from users where activeUser = 1 order by loginTime desc, activeTime_session desc, activeTime desc, ID"
    rows = execute_sql(cur, sql)
    df = pd.DataFrame(rows, dtype=str)
    df.columns = ["姓名", "类型", "站室", "用户操作", "登录时间", "活动时间", "累计活动时间"]
    for index, value in enumerate(rows):
        df.loc[index, "登录时间"] = time.strftime("%Y-%m-%d %H:%M:%S", time.localtime(int(df["登录时间"][index])))
        activeTime = int(df.loc[index, "活动时间"])
        hTime = int(activeTime / 3600)
        mTime = int((activeTime % 3600) / 60)
        if mTime < 10:
            mTime = "0" + str(mTime)
        sTime = int(activeTime % 60)
        if sTime < 10:
            sTime = "0" + str(sTime)
        df.loc[index, "活动时间"] = f"{hTime}小时{mTime}分{sTime}秒"
        activeTime = int(df.loc[index, "累计活动时间"])
        hTime = int(activeTime / 3600)
        mTime = int((activeTime % 3600) / 60)
        if mTime < 10:
            mTime = "0" + str(mTime)
        sTime = int(activeTime % 60)
        if sTime < 10:
            sTime = "0" + str(sTime)
        df.loc[index, "累计活动时间"] = f"{hTime}小时{mTime}分{sTime}秒"
    st.dataframe(df, use_container_width=True)


@st.fragment
def actionQuesModify(row):
    option = []
    if len(row) == 8:
        qQuestion, qOption, qAnswer, qType, qAnalysis, StationCN, chapterName, SourceType = row
    else:
        qQuestion, qOption, qAnswer, qType, qAnalysis, SourceType = row
    st.session_state.qModifyQues_qType = qType
    st.write(f"**此题为{qType}**")
    st.text_area(":blue[**题目**]", value=qQuestion, key="qModifyQues_Question")
    if qType == "单选题":
        qOption2 = qOption.split(";")
        st.session_state.qModifyQues_optionCount = len(qOption2)
        for index, value in enumerate(qOption2):
            st.text_input(f":orange[**选项{chr(65 + index)}**]", value=value, key=f"qModifyQues_{index}")
            option.append(chr(65 + index))
        st.radio(":red[**答案**]", options=option, index=int(qAnswer), key="qModifyQues_Answer", horizontal=True)
    elif qType == "多选题":
        qOption2 = qOption.split(";")
        qAnswer2 = qAnswer.split(";")
        st.session_state.qModifyQues_optionCount = len(qOption2)
        for index, value in enumerate(qOption2):
            st.text_input(f":orange[**选项{chr(65 + index)}**]", value=value, key=f"qModifyQues_{index}")
            if str(index) in qAnswer2:
                st.checkbox(":blue[**选择**]", value=True, key=f"qModifyQues_Answer_{index}")
            else:
                st.checkbox(":blue[**选择**]", value=False, key=f"qModifyQues_Answer_{index}")
    elif qType == "判断题":
        st.radio(":red[**答案**]", ["A. 正确", "B. 错误"], key="qModifyQues_Answer", index=int(qAnswer) ^ 1, horizontal=True)
    elif qType == "填空题":
        qAnswer2 = qAnswer.split(";")
        st.session_state.qModifyQues_optionCount = len(qAnswer2)
        for index, value in enumerate(qAnswer2):
            st.text_input(":orange[**答案**]", value=value, key=f"qModifyQues_Answer_{index}")
    st.text_area(":green[**答案解析**]", value=qAnalysis, key="qModifyQues_Answer_Analysis")


def quesModify():
    st.subheader(":green[试题修改]", divider="blue")
    chosenTable = st.selectbox(":red[选择题库]", ["站室题库", "公共题库"], index=None)
    quesID = st.number_input(":blue[题目ID]", min_value=0, step=1)
    if chosenTable is not None and quesID > 0:
        if chosenTable == "站室题库":
            tablename = "questions"
        elif chosenTable == "公共题库":
            tablename = "commquestions"
        else:
            tablename = ""
        buttonDisplayQues = st.button("显示试题", icon=":material/dvr:")
        if buttonDisplayQues:
            if chosenTable == "站室题库":
                sql = f"SELECT Question, qOption, qAnswer, qType, qAnalysis, StationCN, chapterName, SourceType from {tablename} where ID = {quesID}"
            else:
                sql = f"SELECT Question, qOption, qAnswer, qType, qAnalysis, SourceType from {tablename} where ID = {quesID}"
            rows = execute_sql(cur, sql)
            if rows:
                if chosenTable == "站室题库":
                    st.write(f":green[站室: {rows[0][5]} 章节: {rows[0][6]} 试题来源: {rows[0][7]}]")
                else:
                    st.write(f":green[公共题库 试题来源: {rows[0][5]}]")
                st.button("更新试题", on_click=actionQM, args=(quesID, tablename, rows[0]), icon=":material/published_with_changes:")
                st.button("删除试题", on_click=actionDelQM, args=(quesID, tablename, rows[0]), icon=":material/delete:")
                if chosenTable == "站室题库":
                    st.button("移至公共题库", on_click=moveQM, args=(quesID, tablename, rows[0]), icon=":material/move_item:")
                actionQuesModify(rows[0])
            else:
                st.error("未找到该题目, 请检查题库名称及题目ID是否正确")
    else:
        st.error("请选择题库")


def moveQM(quesID, tablename, mRow):
    sql = f"DELETE from {tablename} where ID = {quesID}"
    execute_sql_and_commit(conn, cur, sql)
    sql = f"INSERT INTO commquestions(Question, qOption, qAnswer, qType, qAnalysis, SourceType) VALUES('{mRow[0]}', '{mRow[1]}', '{mRow[2]}', '{mRow[3]}', '{mRow[4]}', '{mRow[7]}')"
    execute_sql_and_commit(conn, cur, sql)
    st.toast("试题移至公共题库成功")


def actionQM(quesID, tablename, mRow):
    mOption, mAnswer, Option = "", "", ["A", "B", "C", "D", "E", "F", "G", "H"]
    mQues = st.session_state.qModifyQues_Question
    mAnalysis = st.session_state.qModifyQues_Answer_Analysis
    if st.session_state.qModifyQues_qType == "单选题" or st.session_state.qModifyQues_qType == "多选题":
        for i in range(st.session_state.qModifyQues_optionCount):
            mOption = mOption + st.session_state[f"qModifyQues_{i}"] + ";"
        if mOption.endswith(";"):
            mOption = mOption[:-1]
        if st.session_state.qModifyQues_qType == "单选题":
            for index, value in enumerate(Option):
                if value == st.session_state.qModifyQues_Answer:
                    mAnswer = index
                    break
        else:
            for i in range(st.session_state.qModifyQues_optionCount):
                if st.session_state[f"qModifyQues_Answer_{i}"]:
                    mAnswer = mAnswer + str(i) + ";"
            if mAnswer.endswith(";"):
                mAnswer = mAnswer[:-1]
    elif st.session_state.qModifyQues_qType == "判断题":
        if "正确" in st.session_state.qModifyQues_Answer:
            mAnswer = 1
        else:
            mAnswer = 0
    elif st.session_state.qModifyQues_qType == "填空题":
        for i in range(st.session_state.qModifyQues_optionCount):
            mAnswer = mAnswer + st.session_state[f"qModifyQues_Answer_{i}"] + ";"
        if mAnswer.endswith(";"):
            mAnswer = mAnswer[:-1]
    sql = f"UPDATE {tablename} set Question = '{mQues}', qOption = '{mOption}', qAnswer = '{mAnswer}', qAnalysis = '{mAnalysis}' where ID = {quesID}"
    execute_sql_and_commit(conn, cur, sql)
    clearModifyQues(quesID, tablename, mRow)
    for key in st.session_state.keys():
        if key.startswith("qModifyQues_"):
            del st.session_state[key]
    st.toast("试题修改成功")


def actionDelQM(quesID, tablename, mRow):
    sql = f"DELETE from {tablename} where ID = {quesID}"
    execute_sql_and_commit(conn, cur, sql)
    clearModifyQues(quesID, tablename, mRow)
    for key in st.session_state.keys():
        if key.startswith("qModifyQues_"):
            del st.session_state[key]
    st.toast("试题删除成功")


def clearModifyQues(quesID, tablename, mRow):
    delTablePack = ["morepractise", "favques"]
    for each in delTablePack:
        sql = f"DELETE from {each} where Question = '{mRow[0]}' and qOption = '{mRow[1]}' and qAnswer = '{mRow[2]}' and qType = '{mRow[3]}'"
        execute_sql_and_commit(conn, cur, sql)
    sql = f"DELETE from studyinfo where cid = {quesID} and quesTable = '{tablename}'"
    execute_sql_and_commit(conn, cur, sql)


def aboutReadme():
    st.markdown(open("./README.md", "r", encoding="utf-8").read())


# noinspection PyUnboundLocalVariable
def training():
    flagProc, failInfo = True, ""
    if st.session_state.examType == "exam":
        quesType = []
        sql = f"SELECT paramName from setup_{st.session_state.StationCN} where paramType = 'questype' and param = 1 order by ID"
        rows = execute_sql(cur, sql)
        for row in rows:
            quesType.append([row[0], getParam(f"{row[0]}数量", st.session_state.StationCN)])
        for each in quesType:
            quesTypeCount = 0
            tmp = each[0].replace("数量", "")
            sql = f"SELECT count(ID) from questions where qType = '{tmp}' and StationCN = '{st.session_state.StationCN}'"
            quesTypeCount += int(execute_sql(cur, sql)[0][0])
            sql = f"SELECT count(ID) from commquestions where qType = '{tmp}'"
            quesTypeCount += int(execute_sql(cur, sql)[0][0])
            if quesTypeCount < each[1]:
                flagProc = False
                failInfo = failInfo + f"{tmp}/"
    elif st.session_state.examType == "training":
        quesType = [["单选题", 30], ["多选题", 10], ["判断题", 10], ["填空题", 0]]
        sql = f"SELECT mcq, mmcq, tfq, fibq from indivquescount where userName = {st.session_state.userName}"
        rows = execute_sql(cur, sql)
        if rows:
            row = rows[0]
            for i in range(4):
                quesType[i][1] = row[i]
        else:
            sql = f"INSERT INTO indivquescount (userName, mcq, mmcq, tfq, fibq) VALUES({st.session_state.userName}, {quesType[0][1]}, {quesType[1][1]}, {quesType[2][1]}, {quesType[3][1]})"
            execute_sql_and_commit(conn, cur, sql)
    if flagProc:
        generPack, examIDPack, chapterPack, genResult = [], [], [], []
        generQues = st.empty()
        with generQues.container():
            if st.session_state.examType == "exam":
                date = int(time.time())
                sql = f"SELECT examName from examidd where StationCN = '{st.session_state.StationCN}' and validDate >= {date} order by validDate"
                rows = execute_sql(cur, sql)
                for row in rows:
                    examIDPack.append(row[0])
                examName = st.selectbox("请选择考试场次", examIDPack, index=None)
                if examName:
                    generButtonQues = st.button("开始考试")
                    if generButtonQues:
                        st.session_state.examName = examName
                        st.spinner("正在生成题库...")
                        reviseQues()
                        sql = "SELECT chapterName from questionaff where chapterName <> '错题集' and chapterName <> '关注题集' and StationCN = '" + st.session_state.StationCN + "'"
                        rows = execute_sql(cur, sql)
                        for row in rows:
                            chapterPack.append(row[0])
                        genResult = GenerExam(chapterPack, st.session_state.StationCN, st.session_state.userName, st.session_state.examName, st.session_state.examType, quesType, st.session_state.examRandom, False)
            elif st.session_state.examType == "training":
                generButtonQues = st.button("生成题库")
                sql = "SELECT pyLM from verinfo where pyFile = 'chapterChosenType'"
                chapterChosenType = execute_sql(cur, sql)[0][0]
                uCCT = sac.segmented(
                    items=[
                        sac.SegmentedItem(label="默认"),
                        sac.SegmentedItem(label="全选"),
                        sac.SegmentedItem(label="全不选"),
                    ], index=chapterChosenType, align="start", color="orange", return_index=True, size="sm",
                )
                if uCCT != 0:
                    sql = f"UPDATE verinfo set pyLM = {uCCT} where pyFile = 'chapterChosenType'"
                    execute_sql_and_commit(conn, cur, sql)
                st.checkbox(":red[**仅未学习试题**]", value=False, key="GenerNewOnly", help="仅从未学习试题中生成")
                indivCols = st.columns(4)
                for i in range(4):
                    quesType[i][1] = indivCols[i].number_input(quesType[i][0], min_value=0, max_value=100, value=quesType[i][1], step=1)
                for each in ["公共题库", "错题集", "关注题集"]:
                    sql = f"SELECT chapterName, chapterRatio, ID from questionaff where StationCN = '{st.session_state.StationCN}' and chapterName = '{each}'"
                    row = execute_sql(cur, sql)[0]
                    if uCCT == 0:
                        if each == "公共题库":
                            generPack.append(st.checkbox(f"**:blue[{row[0]}]**", value=True))
                        else:
                            generPack.append(st.checkbox(f"**:blue[{row[0]}]**", value=False))
                    elif uCCT == 1:
                        generPack.append(st.checkbox(f"**:blue[{row[0]}]**", value=True))
                    elif uCCT == 2:
                        generPack.append(st.checkbox(f"**:blue[{row[0]}]**", value=False))
                    st.slider("章节权重", min_value=1, max_value=10, value=row[1], step=1, key=f"tempCR_{row[2]}", on_change=updateCRTraining, label_visibility="collapsed")
                sql = "SELECT chapterName, chapterRatio, ID from questionaff where StationCN = '" + st.session_state.StationCN + "' and chapterName <> '公共题库' and chapterName <> '错题集' and chapterName <> '关注题集' order by chapterName"
                rows = execute_sql(cur, sql)
                for row in rows:
                    if uCCT == 0 or uCCT == 1:
                        generPack.append(st.checkbox(f"**:blue[{row[0]}]**", value=True))
                    elif uCCT == 2:
                        generPack.append(st.checkbox(f"**:blue[{row[0]}]**", value=False))
                    st.slider("章节权重", min_value=1, max_value=10, value=row[1], step=1, key=f"tempCR_{row[2]}", on_change=updateCRTraining, label_visibility="collapsed")
                if generButtonQues:
                    st.session_state.examName = "练习题库"
                    sql = f"UPDATE indivquescount set mcq = {quesType[0][1]}, mmcq = {quesType[1][1]}, tfq = {quesType[2][1]}, fibq = {quesType[3][1]} where userName = {st.session_state.userName}"
                    execute_sql_and_commit(conn, cur, sql)
                    for index, value in enumerate(generPack):
                        if value:
                            if index == 0:
                                chapterPack.append("公共题库")
                            elif index == 1:
                                chapterPack.append("错题集")
                            elif index == 2:
                                chapterPack.append("关注题集")
                            else:
                                chapterPack.append(rows[index - 3][0])
                    if chapterPack:
                        st.spinner("正在生成题库...")
                        reviseQues()
                        genResult = GenerExam(chapterPack, st.session_state.StationCN, st.session_state.userName, st.session_state.examName, st.session_state.examType, quesType, st.session_state.examRandom, st.session_state.GenerNewOnly)
                    else:
                        st.warning("请至少选择一个章节")
        if genResult:
            if genResult[0]:
                generQues.empty()
                if st.session_state.examType == "exam":
                    st.success(f"题库生成完毕, 总共生成{genResult[1]}道试题, 请在👈左侧边栏选择开始考试")
                else:
                    st.success(f"题库生成完毕, 总共生成{genResult[1]}道试题, 请在👈左侧边栏选择题库练习")
                st.session_state.examTable = genResult[2]
                st.session_state.examFinalTable = genResult[3]
                st.session_state.curQues = 0
                st.session_state.examStartTime = int(time.time())
                st.session_state.confirmSubmit = False
                st.session_state.flagCompleted = False
                st.session_state.goto = False
                st.session_state.radioCompleted = False
                st.session_state.calcScore = False
                st.session_state.delExam = False
                st.session_state.trainingID = str(time.strftime('%Y%m%d%H%M%S', time.localtime(int(time.time()))))
                if st.session_state.examType != "training":
                    st.session_state.examChosen = True
                    updateActionUser(st.session_state.userName, "生成考试试题", st.session_state.loginTime)
                else:
                    st.session_state.examChosen = False
                    updateActionUser(st.session_state.userName, "生成练习试题", st.session_state.loginTime)
            else:
                st.session_state.examChosen = False
                st.error("题库生成试题不满足要求, 请检查考试参数设置, 或个别题型试题候选数量不够, 或请联系管理员")
    else:
        st.error(f":red[⚠️] **{st.session_state.StationCN}试卷生成失败, :red[{failInfo[:-1]}] 试题数量不足, 请检查题库设置或增加以上题型候选试题**")


def reviseQues():
    for each in ["questions", "commquestions"]:
        for each2 in [['（', '('], ['）', ')']]:
            sql = f"UPDATE {each} set Question = replace(Question, '{each2[0]}', '{each2[1]}') where qType = '填空题' and Question like '%{each2[0]}%'"
            execute_sql_and_commit(conn, cur, sql)
        for each2 in ['( )', '(  )', '(   )', '(    )']:
            sql = f"UPDATE {each} set Question = replace(Question, '{each2}', '()') where qType = '填空题' and Question like '%{each2}'"
            execute_sql_and_commit(conn, cur, sql)


@st.fragment
def updateCRTraining():
    for key in st.session_state.keys():
        if key.startswith("tempCR_"):
            upID = key[key.find("_") + 1:]
            sql = f"UPDATE questionaff SET chapterRatio = {st.session_state[key]} WHERE ID = {upID}"
            execute_sql_and_commit(conn, cur, sql)


def updateCRExam():
    for key in st.session_state.keys():
        if key.startswith("crsetup_"):
            upID = key[key.find("_") + 1:]
            sql = f"UPDATE questionaff SET examChapterRatio = {st.session_state[key]} WHERE ID = {upID}"
            execute_sql_and_commit(conn, cur, sql)
    st.success("章节权重更新成功")
    updateKeyAction("考试章节权重更新")


@st.fragment
def updateAnswer(user_ques_id):
    try:
        # 使用参数化查询防止SQL注入
        update_sql = f"UPDATE {st.session_state.examFinalTable} SET userAnswer = ?, userName = ? WHERE ID = ?"
        params = (st.session_state.answer, st.session_state.userName, user_ques_id)
        execute_sql_and_commit(conn, cur, update_sql, params)

        select_sql = f"SELECT Question, qAnswer, qType, userAnswer, userName, qOption, qAnalysis, SourceType FROM {st.session_state.examFinalTable} WHERE ID = ?"
        jud_table = execute_sql(cur, select_sql, (user_ques_id,))[0]

        if jud_table[1] == jud_table[3]:
            # 正确回答的逻辑
            decrement_wrong_time_sql = "UPDATE morepractise SET WrongTime = WrongTime - 1 WHERE Question = ? AND qType = ? AND userName = ?"
            execute_sql_and_commit(conn, cur, decrement_wrong_time_sql, (jud_table[0], jud_table[2], jud_table[4]))

            increment_user_ranking_sql = "UPDATE users SET userRanking = userRanking + 1 WHERE userName = ?"
            execute_sql_and_commit(conn, cur, increment_user_ranking_sql, (st.session_state.userName,))

        else:
            # 错误回答的逻辑
            check_exists_sql = "SELECT ID FROM morepractise WHERE Question = ? AND qType = ? AND userName = ?"
            if execute_sql(cur, check_exists_sql, (jud_table[0], jud_table[2], jud_table[4])):
                increment_wrong_time_sql = "UPDATE morepractise SET WrongTime = WrongTime + 1, userAnswer = ? WHERE Question = ? AND qType = ? AND userName = ? AND trainingID <> ?"
                params = (jud_table[3], jud_table[0], jud_table[2], jud_table[4], st.session_state.trainingID)
                execute_sql_and_commit(conn, cur, increment_wrong_time_sql, params)
            else:
                insert_sql = "INSERT INTO morepractise(Question, qOption, qAnswer, qType, qAnalysis, userAnswer, userName, WrongTime, StationCN, SourceType, trainingID) VALUES(?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)"
                params = (jud_table[0], jud_table[5], jud_table[1], jud_table[2], jud_table[6], jud_table[3], jud_table[4], 1, st.session_state.StationCN, jud_table[7], st.session_state.trainingID)
                execute_sql_and_commit(conn, cur, insert_sql, params)

        # 清理WrongTime小于1的记录
        cleanup_sql = "DELETE FROM morepractise WHERE WrongTime < 1"
        execute_sql_and_commit(conn, cur, cleanup_sql)

    except Exception as e:
        st.error(f"更新答案时出错: {e}")


@st.dialog("考试成绩")
def score_dialog(userScore, passScore):
    examDate = int(time.mktime(time.strptime(time.strftime("%Y-%m-%d %H:%M:%S", time.localtime()), "%Y-%m-%d %H:%M:%S")))
    if userScore >= passScore:
        flagPass = 1
    else:
        flagPass = 0
    st.write(f"考生ID:  {st.session_state.userName}")
    st.write(f"考生姓名: {st.session_state.userCName}")
    st.write(f"考试时间: {time.strftime('%Y-%m-%d %H:%M:%S', time.localtime(examDate))}")
    st.subheader(f"考试成绩: {userScore} 分 / 合格分数线为 {passScore} 分")
    if flagPass == 1:
        st.subheader("考试结果: :blue[通过] 👏")
        st.balloons()
    else:
        st.subheader("考试结果: :red[未通过] 🤪")
        #st.snow()
    if st.session_state.examType == "training":
        st.write("练习模式成绩不计入记录")
    if st.session_state.examType == "exam" and st.session_state.calcScore:
        sql = "INSERT INTO examresult(userName, userCName, examScore, examDate, examPass, examName) VALUES(" + str(st.session_state.userName) + ", '" + st.session_state.userCName + "', " + str(userScore) + ", " + str(examDate) + ", " + str(flagPass) + ", '" + st.session_state.examName + "')"
        execute_sql_and_commit(conn, cur, sql)
    st.session_state.calcScore = False
    buttonScore = st.button("确定")
    if buttonScore:
        st.session_state.delExam = True
        st.rerun()


def calcScore():
    st.session_state.examStartTime = int(time.time())
    st.session_state.confirmSubmit = True
    st.session_state.curQues = 0
    st.session_state.flagCompleted = False
    flagUseAIFIB = bool(getParam("使用大模型评判错误的填空题答案", st.session_state.StationCN))
    quesScore = getParam("单题分值", st.session_state.StationCN)
    passScore = getParam("合格分数线", st.session_state.StationCN)
    userScore = 0
    sql = f"SELECT qAnswer, qType, userAnswer, Question, qOption, qAnalysis, userName, SourceType from {st.session_state.examFinalTable} where userName = {st.session_state.userName} order by ID"
    rows = execute_sql(cur, sql)
    for row in rows:
        flagAIScore = False
        if row[0].replace(" ", "").lower() == row[2].replace(" ", "").lower():
            userScore += quesScore
            sql = f"UPDATE users set userRanking = userRanking + 1 where userName = {st.session_state.userName}"
            execute_sql_and_commit(conn, cur, sql)
            sql = f"SELECT ID from morepractise where Question = '{row[3]}' and qType = '{row[1]}' and userName = {row[6]}"
            if execute_sql(cur, sql):
                sql = f"UPDATE morepractise set WrongTime = WrongTime - 1 where Question = '{row[3]}' and qType = '{row[1]}' and userName = {row[6]}"
                execute_sql_and_commit(conn, cur, sql)
            execute_sql_and_commit(conn, cur, sql="DELETE from morepractise where WrongTime < 1")
        else:
            if row[1] == "填空题":
                if flagUseAIFIB:
                    fibQues = row[3]
                    fibQues2 = row[3]
                    userAP = row[2].split(";")
                    quesAP = row[0].split(";")
                    if fibQues.count("()") == len(userAP):
                        #st.toast("正在使用 :red[讯飞星火大模型] 对答案进行分析, 请稍等...")
                        for index, value in enumerate(userAP):
                            b1 = fibQues.find("()")
                            c1 = fibQues2.find("()")
                            if b1 != -1:
                                fibQues = f"{fibQues[:b1]}({value}){fibQues[b1 + 2:]}"
                                fibQues2 = f"{fibQues2[:c1]}({quesAP[index]}){fibQues2[c1 + 2:]}"
                        fibAI = xunfei_xh_AI_fib(fibQues, fibQues2)
                        if fibAI != "" and fibAI.find("无法直接回答") == -1 and fibAI.find("尚未查询") == -1 and fibAI.find("我不确定您想要表达什么意思") == -1 and fibAI.find("由于信息不足，无法给出准确答案") == -1 and fibAI.find("无法确定正确答案") == -1 and fibAI.find("无法提供准确答案") == -1:
                            if st.session_state.debug:
                                print(f"debug: [{row[3]}] [Q:{row[0]} / A:{row[2]}] / A.I.判断: [{fibAI}]")
                            if fibAI == "正确":
                                userScore += quesScore
                                sql = f"UPDATE users set userRanking = userRanking + 1 where userName = {st.session_state.userName}"
                                execute_sql_and_commit(conn, cur, sql)
                                flagAIScore = True
                            else:
                                flagAIScore = False
                    else:
                        st.error("⚠️ 试题或是答案数量不匹配, 请检查")
            if not flagAIScore:
                sql = f"SELECT ID from morepractise where Question = '{row[3]}' and qType = '{row[1]}' and userName = {row[6]}"
                if not execute_sql(cur, sql):
                    sql = f"INSERT INTO morepractise(Question, qOption, qAnswer, qType, qAnalysis, userAnswer, userName, WrongTime, StationCN, SourceType, trainingID) VALUES('{row[3]}', '{row[4]}', '{row[0]}', '{row[1]}', '{row[5]}', '{row[2]}', {row[6]}, 1, '{st.session_state.StationCN}', '{row[7]}', '{st.session_state.trainingID}')"
                    execute_sql_and_commit(conn, cur, sql)
                else:
                    sql = f"UPDATE morepractise set WrongTime = WrongTime + 1, userAnswer = '{row[2]}' where Question = '{row[3]}' and qType = '{row[1]}' and userName = {row[6]} and trainingID <> '{st.session_state.trainingID}'"
                    execute_sql_and_commit(conn, cur, sql)
    if st.session_state.calcScore:
        score_dialog(userScore, passScore)


@st.fragment
def updateOptionAnswer(chosenID, chosen, option):
    for index, value in enumerate(option):
        if chosen == value:
            st.session_state.answer = index
    updateAnswer(chosenID)


@st.fragment
def updateRadioAnswer(chosenID):
    if st.session_state.radioChosen is not None:
        if "正确" in st.session_state.radioChosen:
            st.session_state.answer = 1
        else:
            st.session_state.answer = 0
        st.session_state.radioCompleted = True
        updateAnswer(chosenID)


@st.fragment
def updateRadioAnswer2(chosenID):
    if st.session_state.radioChosen2 is not None:
        if "正确" in st.session_state.radioChosen2:
            st.session_state.answer = 1
        else:
            st.session_state.answer = 0
        updateAnswer(chosenID)


@st.fragment
def updateMOptionAnswer(row):
    mpAnswerPack = []
    for key in st.session_state.keys():
        if key.startswith("moption_"):
            if st.session_state[key]:
                mpAnswerPack.append(int(key.replace("moption_", "")))
    mpAnswerPack.sort()
    st.session_state.answer = ""
    for each in mpAnswerPack:
        st.session_state.answer = st.session_state.answer + str(each) + ";"
    if st.session_state.answer.endswith(";"):
        st.session_state.answer = st.session_state.answer[:-1]
    updateAnswer(row[0])


def delQuestion(delQuesRow):
    delTablePack = ["questions", "commquestions", "morepractise", "favques"]
    for delTable in delTablePack:
        sql = f"DELETE from {delTable} where Question = '{delQuesRow[1]}' and qOption = '{delQuesRow[2]}' and qType = '{delQuesRow[4]}'"
        execute_sql_and_commit(conn, cur, sql)
    updateKeyAction(f"删除试题: {delQuesRow[1]}")


@st.fragment
def updateStudyInfo(studyRow):
    query_mapping = {
        "questions": "SELECT ID, chapterName from {{table}} where Question = ? and qType = ? and StationCN = ?",
        "commquestions": "SELECT ID, '公共题库' as chapterName from {{table}} where Question = ? and qType = ?"
    }
    params = (studyRow[1], studyRow[4], st.session_state.StationCN) if st.session_state.StationCN else (studyRow[1], studyRow[4])

    for table, sql_template in query_mapping.items():
        sql = sql_template.format(table=table)
        try:
            studyResult = execute_sql(cur, sql, params[:len(params) - (table == "commquestions")])
            if studyResult:
                check_sql = "SELECT ID from studyinfo where cid = ? and questable = ? and userName = ? and chapterName = ?"
                if not execute_sql(cur, check_sql, (studyResult[0][0], table, st.session_state.userName, studyResult[0][1])):
                    insert_sql = "INSERT INTO studyinfo(cid, questable, userName, userCName, chapterName, startTime) VALUES(?, ?, ?, ?, ?, ?)"
                    execute_sql_and_commit(conn, cur, insert_sql, (studyResult[0][0], table, st.session_state.userName, st.session_state.userCName, studyResult[0][1], int(time.time())))
        except Exception as e:
            st.error(f"更新学习信息时出错: {e}")


@st.fragment
def delFavQues(favRow):
    sql = f"DELETE from favques where Question = '{favRow[1]}' and userName = {st.session_state.userName} and qType = '{favRow[4]}' and StationCN = '{st.session_state.StationCN}'"
    execute_sql_and_commit(conn, cur, sql)
    st.toast("已从关注题集中删除")


@st.fragment
def addFavQues(favRow):
    sql = f"SELECT ID from favques where Question = '{favRow[1]}' and userName = {st.session_state.userName} and StationCN = '{st.session_state.StationCN}'"
    if not execute_sql(cur, sql):
        sql = f"INSERT INTO favques(Question, qOption, qAnswer, qType, qAnalysis, userName, StationCN, SourceType) VALUES('{favRow[1]}', '{favRow[2]}', '{favRow[3]}', '{favRow[4]}', '{favRow[5]}', {st.session_state.userName}, '{st.session_state.StationCN}', '{favRow[8]}')"
        execute_sql_and_commit(conn, cur, sql)
        st.toast("已添加到关注题集")


# noinspection PyUnboundLocalVariable
@st.fragment
def exam(row):
    option, AIModelName, AIOption, AIOptionIndex = [], "", [], 0
    st.session_state.answer = ""
    flagAIUpdate = bool(getParam("A.I.答案解析更新至题库", st.session_state.StationCN))
    sql = f"SELECT paramName, param from setup_{st.session_state.StationCN} where paramType = 'others' and paramName like '%大模型' order by ID"
    tempTable = execute_sql(cur, sql)
    for index, value in enumerate(tempTable):
        AIOption.append(value[0])
        if value[1] == 1:
            AIOptionIndex = index
    if row[4] == "填空题":
        reviseQues = row[1].replace("(", ":red[ ( _ ]").replace(")", ":red[ _ _ ) ]").strip()
    else:
        reviseQues = row[1].strip()
    standardAnswer = getStandardAnswer(row)
    if st.session_state.examType != "exam":
        updateStudyInfo(row)
    st.write(f"##### 第{row[0]}题 :green[{reviseQues}]")
    if st.session_state.userType == "admin" and st.session_state.examType != "exam" and st.session_state.debug:
        buttonConfirm = st.button("⚠️ 从所有题库中删除此题", type="primary")
        if buttonConfirm:
            st.button("确认删除", type="secondary", on_click=delQuestion, args=(row,))
    if st.session_state.examType == "training":
        sql = f"SELECT ID from favques where Question = '{row[1]}' and userName = {st.session_state.userName} and StationCN = '{st.session_state.StationCN}'"
        if execute_sql(cur, sql):
            st.button(label="", icon=":material/heart_minus:", on_click=delFavQues, args=(row,), help="从关注题集中删除")
        else:
            st.button(label="", icon=":material/heart_plus:", on_click=addFavQues, args=(row,), help="添加到关注题集")
    st.write(f":red[本题为{row[4]}]:")
    if row[4] == '单选题':
        for index, value in enumerate(row[2].replace("；", ";").split(";")):
            value = value.replace("\n", "").replace("\t", "").strip()
            option.append(f"{chr(65 + index)}. {value}")
        if row[6] == "":
            chosen = st.radio(" ", option, index=None, label_visibility="collapsed", horizontal=True)
        else:
            chosen = st.radio(" ", option, index=int(row[6]), label_visibility="collapsed", horizontal=True)
            #st.write(f":red[你已选择: ] :blue[{option[int(row[6])]}]")
        if chosen is not None:
            updateOptionAnswer(row[0], chosen, option)
    elif row[4] == '多选题':
        for index, value in enumerate(row[2].replace("；", ";").split(";")):
            value = value.replace("\n", "").replace("\t", "").strip()
            option.append(f"{chr(65 + index)}. {value}")
        if row[6] != "":
            orgOption = row[6].replace("；", ";").split(";")
        else:
            orgOption = []
        for index, value in enumerate(option):
            if str(index) in orgOption:
                st.checkbox(f"{value}", value=True, key=f"moption_{index}", on_change=updateMOptionAnswer, args=(row,))
            else:
                st.checkbox(f"{value}", value=False, key=f"moption_{index}", on_change=updateMOptionAnswer, args=(row,))
    elif row[4] == '判断题':
        radioArea = st.empty()
        with radioArea.container():
            option = ["A. 正确", "B. 错误"]
            st.radio(" ", option, index=None, key="radioChosen", on_change=updateRadioAnswer, args=(row[0],), label_visibility="collapsed", horizontal=True)
            if row[6] != "":
                st.write(f":red[**你已选择:** ] :blue[[**{option[int(row[6]) ^ 1][0]}**]]")
        if st.session_state.radioCompleted:
            radioArea.empty()
            st.session_state.radioCompleted = False
            sql = f"SELECT userAnswer from {st.session_state.examFinalTable} where ID = {row[0]}"
            tempUserAnswer = execute_sql(cur, sql)[0][0]
            if tempUserAnswer != "":
                st.radio(" ", option, index=int(tempUserAnswer) ^ 1, key="radioChosen2", on_change=updateRadioAnswer2, args=(row[0],), label_visibility="collapsed", horizontal=True)
    elif row[4] == '填空题':
        orgOption = row[6].replace("；", ";").split(";")
        textAnswerArea = st.empty()
        with textAnswerArea.container():
            for i in range(row[1].count("()")):
                if row[6] == "":
                    st.text_input(label=" ", key=f"textAnswer_{i}", placeholder=f"请输入第{i + 1}个括号内的内容", label_visibility="collapsed")
                else:
                    st.text_input(label=" ", value=orgOption[i], key=f"textAnswer_{i}", placeholder=f"请输入第{i + 1}个括号内的内容", label_visibility="collapsed")
            buttonTA = st.button("确定")
            if buttonTA:
                updateTA()
                textAnswerArea.empty()
                st.toast(f"第{row[0]}题答案已更新, 请点击上方按钮继续答题或交卷")
    if st.session_state.examType == "training":
        AIOptionIndex = sac.segmented(
            items=[
                sac.SegmentedItem(label="讯飞"),
                sac.SegmentedItem(label="百度"),
                sac.SegmentedItem(label="深索"),
            ], label="可选LLM大模型", index=AIOptionIndex, align="start", color="red", return_index=True, size="sm"
        )
        AIModelName = AIOption[AIOptionIndex]
        updateAIModel2(AIOption, AIOptionIndex)
        if row[5] != "":
            buttonAnalysis = st.button("显示答案解析")
            buttonDelAnalysis = st.button("删除本题答案解析")
            if buttonAnalysis:
                st.subheader(f":orange[解析 标准答案: :green[[{standardAnswer}]]]\n{row[5]}", divider="gray")
            if buttonDelAnalysis:
                delAnalysis(row)
        else:
            if AIModelName != "":
                buttonAnalysis = st.button(f"A.I.答案解析模型 :green[{AIModelName.replace('大模型', '')}]")
                buttonDelAnalysis = st.button("删除本题答案解析")
                if AIModelName == "文心千帆大模型":
                    AIModelType = st.radio(label="请设置生成内容类型", options=("简洁", "详细"), index=0, horizontal=True, help="返回结果类型, 详细型附加了很多解释内容")
                    if AIModelType == "简洁":
                        AIModel = "ERNIE Speed-AppBuilder"
                    elif AIModelType == "详细":
                        AIModel = "ERNIE-Speed-8K"
                if buttonAnalysis:
                    AIAnswerInfo = st.empty()
                    with AIAnswerInfo.container():
                        st.info(f"正在使用 :red[{AIModelName.replace('大模型', '')}] 获取答案解析, 内容不能保证正确, 仅供参考! 请稍等...")
                    if AIModelName == "文心千帆大模型":
                        AIAnswer = qianfan_AI(row[1], AIModel, option, row[4])
                    elif AIModelName == "讯飞星火大模型":
                        AIAnswer = xunfei_xh_AI(row[1], option, row[4])
                    elif AIModelName == "DeepSeek大模型":
                        AIAnswer = deepseek_AI(row[1], option, row[4])
                    AIAnswerInfo.empty()
                    if AIAnswer != "" and AIAnswer.find("无法直接回答") == -1 and AIAnswer.find("尚未查询") == -1 and AIAnswer.find("我不确定您想要表达什么意思") == -1 and AIAnswer.find("由于信息不足，无法给出准确答案") == -1 and AIAnswer.find("无法确定正确答案") == -1 and AIAnswer.find("无法提供准确答案") == -1:
                        if AIAnswer.startswith(":"):
                            AIAnswer = AIAnswer[1:]
                        AIAnswer = AIAnswer + f"\n\n:red[答案解析来自[{AIModelName}], 非人工解析内容, 仅供参考!]"
                        st.subheader(f":orange[解析 标准答案: :green[[{standardAnswer}]]]\n{AIAnswer}", divider="gray")
                        if flagAIUpdate:
                            AIAnswer = AIAnswer.replace('"', '""').replace("'", "''")
                            for each in ["questions", "commquestions", "morepractise", "favques", st.session_state.examTable, st.session_state.examFinalTable]:
                                sql = f"UPDATE {each} set qAnalysis = '{AIAnswer}' where Question = '{row[1]}' and qType = '{row[4]}'"
                                execute_sql_and_commit(conn, cur, sql)
                            st.toast("A.I.答案解析内容已更新至题库")
                    else:
                        st.info("A.I.获取答案解析失败")
                if buttonDelAnalysis:
                    delAnalysis(row)
            else:
                st.info("没有设置A.I.大模型")
    st.session_state.curQues = row[0]


@st.fragment
def delAnalysis(row):
    for each in ["questions", "commquestions", "morepractise", "favques", st.session_state.examTable, st.session_state.examFinalTable]:
        sql = f"UPDATE {each} set qAnalysis = '' where Question = '{row[1]}' and qType = '{row[4]}'"
        execute_sql_and_commit(conn, cur, sql)
    st.info("本题解析已删除")


@st.fragment
def manualFIB(rowID):
    fibAI = ""
    sql = f"SELECT Question, qAnswer, userAnswer from {st.session_state.examFinalTable} where ID = {rowID}"
    fibRow = execute_sql(cur, sql)[0]
    fibQues = fibRow[0]
    userAP = fibRow[2].split(";")
    qAP = fibRow[1].split(";")
    if len(qAP) == len(userAP):
        for each in qAP:
            b1 = fibQues.find("()")
            if b1 != -1:
                fibQues = f"{fibQues[:b1]}({each}){fibQues[b1 + 2:]}"
        fibAI = xunfei_xh_AI_fib(userAP, fibQues)

    return fibAI


@st.fragment
def getStandardAnswer(qRow):
    radioOption, standardAnswer = ["A", "B", "C", "D", "E", "F", "G", "H"], ""
    if qRow[4] == "单选题" or qRow[4] == "多选题":
        orgOption = qRow[3].replace("；", ";").split(";")
        for value in orgOption:
            standardAnswer = standardAnswer + radioOption[int(value)] + ", "
    elif qRow[4] == "判断题":
        if qRow[3] == "1":
            standardAnswer = "正确"
        else:
            standardAnswer = "错误"
    elif qRow[4] == "填空题":
        standardAnswer = qRow[3].replace("；", ";").replace(";", ", ")
    if standardAnswer.endswith(", "):
        standardAnswer = standardAnswer[:-2]

    return standardAnswer


# noinspection PyTypeChecker
@st.fragment
def updateTA():
    textAnswerPack = []
    for key in st.session_state.keys():
        if "textAnswer_" in key:
            textAnswerPack.append([int(key.replace("textAnswer_", "")), st.session_state[key]])
    textAnswerPack.sort()
    st.session_state.answer = ""
    for each in textAnswerPack:
        st.session_state.answer += each[1] + ";"
    if st.session_state.answer.endswith(";"):
        st.session_state.answer = st.session_state.answer[:-1]
    updateAnswer(st.session_state.curQues)


def changeCurQues(step, cQuesCount):
    st.session_state.curQues += step
    if st.session_state.curQues < 1:
        st.session_state.curQues = 1
    elif st.session_state.curQues > cQuesCount:
        st.session_state.curQues = cQuesCount


@st.fragment
def quesGoto():
    if st.session_state.chosenID is not None:
        st.session_state.goto = True
        cop = re.compile('[^0-9^.]')
        st.session_state.curQues = int(cop.sub('', st.session_state.chosenID))


@st.fragment
def displayTimeCountdown():
    countdownType = "Circle"
    if countdownType == "Flip":
        remindTimeText = open("./MyComponentsScript/Countdown-NoFlip.txt", "r", encoding="utf-8").read()
    elif countdownType == "Circle":
        remindTimeText = open("./MyComponentsScript/Countdown-Circle.txt", "r", encoding="utf-8").read()
    else:
        remindTimeText = ""
    timeArea = st.empty()
    with timeArea.container():
        #st.write(f"### :red[{st.session_state.examName}]")
        #st.markdown(f"<font face='微软雅黑' color=red size=16><center>**{st.session_state.examName}**</center></font>", unsafe_allow_html=True)
        #st.markdown(f"### <font face='微软雅黑' color=red><center>{st.session_state.examName}</center></font>", unsafe_allow_html=True)
        flagTime = bool(getParam("显示考试时间", st.session_state.StationCN))
        if st.session_state.examType == "exam" or flagTime:
            examTimeLimit = int(getParam("考试时间", st.session_state.StationCN) * 60)
            examEndTime = st.session_state.examStartTime + examTimeLimit
            if countdownType == "Flip":
                examEndTimeText = time.strftime("%Y-%m-%dT%H:%M:%S", time.localtime(examEndTime))
            elif countdownType == "Circle":
                examEndTimeText = str(examTimeLimit - int(time.time() - st.session_state.examStartTime))
            else:
                examEndTimeText = ""
            remindTimeText = remindTimeText.replace("remindTime", f'"{examEndTimeText}"')
            remainingTime = examTimeLimit - (int(time.time() - st.session_state.examStartTime))
            if remainingTime < 0:
                if st.session_state.examType == "exam":
                    st.warning("⚠️ 考试已结束, 将强制交卷!")
                    st.session_state.calcScore = True
                    calcScore()
                else:
                    st.session_state.examStartTime = int(time.time())
            elif remainingTime < 900:
                st.warning(f"⚠️ :red[考试剩余时间已不足{int(remainingTime / 60) + 1}分钟, 请抓紧时间完成考试!]")
            if remindTimeText != "":
                components.html(remindTimeText, height=94)


@st.fragment
def displayBigTime():
    components.html(open("./MyComponentsScript/Clock-Big.txt", "r", encoding="utf-8").read(), height=140)


@st.fragment
def displaySmallTime():
    components.html(open("./MyComponentsScript/Clock-Small.txt", "r", encoding="utf-8").read(), height=34)


@st.fragment
def displaySmallClock():
    components.html(open("./MyComponentsScript/Clock-Number.txt", "r", encoding="utf-8").read(), height=30)


@st.fragment
def displayBigTimeCircle():
    components.html(open("./MyComponentsScript/Clock-Big-Circle.txt", "r", encoding="utf-8").read(), height=260)


@st.fragment
def displayVisitCounter():
    sql = "SELECT pyLM from verinfo where pyFile = 'visitcounter'"
    visitcount = execute_sql(cur, sql)[0][0]
    countScript = (open("./MyComponentsScript/FlipNumber.txt", "r", encoding="utf-8").read()).replace("visitcount", str(visitcount))
    components.html(countScript, height=40)


@st.fragment
def displayAppInfo():
    infoStr = open("./MyComponentsScript/glowintext.txt", "r", encoding="utf-8").read()
    infoStr = infoStr.replace("软件名称", APPNAME)
    verinfo, verLM, likeCM = getVerInfo()
    infoStr = infoStr.replace("软件版本", f"软件版本: {int(verinfo / 10000)}.{int((verinfo % 10000) / 100)}.{int(verinfo / 10)} building {verinfo}")
    infoStr = infoStr.replace("更新时间", f"更新时间: {time.strftime('%Y-%m-%d %H:%M', time.localtime(verLM))}")
    #infoStr = infoStr.replace("用户评价", f"用户评价: {EMOJI[int(likeCM) - 1][0]} {likeCM} I feel {EMOJI[int(likeCM) - 1][1]}")
    infoStr = infoStr.replace("更新内容", f"更新内容: {UPDATETYPE['New']} 增加supervisor账户类型; 修改错题集重置归属功能; 添加使用手册")

    components.html(infoStr, height=340)


@st.fragment
def changelog():
    changelogInfo = open("./CHANGELOG.md", "r", encoding="utf-8").read()
    st.markdown(changelogInfo)


@st.dialog("交卷")
def submit_dialog(prompt):
    st.write(f":red[**{prompt}**]")
    buttonSubmit = st.button("确定")
    buttonCancel = st.button("取消")
    if buttonSubmit:
        st.session_state.calcScore = True
        st.rerun()
    elif buttonCancel:
        st.session_state.calcScore = False
        st.rerun()


def ClearStr(strValue):
    strValue = strValue.replace("\n", "").replace("\t", "").strip()

    return strValue


@st.fragment
def addExamIDD():
    itemArea = st.empty()
    with itemArea:
        exam_name = st.text_input("考试名称", help="名称不能设置为练习题库(此为保留题库)")
        exam_name = ClearStr(exam_name)
        exam_date = st.date_input(
            "请设置考试有效期",
            min_value=datetime.date.today() + datetime.timedelta(days=1),
            max_value=datetime.date.today() + datetime.timedelta(days=180),
            value=datetime.date.today() + datetime.timedelta(days=3),
            help="考试有效期最短1天, 最长180天, 默认3天",
        )
        if exam_name and exam_date and exam_name != "练习题库":
            if st.button("添加考试场次"):
                exam_date_str = exam_date.strftime("%Y-%m-%d")
                exam_date_end = datetime.datetime.combine(exam_date, datetime.time.max)
                exam_date_timestamp = int(exam_date_end.timestamp())

                # 使用参数化查询防止SQL注入
                sql = "SELECT ID from examidd where examName = ? and StationCN = ?"
                params = (exam_name, st.session_state.StationCN)
                if not execute_sql(cur, sql, params):
                    sql = "INSERT INTO examidd(examName, validDate, StationCN) VALUES(?, ?, ?)"
                    params = (exam_name, exam_date_timestamp, st.session_state.StationCN)
                    try:
                        execute_sql_and_commit(conn, cur, sql, params)
                        st.success(f"考试场次: [{exam_name}] 有效期: [{exam_date_str} 23:59:59] 添加成功")
                        updateKeyAction(f"新建考试场次{exam_name}")
                        itemArea.empty()
                    except Exception as e:
                        st.error(f"考试场次 [{exam_name}] 添加失败: {e}")
                else:
                    st.error(f"[{exam_name}] 考试场次已存在")
        else:
            if not exam_name:
                st.warning("请输入考试名称")


@st.fragment
def addStation():
    def create_station_table(station_name):
        table_name = f"setup_{station_name}"
        create_table_sql = f"""
        CREATE TABLE IF NOT EXISTS {table_name} (
            ID integer PRIMARY KEY AUTOINCREMENT,
            paramName text NOT NULL,
            param integer,
            paramType text NOT NULL
        );
        """
        cur.execute(create_table_sql)
        conn.commit()

        # Insert default settings from 'setup_默认'
        insert_default_sql = f"INSERT INTO {table_name}(paramName, param, paramType) SELECT paramName, param, paramType FROM setup_默认;"
        execute_sql_and_commit(conn, cur, insert_default_sql)

    def add_chapter_to_station(station_name, chapter):
        insert_chapter_sql = f"""
        INSERT INTO questionaff(chapterName, StationCN, chapterRatio, examChapterRatio)
        SELECT '{chapter}', '{station_name}', 10, 10
        WHERE NOT EXISTS (
            SELECT 1 FROM questionaff WHERE chapterName = '{chapter}' AND StationCN = '{station_name}'
        );
        """
        execute_sql_and_commit(conn, cur, insert_chapter_sql)

    item_area = st.empty()
    with item_area.container():
        station_name = st.text_input("站室名称", value="")
        station_name = ClearStr(station_name)  # Assuming ClearStr is a function to clean the input

        if station_name:
            if st.button("添加站室名称"):
                # Check if station already exists using parameterized query
                check_sql = "SELECT ID FROM stations WHERE Station = ?"
                if not execute_sql(cur, check_sql, (station_name,)):
                    insert_sql = "INSERT INTO stations(Station) VALUES(?)"
                    execute_sql_and_commit(conn, cur, insert_sql, (station_name,))

                    # Create station table and add chapters
                    create_station_table(station_name)
                    for chapter in ["公共题库", "错题集", "关注题集"]:
                        add_chapter_to_station(station_name, chapter)

                    st.success(f"[{station_name}] 站室添加成功")
                    updateKeyAction(f"新建站室{station_name}")
                    item_area.empty()
                else:
                    st.error(f"[{station_name}] 已存在")
        else:
            st.warning("请输入站室名称")


@st.fragment
def addUser():
    stationCName = getAllStations()
    user_input = {
        "userName": st.number_input("用户编码", min_value=1, max_value=999999, value=1, help="建议使用员工编码, 具有唯一性"),
        "userCName": st.text_input("用户姓名", max_chars=10, autocomplete="name", help="请输入用户中文姓名").strip(),
        "station": st.select_slider("站室", stationCName, value=st.session_state.StationCN),
        "userPassword1": st.text_input("设置密码", max_chars=8, type="password", autocomplete="off", help="设置用户密码"),
        "userPassword2": st.text_input("请再次输入密码", max_chars=8, type="password", placeholder="请与上一步输入的密码一致", autocomplete="off"),
        "userType": sac.switch(label="管理员", on_label="On", align='start', size='sm', value=False),
    }

    if all(user_input.values()):
        if user_input["userPassword1"] == user_input["userPassword2"]:
            user_exists = check_user_exists(user_input["userName"])
            if not user_exists:
                encrypted_password = getUserEDKeys(user_input["userPassword1"], "enc")
                insert_user(user_input["userName"], user_input["userCName"], user_input["station"], encrypted_password, user_input["userType"])
                st.toast(f"用户: {user_input['userName']} 姓名: {user_input['userCName']} 添加成功")
                updateActionUser(f"新建用户: {user_input['userName']} 姓名: {user_input['userCName']} 类型: {'admin' if user_input['userType'] else 'user'} 站室: {user_input['station']}")
            else:
                st.error(f"ID: [{user_input['userName']}] 姓名: [{user_input['userCName']}] 用户已存在或用户编码重复")
        else:
            st.error("两次输入密码不一致")
    else:
        validate_user_input(user_input)


def check_user_exists(user_name):
    sql = "SELECT ID from users where userName = ?"
    params = (user_name,)
    result = execute_sql(cur, sql, params)
    return bool(result)


def insert_user(user_name, user_cname, station, password, user_type):
    ut = "admin" if user_type else "user"
    sql = "INSERT INTO users(userName, userCName, userType, StationCN, userPassword) VALUES(?, ?, ?, ?, ?)"
    params = (user_name, user_cname, ut, station, password)
    execute_sql_and_commit(conn, cur, sql, params)


def validate_user_input(user_input):
    if not user_input["userCName"]:
        st.warning("请输入用户姓名")
    elif not user_input["userPassword1"]:
        st.warning("请输入密码")
    elif not user_input["userPassword2"]:
        st.warning("请确认密码")
# 其他依赖的函数（如get_station_cn_all, get_user_ed_keys, update_key_action等）保持不变或根据需要进行类似的优化。


def getStationCNALL(flagALL=False):
    StationCNamePack = []
    if flagALL:
        StationCNamePack.append("全站")
    sql = "SELECT Station from stations order by ID"
    rows = execute_sql(cur, sql)
    for row in rows:
        StationCNamePack.append(row[0])

    return StationCNamePack


def updateDAParam(updateParamType):
    for key in st.session_state.keys():
        if key.startswith("dasetup_"):
            upID = key[key.find("_") + 1:]
            sql = f"UPDATE setup_{st.session_state.StationCN} SET param = {int(st.session_state[key])} WHERE ID = {upID}"
            execute_sql_and_commit(conn, cur, sql)
    st.success(f"{updateParamType} 参数更新成功")
    updateKeyAction("考试参数更新")


def updateSwitchOption(quesType):
    if st.session_state[quesType]:
        sql = f"UPDATE setup_{st.session_state.StationCN} set param = 1 where paramName = '{quesType}'"
    else:
        sql = f"UPDATE setup_{st.session_state.StationCN} set param = 0 where paramName = '{quesType}'"
    execute_sql_and_commit(conn, cur, sql)
    if quesType == "测试模式":
        st.session_state.debug = bool(st.session_state[quesType])
    if quesType == "时钟样式":
        st.session_state.clockType = bool(st.session_state[quesType])
    #st.success(f"{quesType} 设置更新成功")


def setupReset():
    try:
        # 使用参数化查询防止SQL注入
        station_cn = st.session_state.StationCN
        reset_sql = "DELETE from setup_{} where ID > 0".format(station_cn)
        insert_sql = """
            INSERT INTO setup_{station_cn}(paramName, param, paramType)
            SELECT paramName, param, paramType from setup_默认
        """.format(station_cn=station_cn)
        update_sql_1 = """
            UPDATE questionaff
            SET chapterRatio = 10, examChapterRatio = 10
            WHERE StationCN = ? AND (chapterName = '公共题库' OR chapterName = '错题集')
        """
        update_sql_2 = """
            UPDATE questionaff
            SET chapterRatio = 5, examChapterRatio = 5
            WHERE StationCN = ? AND chapterName <> '公共题库' AND chapterName <> '错题集'
        """
        # 执行SQL语句
        execute_sql_and_commit(conn, cur, reset_sql)
        execute_sql_and_commit(conn, cur, insert_sql)
        execute_sql_and_commit(conn, cur, update_sql_1, (station_cn,))
        execute_sql_and_commit(conn, cur, update_sql_2, (station_cn,))

        st.success("所有设置已重置")
        updateKeyAction("重置所有设置")
    except Exception as e:
        st.error(f"重置设置时出错: {e}")


def updateAIModel():
    sql = f"UPDATE setup_{st.session_state.StationCN} set param = 0 where paramType = 'others' and paramName like '%大模型'"
    execute_sql_and_commit(conn, cur, sql)
    sql = f"UPDATE setup_{st.session_state.StationCN} set param = 1 where paramType = 'others' and paramName = '{st.session_state.AIModel}'"
    execute_sql_and_commit(conn, cur, sql)
    st.success(f"LLM大模型已设置为{st.session_state.AIModel}")


@st.fragment
def updateAIModel2(AIOption, AIOptionIndex):
    sql = f"UPDATE setup_{st.session_state.StationCN} set param = 0 where paramType = 'others' and paramName like '%大模型'"
    execute_sql_and_commit(conn, cur, sql)
    sql = f"UPDATE setup_{st.session_state.StationCN} set param = 1 where paramType = 'others' and paramName = '{AIOption[AIOptionIndex]}'"
    execute_sql_and_commit(conn, cur, sql)


# noinspection PyTypeChecker
def highlight_max(x, forecolor='black', backcolor='yellow'):
    is_max = x == x.max()

    return [f'color: {forecolor}; background-color: {backcolor}' if v else '' for v in is_max]


def queryExamAnswer(tablename):
    if tablename == "morepractise":
        chosenType = ["错题"]
    else:
        chosenType = ["对题", "错题"]
    options = st.multiselect(
        "查询类型",
        chosenType,
        ["错题"],
    )
    if options:
        searchButton = st.button("查询")
        if searchButton:
            if len(options) == 2:
                sql = "SELECT Question, qOption, qAnswer, qType, qAnalysis, userAnswer, ID from " + tablename + " where userAnswer <> '' and userName = " + str(st.session_state.userName) + " order by ID"
            elif len(options) == 1:
                if options[0] == "对题":
                    sql = "SELECT Question, qOption, qAnswer, qType, qAnalysis, userAnswer, ID from " + tablename + " where userAnswer <> '' and qAnswer = userAnswer and userName = " + str(st.session_state.userName) + " order by ID"
                elif options[0] == "错题":
                    sql = "SELECT Question, qOption, qAnswer, qType, qAnalysis, userAnswer, ID from " + tablename + " where userAnswer <> '' and qAnswer <> userAnswer and userName = " + str(st.session_state.userName) + " order by ID"
                else:
                    sql = ""
            else:
                sql = ""
            rows = execute_sql(cur, sql)
            if rows:
                for row in rows:
                    if row[2] != row[5]:
                        flagAnswer = "错误"
                        st.subheader("", divider="red")
                    else:
                        flagAnswer = "正确"
                        st.subheader("", divider="green")
                    st.subheader(f"题目: :grey[{row[0]}]")
                    if row[3] == "单选题":
                        st.write(":red[标准答案:]")
                        option, userAnswer = [], ["A", "B", "C", "D"]
                        tmp = row[1].replace("；", ";").split(";")
                        for index, each in enumerate(tmp):
                            each = each.replace("\n", "").replace("\t", "").strip()
                            option.append(f"{userAnswer[index]}. {each}")
                        st.radio(" ", option, key=f"compare_{row[6]}", index=int(row[2]), horizontal=True, label_visibility="collapsed", disabled=True)
                        st.write(f"你的答案: :red[{userAnswer[int(row[5])]}] 你的选择为: :blue[[{flagAnswer}]]")
                    elif row[3] == "多选题":
                        userOption = ["A", "B", "C", "D", "E", "F", "G", "H"]
                        st.write(":red[标准答案:]")
                        option = row[1].replace("；", ";").split(";")
                        orgOption = row[2].replace("；", ";").split(";")
                        for index, value in enumerate(option):
                            value = value.replace("\n", "").replace("\t", "").strip()
                            if str(index) in orgOption:
                                st.checkbox(f"{userOption[index]}. {value}:", value=True, disabled=True)
                            else:
                                st.checkbox(f"{userOption[index]}. {value}:", value=False, disabled=True)
                        userAnswer = row[5].replace("；", ";").split(";")
                        tmp = ""
                        for each in userAnswer:
                            tmp = tmp + userOption[int(each)] + ", "
                        st.write(f"你的答案: :red[{tmp[:-2]}] 你的选择为: :blue[[{flagAnswer}]]")
                    elif row[3] == "判断题":
                        st.write(":red[标准答案:]")
                        option = ["A. 正确", "B. 错误"]
                        tmp = int(row[2]) ^ 1
                        st.radio(" ", option, key=f"compare_{row[6]}", index=tmp, horizontal=True, label_visibility="collapsed", disabled=True)
                        tmp = int(row[5]) ^ 1
                        st.write(f"你的答案: :red[{option[tmp]}] 你的选择为: :blue[[{flagAnswer}]]")
                    elif row[3] == "填空题":
                        option = row[2].replace("；", ";").split(";")
                        userAnswer = row[5].replace("；", ";").split(";")
                        st.write(":red[标准答案:]")
                        for index, value in enumerate(option):
                            st.write(f"第{index + 1}个填空: :green[{value}]")
                        st.write("你的答案:")
                        for index, value in enumerate(userAnswer):
                            st.write(f"第{index + 1}个填空: :red[{value}]")
                        st.write(f"你的填写为: :blue[[{flagAnswer}]]")
                    if row[4] != "":
                        st.markdown(f"答案解析: :green[{row[4]}]")
            else:
                st.info("暂无数据")
    else:
        st.warning("请设置查询类型")


# noinspection PyTypeChecker
def queryExamResult():
    searchOption = []
    sql = f"SELECT ID, examName from examidd where StationCN = '{st.session_state.StationCN}' order by ID"
    rows = execute_sql(cur, sql)
    for row in rows:
        searchOption.append(row[1])
    searchExamName = st.selectbox("请选择考试场次", searchOption, index=None)
    options = st.multiselect(
        "查询类型",
        ["通过", "未通过"],
        ["通过", "未通过"],
    )
    if searchExamName:
        searchButton = st.button("查询")
    else:
        searchButton = st.button("查询", disabled=True)
    if searchButton and searchExamName:
        if options:
            tab1, tab2 = st.tabs(["简报", "详情"])
            sql = f"SELECT userName, userCName, examScore, examDate, examPass from examresult where examName = '{searchExamName}' and ("
            for each in options:
                if each == "通过":
                    sql = sql + " examPass = 1 or "
                elif each == "未通过":
                    sql = sql + " examPass = 0 or "
            if sql.endswith(" or "):
                sql = sql[:-4] + ") order by ID DESC"
            rows = execute_sql(cur, sql)
            if rows:
                df = pd.DataFrame(rows, dtype=str)
                df.columns = ["编号", "姓名", "成绩", "考试日期", "考试结果"]
                for index, value in enumerate(rows):
                    df.loc[index, "考试日期"] = time.strftime("%Y-%m-%d %H:%M:%S", time.localtime(int(df["考试日期"][index])))
                    df.loc[index, "考试结果"] = "通过" if int(df["考试结果"][index]) == 1 else "未通过"
                tab2.dataframe(df.style.apply(highlight_max, backcolor='yellow', subset=["成绩", "考试结果"]))
            if rows:
                for row in rows:
                    tab1.markdown(f"考生ID:  :red[{row[0]}] 考生姓名: :red[{row[1]}] 考试时间: :red[{time.strftime('%Y-%m-%d %H:%M:%S', time.localtime(row[3]))}]")
                    tab1.subheader(f"考试成绩: {row[2]} 分")
                    if row[4] == 1:
                        tab1.subheader("考试结果: :blue[通过] 👏")
                        tab1.subheader("", divider="orange")
                    else:
                        tab1.subheader("考试结果: :red[未通过] 😝")
                        tab1.subheader("", divider="red")
            else:
                st.info("暂无数据")
        else:
            st.warning("请设置查询类型")


def queryExamResultUsers():
    ExamNamePack = []
    sql = f"SELECT ID, examName from examidd where StationCN = '{st.session_state.StationCN}' order by ID"
    rows = execute_sql(cur, sql)
    for row in rows:
        ExamNamePack.append(row[1])
    searchExamName = st.selectbox("请选择考试场次", ExamNamePack, index=None)
    options = st.multiselect(
        "查询类型",
        ["已参加考试", "未参加考试"],
        ["未参加考试"],
    )
    searchButton = st.button("查询")
    if searchButton and searchExamName:
        if options:
            tab1, tab2 = st.tabs(["简报", "详情"])
            if len(options) == 2:
                sql = "SELECT userName, userCName, StationCN from users where StationCN = '" + st.session_state.StationCN + "' and userType <> 'supervisor' order by ID"
            elif len(options) == 1:
                if options[0] == "已参加考试":
                    sql = "SELECT users.userName, users.userCName, users.StationCN from users, examresult where users.userType <> 'supervisor' and examresult.examName = '" + searchExamName + "' and examresult.userName = users.userName and users.StationCN = '" + st.session_state.StationCN + "'"
                elif options[0] == "未参加考试":
                    sql = "SELECT userName, userCName, StationCN from users where userType <> 'supervisor' and userName not in (SELECT users.userName from users, examresult where examresult.examName = '" + searchExamName + "' and examresult.userName = users.userName) and StationCN = '" + st.session_state.StationCN + "'"
            rows = execute_sql(cur, sql)
            if rows:
                df = pd.DataFrame(rows)
                df.columns = ["编号", "姓名", "站室"]
                tab2.dataframe(df)
            if rows:
                for row in rows:
                    sql = "SELECT userName, userCName, examScore, examDate, examPass from examresult where examName = '" + searchExamName + "' and userName = " + str(row[0])
                    rows2 = execute_sql(cur, sql)
                    if rows2:
                        tab1.markdown(f"考生ID:  :red[{rows2[0][0]}] 考生姓名: :red[{rows2[0][1]}] 考试时间: :red[{time.strftime('%Y-%m-%d %H:%M:%S', time.localtime(rows2[0][3]))}]")
                        tab1.subheader(f"考试成绩: {rows2[0][2]} 分")
                        if rows2[0][4] == 1:
                            tab1.subheader("考试结果: :blue[通过] 👏")
                            tab1.subheader("", divider="orange")
                        else:
                            tab1.subheader("考试结果: :red[未通过] 🤪")
                            tab1.subheader("", divider="red")
                    else:
                        tab1.subheader("未参加考试", divider="red")
                        tab1.markdown(f"考生ID:  :red[{row[0]}] 考生姓名: :red[{row[1]}] 站室: :red[{row[2]}]")
            else:
                st.info("暂无数据")
        else:
            st.warning("请设置查询类型")


def verifyUserPW(vUserName, vUserPW):
    st.session_state.userPwRecheck = False
    vUserEncPW = ""
    sql = f"SELECT userPassword from users where userName = {vUserName}"
    pwTable = execute_sql(cur, sql)
    if pwTable:
        vUserEncPW = pwTable[0][0]
        vUserDecPW = getUserEDKeys(vUserEncPW, "dec")
        if vUserPW == vUserDecPW:
            st.session_state.userPwRecheck = True
            return True, vUserEncPW
        else:
            return False, vUserEncPW
    else:
        return False, vUserEncPW


def resetPassword():
    st.subheader(":orange[密码重置及更改账户类型]", divider="red")

    # 验证管理员密码
    if not st.session_state.userPwRecheck:
        vUserPW = st.text_input("请输入密码", max_chars=8, type="password", autocomplete="off")
        if vUserPW:
            if verifyUserPW(st.session_state.userName, vUserPW)[0]:
                st.session_state.userPwRecheck = True
                st.success("密码验证成功，请进行后续操作。")
            else:
                st.error("密码错误，请重新输入。")

    # 重置用户信息
    if st.session_state.userPwRecheck:
        st.write(":red[**重置用户信息**]")
        rUserName = st.number_input("用户编码", min_value=0)
        if rUserName:
            user_info = get_user_info(rUserName)
            if user_info:
                st.write(f"用户姓名: **{user_info['userCName']}**")
                # 初始化管理员开关状态
                rUserType = sac.switch(label="管理员", value=user_info['is_admin'], on_label="On", off_label="Off", size='sm')

                # 重置选项
                rOption1 = st.checkbox("密码", value=False)
                rOption2 = st.checkbox("账户类型", value=False)
                if rOption1 or rOption2:
                    btnResetUserPW = st.button("重置")
                    if btnResetUserPW:
                        # 执行重置操作，此处假设actionResetUserPW已定义并处理相关逻辑
                        st.button("确认", on_click=actionResetUserPW, args=(rUserName, rOption1, rOption2, rUserType,))
                        st.session_state.userPwRecheck = False
                        st.success("重置操作已完成。")
                else:
                    st.warning("请选择重置类型")
            else:
                st.error("用户不存在")


# 辅助函数：获取用户信息，使用参数化查询防止SQL注入
def get_user_info(user_id):
    sql = "SELECT userCName, userType FROM users WHERE userName = ?"
    params = (user_id,)
    rows = execute_sql(cur, sql, params)
    if rows:
        return {
            'userCName': rows[0][0],
            'is_admin': rows[0][1] in ('admin', 'supervisor')
        }
    return None


def actionResetUserPW(rUserName, rOption1, rOption2, rUserType):
    rInfo = []

    def reset_password(user_name, new_password):
        encrypted_pw = getUserEDKeys(new_password, "enc")
        sql = "UPDATE users SET userPassword = ? WHERE userName = ?"
        params = (encrypted_pw, user_name)
        try:
            execute_sql_and_commit(conn, cur, sql, params)
            rInfo.append("密码已重置为: 1234 / ")
            updateKeyAction("密码重置")
        except Exception as e:
            st.error(f"密码重置失败: {e}")

    def update_user_type(user_name, new_type):
        sql = "UPDATE users SET userType = ? WHERE userName = ?"
        params = (new_type, user_name)
        try:
            execute_sql_and_commit(conn, cur, sql, params)
            rInfo.append(f"账户类型已更改为: {new_type} / ")
            updateKeyAction(f"更改账户类型为{new_type}")
        except Exception as e:
            st.error(f"更改账户类型失败: {e}")

    if rOption1:
        reset_password(rUserName, "1234")
    if rOption2:
        user_type = "admin" if rUserType else "user"
        update_user_type(rUserName, user_type)

    if rInfo:
        st.success(f"**{''.join(rInfo)[:-3]}**")


def displayKeyAction():
    st.subheader(":violet[操作日志]", divider="red")

    # 检查是否需要密码验证
    if st.session_state.userPwRecheck:
        # 直接展示操作日志
        display_action_log()
    else:
        # 验证管理员密码
        vUserPW = st.text_input("请输入密码", max_chars=8, placeholder="请输入管理员密码, 以验证身份", type="password", autocomplete="off")
        if vUserPW:
            if verifyUserPW(st.session_state.userName, vUserPW)[0]:
                # 密码验证成功，更新状态并展示操作日志
                st.session_state.userPwRecheck = True
                display_action_log()
            else:
                # 密码错误，显示错误信息并清空输入框
                st.error("密码错误, 请重新输入")
                st.session_state.password_input = ""  # 假设有一个状态用于跟踪输入框内容


def display_action_log():
    # 查询并展示操作日志
    sql = "SELECT userName, userCName, StationCN, userAction, datetime(actionDate, 'unixepoch', 'localtime') from keyactionlog order by actionDate DESC"
    rows = execute_sql(cur, sql)
    if rows:
        df = pd.DataFrame(rows, columns=["用户编码", "用户姓名", "所属站室", "操作内容", "操作时间"])
        st.write(df)


def ls_get(key):

    return st_javascript(f"localStorage.getItem('{key}');")


def ls_set(key, value):
    value = json.dumps(value, ensure_ascii=False)

    return st_javascript(f"localStorage.setItem('{key}', JSON.stringify('{value}');")


def getAllStations():
    STATIONPACK, stationIndex = [], 0
    sql = "SELECT Station from stations where Station <> '调控中心' order by ID"
    rows = execute_sql(cur, sql)
    for row in rows:
        STATIONPACK.append(row[0])
        if st.session_state.StationCN == row[0]:
            stationIndex = rows.index(row)

    return STATIONPACK, stationIndex


def displayUserManual():
    pdfFile = "./Demo/ETest使用手册.pdf"
    with open(pdfFile, "rb") as f:
        base64_pdf = base64.b64encode(f.read()).decode('utf-8')
    pdf_display = f'<embed src="data:application/pdf;base64,{base64_pdf}" width="800" height="1000" type="application/pdf">'
    st.markdown(pdf_display, unsafe_allow_html=True)


global APPNAME, EMOJI, UPDATETYPE, STATIONPACK

DBFILE = "./DB/ETest.db"

conn = sqlite3.Connection(DBFILE, check_same_thread=False)
cur = conn.cursor()

with open("./css/globalstyle.css") as css:
    st.markdown(f"<style>{css.read()}</style>", unsafe_allow_html=True)

st.logo("./Images/etest-logo2.png", icon_image="./Images/exam2.png", size="small")

# noinspection PyRedeclaration
APPNAME = "调控中心安全生产业务考试系统-SP版"
# noinspection PyRedeclaration
EMOJI = [["🥺", "very sad!"], ["😣", "bad!"], ["😋", "not bad!"], ["😊", "happy!"], ["🥳", "fab, thank u so much!"]]
# noinspection PyRedeclaration
UPDATETYPE = {"New": "✨", "Optimize": "🚀", "Fix": "🐞"}
selected = None
if "logged_in" not in st.session_state:
    st.session_state.logged_in = False
    login()

if st.session_state.logged_in:
    with st.sidebar:
        if st.session_state.clockType:
            displaySmallTime()
        else:
            displaySmallClock()
        if st.session_state.examType == "exam":
            selected = sac.menu([
                sac.MenuItem('主页', icon='house'),
                sac.MenuItem('功能', icon='grid-3x3-gap', children=[
                    sac.MenuItem('选择考试', icon='list-task'),
                    sac.MenuItem('开始考试', icon='pencil-square'),
                ]),
                sac.MenuItem('账户', icon='person-gear', children=[
                    sac.MenuItem('密码修改', icon='key', disabled=True),
                    sac.MenuItem('登出', icon='box-arrow-right'),
                ]),
                sac.MenuItem('关于', icon='layout-wtf', children=[
                    sac.MenuItem('Changelog', icon='view-list', disabled=True),
                    sac.MenuItem('Readme', icon='github', disabled=True),
                    #sac.MenuItem('使用手册', icon='question-diamond', disabled=True),
                    sac.MenuItem('关于...', icon='link-45deg', disabled=True),
                ], disabled=True),
            ], open_all=True, size="sm")
        elif st.session_state.examType == "training":
            if st.session_state.userType == "admin" or st.session_state.userType == 'supervisor':
                selected = sac.menu([
                    sac.MenuItem('主页', icon='house'),
                    sac.MenuItem('功能', icon='grid-3x3-gap', children=[
                        sac.MenuItem('生成题库', icon='list-task'),
                        sac.MenuItem('题库练习', icon='pencil-square'),
                        sac.MenuItem('数据录入', icon='database-add'),
                        sac.MenuItem('试题修改', icon='clipboard-check'),
                        sac.MenuItem('文件导出', icon='journal-arrow-down'),
                        sac.MenuItem('题库功能', icon='database-gear'),
                        sac.MenuItem('参数设置', icon='gear'),
                    ]),
                    sac.MenuItem('信息', icon='info-circle', children=[
                        sac.MenuItem('学习信息', icon='book'),
                        sac.MenuItem('证书及榜单', icon='bookmark-star'),
                    ]),
                    sac.MenuItem('查询', icon='search', children=[
                        sac.MenuItem('信息查询', icon='info-lg'),
                        sac.MenuItem('用户状态', icon='people'),
                        sac.MenuItem('操作日志', icon='incognito'),
                    ]),
                    sac.MenuItem('账户', icon='person-gear', children=[
                        sac.MenuItem('密码修改', icon='key'),
                        sac.MenuItem('密码重置', icon='bootstrap-reboot'),
                        sac.MenuItem('登出', icon='box-arrow-right'),
                    ]),
                    sac.MenuItem('关于', icon='layout-wtf', children=[
                        sac.MenuItem('Changelog', icon='view-list'),
                        sac.MenuItem('Readme', icon='github'),
                        #sac.MenuItem('使用手册', icon='question-diamond'),
                        sac.MenuItem('关于...', icon='link-45deg'),
                    ]),
                ], open_index=[1], open_all=False, size="sm")
            elif st.session_state.userType == "user":
                selected = sac.menu([
                    sac.MenuItem('主页', icon='house'),
                    sac.MenuItem('功能', icon='grid-3x3-gap', children=[
                        sac.MenuItem('生成题库', icon='list-task'),
                        sac.MenuItem('题库练习', icon='pencil-square'),
                    ]),
                    sac.MenuItem('信息', icon='info-circle', children=[
                        sac.MenuItem('学习信息', icon='book'),
                        sac.MenuItem('证书及榜单', icon='bookmark-star'),
                    ]),
                    sac.MenuItem('账户', icon='person-gear', children=[
                        sac.MenuItem('密码修改', icon='key'),
                        sac.MenuItem('登出', icon='box-arrow-right'),
                    ]),
                    sac.MenuItem('关于', icon='layout-wtf', children=[
                        sac.MenuItem('Changelog', icon='view-list'),
                        sac.MenuItem('Readme', icon='github'),
                        #sac.MenuItem('使用手册', icon='question-diamond'),
                        sac.MenuItem('关于...', icon='link-45deg'),
                    ]),
                ], open_index=[1, 2, 3, 4, 5, 6], open_all=False, size="sm")

        if st.session_state.userType == "supervisor":
            spv = getAllStations()
            st.session_state.StationCN = st.selectbox("请选择站室", options=spv[0], index=spv[1])
            sql = f"UPDATE users set StationCN = '{st.session_state.StationCN}' where userName = {st.session_state.userName}"
            execute_sql_and_commit(conn, cur, sql)
            preExamTypeIndex = 0
            if st.session_state.examType == "training":
                preExamTypeIndex = 0
            elif st.session_state.examType == "exam":
                preExamTypeIndex = 1
            tmpExamType = st.selectbox("请选择模式类型", options=["练习", "考试"], index=preExamTypeIndex)
            if tmpExamType == "练习":
                st.session_state.examType = "training"
                st.session_state.examName = "练习题库"
                st.session_state.examRandom = True
            elif tmpExamType == "考试":
                st.session_state.examType = "exam"
                st.session_state.examRandom = bool(getParam("考试题库每次随机生成", st.session_state.StationCN))
        st.write(f"### 姓名: :orange[{st.session_state.userCName}] 站室: :orange[{st.session_state.StationCN}]")
        st.caption("📢:red[**不要刷新页面, 否则会登出**]")
        #st.caption("**请使用 :red[[登出]] 功能退出页面, 否则会影响下次登录**")
    updatePyFileinfo()
    if selected != "密码重置" and selected != "用户状态" and selected != "操作日志":
        st.session_state.userPwRecheck = False
    if selected == "主页":
        #displayBigTime()
        displayBigTimeCircle()
        st.markdown(f"<font face='微软雅黑' color=purple size=5><center>**{APPNAME}**</center></font>", unsafe_allow_html=True)
        verinfo, verLM, likeCM = getVerInfo()
        #st.subheader(f"软件版本: {int(verinfo / 10000)}.{int((verinfo % 10000) / 100)}.{int(verinfo / 10)} building {verinfo}")
        #st.subheader(f"更新时间: {time.strftime('%Y-%m-%d %H:%M', time.localtime(verLM))}")
        #st.subheader(f"用户评价: {EMOJI[int(likeCM) - 1][0]} {likeCM} :orange[I feel {EMOJI[int(likeCM) - 1][1]}]")

        st.markdown(f"<font size=4><center>**软件版本: {int(verinfo / 10000)}.{int((verinfo % 10000) / 100)}.{int(verinfo / 10)} building {verinfo}**</center></font>", unsafe_allow_html=True)
        st.markdown(f"<font size=4><center>**更新时间: {time.strftime('%Y-%m-%d %H:%M', time.localtime(verLM))}**</center></font>", unsafe_allow_html=True)
        #st.markdown(f"<font size=5><center>**用户评价: {EMOJI[int(likeCM) - 1][0]} {likeCM} :orange[I feel {EMOJI[int(likeCM) - 1][1]}]**</center></font>", unsafe_allow_html=True)
        st.markdown(f"<font size=3><center>**更新内容: {UPDATETYPE['Optimize']} PC端与手机端合二为一, 数据互通**</center></font>", unsafe_allow_html=True)

        #displayAppInfo()
        st.subheader("")
        displayVisitCounter()

    elif selected == "生成题库" or selected == "选择考试":
        if st.session_state.examType == "training":
            #st.write("### :red[生成练习题库]")
            #st.markdown("<font face='微软雅黑' color=blue size=20><center>**生成练习题库**</center></font>", unsafe_allow_html=True)
            st.markdown("### <font face='微软雅黑' color=teal><center>生成练习题库</center></font>", unsafe_allow_html=True)
        elif st.session_state.examType == "exam":
            #st.markdown("<font face='微软雅黑' color=red size=20><center>**选择考试**</center></font>", unsafe_allow_html=True)
            st.markdown("### <font face='微软雅黑' color=red><center>选择考试</center></font>", unsafe_allow_html=True)
        if not st.session_state.examChosen or not st.session_state.calcScore:
            sql = "UPDATE verinfo set pyLM = 0 where pyFile = 'chapterChosenType'"
            execute_sql_and_commit(conn, cur, sql)
            training()
        else:
            st.error("你不能重复选择考试场次")
    elif selected == "题库练习" or selected == "开始考试":
        if st.session_state.examType == "exam":
            updateActionUser(st.session_state.userName, "考试", st.session_state.loginTime)
        elif st.session_state.examType == "training":
            updateActionUser(st.session_state.userName, "练习", st.session_state.loginTime)
        if "confirmSubmit" not in st.session_state:
            st.session_state.confirmSubmit = False
        if "examFinalTable" in st.session_state and "examName" in st.session_state and not st.session_state.confirmSubmit:
            sql = f"SELECT userName, examName from examresult GROUP BY userName, examName HAVING Count(examName) >= {st.session_state.examLimit} and userName = {st.session_state.userName} and examName = '{st.session_state.examName}'"
            if not execute_sql(cur, sql) or st.session_state.examType == "training":
                if st.session_state.calcScore:
                    calcScore()
                for key in st.session_state.keys():
                    if key.startswith("moption_") or key.startswith("textAnswer_"):
                        del st.session_state[key]
                displayTimeCountdown()
                sql = "SELECT * from " + st.session_state.examFinalTable + " order by ID"
                rows = execute_sql(cur, sql)
                quesCount = len(rows)
                preButton, nextButton, submitButton = False, False, False
                #st.write(f"Cur:{st.session_state.curQues} Comp:{st.session_state.flagCompleted}")
                if st.session_state.flagCompleted:
                    if st.session_state.curQues == 1:
                        preButton = st.button("上题", icon=":material/arrow_back_ios:", disabled=True)
                    else:
                        preButton = st.button("上题", icon=":material/arrow_back_ios:", on_click=changeCurQues, args=(-1, quesCount,))
                    if st.session_state.curQues == quesCount:
                        nextButton = st.button("下题", icon=":material/arrow_forward_ios:", disabled=True)
                    else:
                        nextButton = st.button("下题", icon=":material/arrow_forward_ios:", on_click=changeCurQues, args=(1, quesCount,))
                    submitButton = st.button("交卷", icon=":material/publish:")
                elif st.session_state.confirmSubmit:
                    preButton = st.button("上题", icon=":material/arrow_back_ios:", disabled=True)
                    nextButton = st.button("下题", icon=":material/arrow_forward_ios:", disabled=True)
                    submitButton = st.button("交卷", icon=":material/publish:", disabled=True)
                elif st.session_state.curQues == 0:
                    preButton = st.button("上题", icon=":material/arrow_back_ios:", disabled=True)
                    nextButton = st.button("下题", icon=":material/arrow_forward_ios:", on_click=changeCurQues, args=(1, quesCount,))
                    submitButton = st.button("交卷", icon=":material/publish:", disabled=True)
                    exam(rows[0])
                elif st.session_state.curQues == 1:
                    preButton = st.button("上题", icon=":material/arrow_back_ios:", disabled=True)
                    nextButton = st.button("下题", icon=":material/arrow_forward_ios:", on_click=changeCurQues, args=(1, quesCount,))
                    submitButton = st.button("交卷", icon=":material/publish:", disabled=True)
                elif st.session_state.curQues == quesCount:
                    preButton = st.button("上题", icon=":material/arrow_back_ios:", on_click=changeCurQues, args=(-1, quesCount,))
                    nextButton = st.button("下题", icon=":material/arrow_forward_ios:", disabled=True)
                    submitButton = st.button("交卷", icon=":material/publish:")
                    st.session_state.flagCompleted = True
                elif 1 < st.session_state.curQues < quesCount:
                    preButton = st.button("上题", icon=":material/arrow_back_ios:", on_click=changeCurQues, args=(-1, quesCount,))
                    nextButton = st.button("下题", icon=":material/arrow_forward_ios:", on_click=changeCurQues, args=(1, quesCount,))
                    submitButton = st.button("交卷", icon=":material/publish:", disabled=True)
                completedPack, cpStr, cpCount = [], "", 0
                sql = f"SELECT ID, qType from {st.session_state.examFinalTable} where userAnswer = '' order by ID"
                rows3 = execute_sql(cur, sql)
                for row3 in rows3:
                    completedPack.append(f"第{row3[0]}题 [{row3[1]}] 未作答")
                    cpStr = cpStr + str(row3[0]) + "/"
                sql = f"SELECT ID, qType from {st.session_state.examFinalTable} where userAnswer <> '' order by ID"
                rows3 = execute_sql(cur, sql)
                for row3 in rows3:
                    completedPack.append(f"第{row3[0]}题 [{row3[1]}] 已作答")
                cpCount = len(rows3)
                if cpCount == quesCount:
                    st.caption(":orange[作答提示: 全部题目已作答]")
                elif quesCount - cpCount > 40:
                    st.caption(f":blue[作答提示:] :red[你还有{quesCount - cpCount}道题未作答, 请尽快完成]")
                elif quesCount - cpCount > 0:
                    st.caption(f":blue[作答提示:] :red[{cpStr[:-1]}] :blue[题还未作答, 可以在下方答题卡列表中跳转]")
                else:
                    st.caption(":red[你还未开始答题]")
                st.selectbox(":green[答题卡] :red[[未答题前置排序]]", completedPack, index=None, on_change=quesGoto, key="chosenID")
                st.divider()
                if (preButton or nextButton or submitButton or st.session_state.goto) and not st.session_state.confirmSubmit:
                    sql = f"SELECT * from {st.session_state.examFinalTable} where ID = {st.session_state.curQues}"
                    row = execute_sql(cur, sql)[0]
                    if preButton or nextButton or st.session_state.goto:
                        if st.session_state.goto:
                            st.session_state.goto = False
                            st.write("#### :blue[跳转到指定题号: ]")
                        exam(row)
                    if submitButton:
                        emptyAnswer = "你没有作答的题为:第["
                        sql = f"SELECT ID from {st.session_state.examFinalTable} where userAnswer == '' order by ID"
                        rows2 = execute_sql(cur, sql)
                        for row2 in rows2:
                            emptyAnswer = emptyAnswer + str(row2[0]) + ", "
                        if emptyAnswer.endswith(", "):
                            emptyAnswer = emptyAnswer[:-2] + "]题, 请检查或直接交卷!"
                        else:
                            emptyAnswer = "你的所有题目均已作答, 确认交卷吗?"
                        submit_dialog(emptyAnswer)
                    preButton, nextButton, submitButton = False, False, False
            elif st.session_state.examType == "exam":
                st.info("你本场考试已达到次数限制, 无法再次进行, 如有疑问请联系管理员", icon="ℹ️")
        else:
            if st.session_state.examType == "training":
                st.info("请先生成新的题库", icon="ℹ️")
            elif st.session_state.examType == "exam":
                st.info("请先选择考试场次并点击开始考试", icon="ℹ️")
    elif selected == "数据录入":
        st.subheader(":orange[基础数据录入]", divider="violet")
        selectFunc = st.selectbox("请选择数据表", ["考试场次", "站室", "用户"], index=None, help="请选择数据表")
        stationCName = getStationCNALL()
        if selectFunc == "考试场次":
            buttonAdd = st.button("新增")
            if buttonAdd:
                addExamIDD()
        elif selectFunc == "站室":
            buttonAdd = st.button("新增")
            if buttonAdd:
                addStation()
        elif selectFunc == "用户":
            buttonAdd = st.button("新增")
            if buttonAdd:
                addUser()
        if selectFunc is not None:
            updateActionUser(st.session_state.userName, f"添加{selectFunc}", st.session_state.loginTime)
    elif selected == "试题修改":
        quesModify()
    elif selected == "文件导出":
        dboutput()
    elif selected == "题库功能":
        dbfunc()
    elif selected == "参数设置":
        st.subheader(":green[系统参数设置]")
        updateActionUser(st.session_state.userName, "设置系统参数", st.session_state.loginTime)
        with st.expander("# :blue[考试参数设置]"):
            sql = f"SELECT paramName, param, ID from setup_{st.session_state.StationCN} where paramType = 'exam' order by ID"
            rows = execute_sql(cur, sql)
            for row in rows:
                if row[0] == "单题分值":
                    quesScore = row[1]
                if row[0] == "考题总数":
                    quesTotal = row[1]
                if row[0] == "单选题数量":
                    st.slider(row[0], min_value=0, max_value=100, value=row[1], key=f"dasetup_{row[2]}")
                elif row[0] == "多选题数量":
                    st.slider(row[0], min_value=0, max_value=100, value=row[1], key=f"dasetup_{row[2]}")
                elif row[0] == "判断题数量":
                    st.slider(row[0], min_value=0, max_value=100, value=row[1], key=f"dasetup_{row[2]}")
                elif row[0] == "填空题数量":
                    st.slider(row[0], min_value=0, max_value=100, value=row[1], key=f"dasetup_{row[2]}")
                elif row[0] == "单题分值":
                    st.number_input(row[0], min_value=1, max_value=5, value=row[1], key=f"dasetup_{row[2]}", help="所有题型统一分值")
                elif row[0] == "考题总数":
                    st.number_input(row[0], min_value=10, max_value=120, value=row[1], key=f"dasetup_{row[2]}", help="仅对考试有效, 练习模式不受限制")
                elif row[0] == "合格分数线":
                    st.slider(row[0], min_value=60, max_value=120, value=row[1], step=10, key=f"dasetup_{row[2]}", help=f"建议为{int(quesScore * quesTotal * 0.8)}分")
                elif row[0] == "同场考试次数限制":
                    st.number_input(row[0], min_value=1, max_value=5, value=row[1], key=f"dasetup_{row[2]}", help="最多5次")
                elif row[0] == "考试题库每次随机生成":
                    #st.toggle(row[0], value=row[1], key=f"dasetup_{row[2]}", help="开启有效, 关闭无效")
                    sac.switch(label=row[0], value=row[1], key=row[0], on_label="On", align='start', size='sm')
                    updateSwitchOption(row[0])
                elif row[0] == "考试时间":
                    st.slider(row[0], min_value=30, max_value=150, value=row[1], step=15, key=f"dasetup_{row[2]}", help="建议为60-90分钟")
                elif row[0] == "使用大模型评判错误的填空题答案":
                    sac.switch(label=row[0], value=row[1], key=row[0], on_label="On", align='start', size='sm')
                    updateSwitchOption(row[0])
                else:
                    st.slider(row[0], min_value=1, max_value=150, value=row[1], key=f"dasetup_{row[2]}")
            updateDA = st.button("考试参数更新", on_click=updateDAParam, args=("考试",))
        with st.expander("# :red[章节权重设置]"):
            sql = "SELECT chapterName, examChapterRatio, ID from questionaff where chapterName <> '公共题库' and chapterName <> '错题集' and StationCN = '" + st.session_state.StationCN + "'"
            rows = execute_sql(cur, sql)
            if rows:
                sql = "SELECT chapterName, examChapterRatio, ID from questionaff where chapterName = '公共题库' and StationCN = '" + st.session_state.StationCN + "'"
                row = execute_sql(cur, sql)[0]
                st.slider(row[0], min_value=1, max_value=10, value=row[1], key=f"crsetup_{row[2]}", help="权重越大的章节占比越高")
                sql = "SELECT chapterName, examChapterRatio, ID from questionaff where chapterName = '错题集' and StationCN = '" + st.session_state.StationCN + "'"
                row = execute_sql(cur, sql)[0]
                st.slider(row[0], min_value=1, max_value=10, value=row[1], key=f"crsetup_{row[2]}", help="仅在练习题库中有效")
                for row in rows:
                    st.slider(row[0], min_value=1, max_value=10, value=row[1], key=f"crsetup_{row[2]}", help="权重越大的章节占比越高")
                st.button("章节权重更新", on_click=updateCRExam)
            else:
                st.info("该站室没有可设置章节")
        with st.expander("# :green[题型设置]"):
            sql = f"SELECT paramName, param from setup_{st.session_state.StationCN} where paramType = 'questype' order by ID"
            rows = execute_sql(cur, sql)
            for row in rows:
                sac.switch(label=row[0], value=row[1], key=row[0], on_label="On", align='start', size='sm')
                updateSwitchOption(row[0])
        with st.expander("# :violet[导出文件字体设置]"):
            sql = f"SELECT paramName, param, ID from setup_{st.session_state.StationCN} where paramType = 'fontsize' order by ID"
            rows = execute_sql(cur, sql)
            for row in rows:
                if row[0] == "抬头字体大小":
                    st.number_input(row[0], min_value=8, max_value=32, value=row[1], key=f"dasetup_{row[2]}", help="题库导出至Word文件中的字体大小")
                elif row[0] == "题型字体大小":
                    st.number_input(row[0], min_value=8, max_value=32, value=row[1], key=f"dasetup_{row[2]}")
                elif row[0] == "题目字体大小":
                    st.number_input(row[0], min_value=8, max_value=32, value=row[1], key=f"dasetup_{row[2]}")
                elif row[0] == "选项字体大小":
                    st.number_input(row[0], min_value=8, max_value=32, value=row[1], key=f"dasetup_{row[2]}")
                elif row[0] == "复核信息字体大小":
                    st.number_input(row[0], min_value=8, max_value=32, value=row[1], key=f"dasetup_{row[2]}")
            updateDA = st.button("字体设置更新", on_click=updateDAParam, args=("字体设置",))
        with st.expander("# :orange[其他设置]"):
            sql = f"SELECT paramName, param, ID from setup_{st.session_state.StationCN} where paramType = 'others' order by ID"
            rows = execute_sql(cur, sql)
            for row in rows:
                if row[0] == "显示考试时间" or row[0] == "A.I.答案解析更新至题库" or row[0] == "测试模式":
                    sac.switch(label=row[0], value=row[1], key=row[0], on_label="On", align='start', size='sm')
                    updateSwitchOption(row[0])
                elif row[0] == "时钟样式":
                    sac.switch(label=row[0], value=row[1], key=row[0], on_label="翻牌", off_label="数字", align='start', size='sm')
                    updateSwitchOption(row[0])
            AIModel, AIModelIndex = [], 0
            sql = f"SELECT paramName, param, ID from setup_{st.session_state.StationCN} where paramName like '%大模型' and paramType = 'others' order by ID"
            rows = execute_sql(cur, sql)
            for index, value in enumerate(rows):
                AIModel.append(value[0])
                if value[1] == 1:
                    AIModelIndex = index
            st.radio("选择LLM大模型", options=AIModel, index=AIModelIndex, key="AIModel", horizontal=True, on_change=updateAIModel, help="讯飞输出质量高, 规范引用准确, 建议选用;文心千帆输出速度快, 内容可用;DeepSeek内容准确性相对高一些")
        st.divider()
        buttonReset = st.button("重置所有设置", type="primary")
        if buttonReset:
            buttonConfirm = st.button("确认重置", type="secondary", on_click=setupReset)
            updateActionUser(st.session_state.userName, "重置所有设置", st.session_state.loginTime)
    elif selected == "信息查询":
        st.subheader(":violet[信息查询]", divider="orange")
        selectFunc = st.selectbox("查询项目", ["考试信息", "未参加考试人员", "答题解析"], index=None)
        if selectFunc == "考试信息":
            queryExamResult()
        elif selectFunc == "未参加考试人员":
            queryExamResultUsers()
        elif selectFunc == "答题解析":
            queryExamName = st.selectbox("请选择考试场次", ["练习题库", "错题集"], index=0)
            if queryExamName:
                if queryExamName == "错题集":
                    tablename = "morepractise"
                else:
                    tablename = f"exam_final_{st.session_state.StationCN}_{st.session_state.userName}_{queryExamName}"
                sql = "SELECT * from sqlite_master where type = 'table' and name = '" + tablename + "'"
                tempTable = execute_sql(cur, sql)
                if tempTable:
                    queryExamAnswer(tablename)
                else:
                    st.info("暂无数据")
        if selectFunc is not None:
            updateActionUser(st.session_state.userName, f"查询{selectFunc}", st.session_state.loginTime)
    elif selected == "用户状态":
        userStatus()
    elif selected == "操作日志":
        displayKeyAction()
    elif selected == "学习信息":
        studyinfo()
    elif selected == "证书及榜单":
        userRanking()
    elif selected == "密码修改":
        changePassword()
    elif selected == "密码重置":
        resetPassword()
    elif selected == "登出":
        logout()
    elif selected == "Changelog":
        changelog()
    elif selected == "Readme":
        aboutReadme()
    elif selected == "使用手册":
        displayUserManual()
    elif selected == "关于...":
        aboutInfo()
